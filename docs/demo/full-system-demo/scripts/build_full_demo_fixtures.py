from __future__ import annotations

import argparse
import textwrap
from dataclasses import dataclass
from pathlib import Path

import fitz
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


REPO_ROOT = Path(__file__).resolve().parents[4]
DEFAULT_SOURCE = REPO_ROOT / "docs/demo/full-system-demo/source-contracts.md"
DEFAULT_OUTPUT_DIR = REPO_ROOT / "output/full-system-demo/fixtures"

DOCX_FIXTURE_IDS = {"MSA_V1", "MSA_V2", "SOW_V1", "SOW_V2"}
PDF_FIXTURE_ID = "SECURITY_ADDENDUM"


@dataclass(frozen=True)
class SourceFixture:
    fixture_id: str
    title: str
    lines: list[str]


def parse_source(source_path: Path) -> list[SourceFixture]:
    fixtures: list[SourceFixture] = []
    current_id: str | None = None
    current_title: str | None = None
    current_lines: list[str] = []

    for raw_line in source_path.read_text(encoding="utf-8").splitlines():
        if raw_line.startswith("## ") and " - " in raw_line:
            if current_id and current_title:
                fixtures.append(SourceFixture(current_id, current_title, _trim_blank_edges(current_lines)))
            header = raw_line.removeprefix("## ").strip()
            current_id, current_title = [part.strip() for part in header.split(" - ", 1)]
            current_lines = []
            continue
        if current_id:
            current_lines.append(raw_line)

    if current_id and current_title:
        fixtures.append(SourceFixture(current_id, current_title, _trim_blank_edges(current_lines)))

    if not fixtures:
        raise ValueError(f"No fixtures found in {source_path}")
    return fixtures


def build_all_fixtures(source_path: Path, output_dir: Path) -> list[Path]:
    output_dir.mkdir(parents=True, exist_ok=True)
    generated: list[Path] = []
    fixtures = {fixture.fixture_id: fixture for fixture in parse_source(source_path)}

    for fixture_id in ("MSA_V1", "MSA_V2", "SOW_V1", "SOW_V2"):
        generated.append(write_docx_fixture(fixtures[fixture_id], output_dir))

    pdf_fixture = fixtures[PDF_FIXTURE_ID]
    generated.append(write_text_pdf_fixture(pdf_fixture, output_dir))
    generated.append(write_scanned_pdf_fixture(pdf_fixture, output_dir))
    return generated


def write_docx_fixture(fixture: SourceFixture, output_dir: Path) -> Path:
    document = Document()
    _configure_document(document)

    title = document.add_paragraph()
    title.style = "Title"
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title_run = title.add_run(fixture.title)
    title_run.bold = True

    subtitle = document.add_paragraph()
    subtitle_run = subtitle.add_run("Redline full-system demo fixture")
    subtitle_run.italic = True
    subtitle_run.font.color.rgb = RGBColor(91, 104, 124)

    _add_metadata_block(document, fixture)
    _add_body_from_lines(document, fixture.lines)

    output_path = output_dir / f"redline-full-demo-{_slugify(fixture.fixture_id)}-{_slugify(fixture.title)}.docx"
    document.save(output_path)
    return output_path


def _configure_document(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10)
    normal.paragraph_format.space_after = Pt(6)

    for style_name, size, color in (
        ("Title", 18, RGBColor(25, 30, 39)),
        ("Heading 1", 13, RGBColor(25, 30, 39)),
        ("Heading 2", 11, RGBColor(48, 58, 76)),
    ):
        style = styles[style_name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(10)
        style.paragraph_format.space_after = Pt(4)


def _add_metadata_block(document: Document, fixture: SourceFixture) -> None:
    metadata = [
        ("Demo owner", "MedNova Clinics Group"),
        ("Counterparty", "Aster Cloud Solutions"),
        ("Fixture ID", fixture.fixture_id),
        ("Review focus", _review_focus(fixture.fixture_id)),
    ]
    for label, value in metadata:
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(1)
        label_run = paragraph.add_run(f"{label}: ")
        label_run.bold = True
        label_run.font.size = Pt(9)
        value_run = paragraph.add_run(value)
        value_run.font.size = Pt(9)
        value_run.font.color.rgb = RGBColor(65, 75, 94)
    spacer = document.add_paragraph()
    spacer.paragraph_format.space_after = Pt(8)


def _add_body_from_lines(document: Document, lines: list[str]) -> None:
    index = 0
    while index < len(lines):
        line = lines[index].rstrip()
        stripped = line.strip()
        if not stripped:
            index += 1
            continue
        if stripped.startswith("### "):
            document.add_heading(stripped.removeprefix("### ").strip(), level=1)
            index += 1
            continue
        if stripped.startswith("#### "):
            document.add_heading(stripped.removeprefix("#### ").strip(), level=2)
            index += 1
            continue
        if stripped.startswith("|"):
            table_lines: list[str] = []
            while index < len(lines) and lines[index].strip().startswith("|"):
                table_lines.append(lines[index].strip())
                index += 1
            _add_markdown_table(document, table_lines)
            continue
        if stripped.startswith("- "):
            paragraph = document.add_paragraph(style="List Bullet")
            paragraph.add_run(stripped.removeprefix("- ").strip())
            index += 1
            continue
        document.add_paragraph(stripped)
        index += 1


def _add_markdown_table(document: Document, table_lines: list[str]) -> None:
    rows = [
        [cell.strip() for cell in line.strip("|").split("|")]
        for line in table_lines
        if not set(line.replace("|", "").strip()) <= {"-", " "}
    ]
    if not rows:
        return
    label = document.add_paragraph()
    label.paragraph_format.space_after = Pt(2)
    label_run = label.add_run("Structured schedule")
    label_run.bold = True
    label_run.font.size = Pt(9)
    for row_index, row_values in enumerate(rows):
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.left_indent = Inches(0.18)
        paragraph.paragraph_format.space_after = Pt(2)
        run = paragraph.add_run(" | ".join(row_values))
        run.font.name = "Arial"
        run.font.size = Pt(8.5)
        run.bold = row_index == 0
    document.add_paragraph()


def write_text_pdf_fixture(fixture: SourceFixture, output_dir: Path) -> Path:
    output_path = output_dir / "redline-full-demo-security-addendum-text.pdf"
    pdf = fitz.open()
    current_page = pdf.new_page(width=595, height=842)
    y = 54
    y = _insert_pdf_line(current_page, fixture.title, x=54, y=y, size=15, bold=True)
    y = _insert_pdf_line(current_page, "Redline full-system demo text-layer PDF", x=54, y=y + 8, size=10)
    for line in _pdf_lines(fixture):
        if y > 780:
            current_page = pdf.new_page(width=595, height=842)
            y = 54
        size = 12 if line.startswith(("1.", "2.", "3.", "4.", "5.", "6.", "7.", "8.")) else 9.5
        y = _insert_pdf_line(current_page, line, x=54, y=y, size=size, bold=size == 12)
    pdf.save(output_path)
    pdf.close()
    return output_path


def write_scanned_pdf_fixture(fixture: SourceFixture, output_dir: Path) -> Path:
    output_path = output_dir / "redline-full-demo-security-addendum-scan.pdf"
    pages: list[Image.Image] = []
    font_regular = _load_font(24)
    font_heading = _load_font(30)
    page_width, page_height = 1654, 2339
    margin_x, margin_y = 150, 140

    current = Image.new("RGB", (page_width, page_height), "white")
    draw = ImageDraw.Draw(current)
    y = margin_y
    y = _draw_wrapped(draw, fixture.title, margin_x, y, font_heading, width=58, line_gap=12)
    y = _draw_wrapped(draw, "Scanned signature-copy demo fixture for OCR fallback.", margin_x, y + 20, font_regular, width=78, line_gap=8)

    for line in _pdf_lines(fixture):
        if y > page_height - 220:
            pages.append(current)
            current = Image.new("RGB", (page_width, page_height), "white")
            draw = ImageDraw.Draw(current)
            y = margin_y
        font = font_heading if line.startswith(("1.", "2.", "3.", "4.", "5.", "6.", "7.", "8.")) else font_regular
        width = 58 if font is font_heading else 82
        y = _draw_wrapped(draw, line, margin_x, y, font, width=width, line_gap=8)
        y += 10
    pages.append(current)

    first, *rest = pages
    first.save(output_path, "PDF", resolution=200.0, save_all=True, append_images=rest)
    return output_path


def _insert_pdf_line(page: fitz.Page, text: str, *, x: int, y: int, size: float, bold: bool = False) -> int:
    font = "helv"
    chunks = textwrap.wrap(text, width=92 if size < 11 else 62) or [""]
    for chunk in chunks:
        page.insert_text((x, y), chunk, fontsize=size, fontname=font)
        y += int(size + 6)
    return y + (4 if bold else 2)


def _pdf_lines(fixture: SourceFixture) -> list[str]:
    lines: list[str] = []
    for line in fixture.lines:
        stripped = line.strip()
        if not stripped:
            continue
        if stripped.startswith("### "):
            lines.append(stripped.removeprefix("### ").strip())
        elif stripped.startswith("|"):
            continue
        else:
            lines.append(stripped)
    return lines


def _draw_wrapped(draw: ImageDraw.ImageDraw, text: str, x: int, y: int, font: ImageFont.FreeTypeFont, *, width: int, line_gap: int) -> int:
    for chunk in textwrap.wrap(text, width=width) or [""]:
        draw.text((x, y), chunk, fill=(25, 30, 39), font=font)
        y += font.size + line_gap
    return y


def _load_font(size: int) -> ImageFont.FreeTypeFont:
    candidates = [
        Path("C:/Windows/Fonts/arial.ttf"),
        Path("C:/Windows/Fonts/calibri.ttf"),
        Path("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf"),
    ]
    for path in candidates:
        if path.exists():
            return ImageFont.truetype(str(path), size=size)
    return ImageFont.load_default()


def _review_focus(fixture_id: str) -> str:
    return {
        "MSA_V1": "Baseline legal position",
        "MSA_V2": "Vendor-favorable changes in confidentiality, security, IP, liability, termination",
        "SOW_V1": "Baseline commercial delivery terms",
        "SOW_V2": "Vendor-favorable changes in acceptance, payment, IP, change control",
    }.get(fixture_id, "PDF parser and OCR smoke")


def _slugify(value: str) -> str:
    result = []
    for char in value.lower():
        if char.isalnum():
            result.append(char)
        elif char in {" ", "-", "_"}:
            result.append("-")
    slug = "".join(result)
    while "--" in slug:
        slug = slug.replace("--", "-")
    return slug.strip("-")


def _trim_blank_edges(lines: list[str]) -> list[str]:
    start = 0
    end = len(lines)
    while start < end and not lines[start].strip():
        start += 1
    while end > start and not lines[end - 1].strip():
        end -= 1
    return lines[start:end]


def main() -> int:
    parser = argparse.ArgumentParser(description="Build realistic Redline full-system demo fixtures.")
    parser.add_argument("--source", type=Path, default=DEFAULT_SOURCE, help="Path to source-contracts.md")
    parser.add_argument("--out", type=Path, default=DEFAULT_OUTPUT_DIR, help="Output directory for generated fixtures")
    args = parser.parse_args()

    generated = build_all_fixtures(args.source, args.out)
    for path in generated:
        print(path)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
