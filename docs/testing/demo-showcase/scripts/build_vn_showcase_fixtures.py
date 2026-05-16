from __future__ import annotations

import argparse
from dataclasses import dataclass
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt


REPO_ROOT = Path(__file__).resolve().parents[4]
DEFAULT_SOURCE = REPO_ROOT / "docs/testing/demo-showcase/vn-sample-contract-notes.md"
DEFAULT_OUTPUT_DIR = REPO_ROOT / "output/demo-showcase/fixtures"


@dataclass(frozen=True)
class DraftFixture:
    title: str
    filename: str
    lines: list[str]


def parse_fixtures(source_path: Path) -> list[DraftFixture]:
    fixtures: list[DraftFixture] = []
    current_title: str | None = None
    current_lines: list[str] = []

    for raw_line in source_path.read_text(encoding="utf-8").splitlines():
        if raw_line.startswith("## ") and " - " in raw_line:
            if current_title is not None:
                fixtures.append(_build_fixture(current_title, current_lines))
            current_title = raw_line.removeprefix("## ").strip()
            current_lines = []
            continue
        if current_title is not None:
            current_lines.append(raw_line)

    if current_title is not None:
        fixtures.append(_build_fixture(current_title, current_lines))

    if not fixtures:
        raise ValueError(f"No VN showcase draft fixtures found in {source_path}")
    return fixtures


def _build_fixture(title: str, lines: list[str]) -> DraftFixture:
    filename = f"redline-vn-showcase-{_slugify_title(title)}.docx"
    return DraftFixture(title=title, filename=filename, lines=_trim_blank_edges(lines))


def _slugify_title(title: str) -> str:
    parts = []
    for char in title.lower():
        if char.isalnum():
            parts.append(char)
        elif char in {" ", "-"}:
            parts.append("-")
    slug = "".join(parts)
    while "--" in slug:
        slug = slug.replace("--", "-")
    return slug.strip("-")


def _trim_blank_edges(lines: list[str]) -> list[str]:
    start = 0
    end = len(lines)
    while start < end and not lines[start].strip():
        start += 1
    while end > start and not lines[end - 1].strip():
        end -= 1
    return lines[start:end]


def write_docx_fixture(fixture: DraftFixture, output_dir: Path) -> Path:
    document = Document()
    _configure_document(document)
    document.add_heading(fixture.title, level=0)

    for line in fixture.lines:
        stripped = line.strip()
        if not stripped:
            continue
        if stripped.startswith("### "):
            document.add_heading(stripped.removeprefix("### ").strip(), level=1)
        else:
            document.add_paragraph(stripped)

    output_dir.mkdir(parents=True, exist_ok=True)
    output_path = output_dir / fixture.filename
    document.save(output_path)
    return output_path


def _configure_document(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10)


def build_all_fixtures(source_path: Path, output_dir: Path) -> list[Path]:
    fixtures = parse_fixtures(source_path)
    return [write_docx_fixture(fixture, output_dir) for fixture in fixtures]


def main() -> int:
    parser = argparse.ArgumentParser(description="Build local DOCX fixtures for Redline VN demo showcase.")
    parser.add_argument("--source", type=Path, default=DEFAULT_SOURCE, help="Path to vn-sample-contract-notes.md")
    parser.add_argument("--out", type=Path, default=DEFAULT_OUTPUT_DIR, help="Output directory for generated DOCX files")
    args = parser.parse_args()

    generated_paths = build_all_fixtures(args.source, args.out)
    for path in generated_paths:
        print(path)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
