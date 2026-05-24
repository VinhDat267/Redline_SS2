from io import BytesIO
from pathlib import Path

from docx import Document as DocxDocument
from docx.shared import Inches, Pt, RGBColor
from sqlalchemy import select
from sqlalchemy.orm import Session

from app.core.config import settings
from app.models import Document, DocumentVersion, Project, ProjectMember, User
from app.services.upload_storage import store_bytes


WORKSPACE_PROJECT_NAME = "Redline Review Workspace"
LEGACY_PROJECT_NAMES = {WORKSPACE_PROJECT_NAME, "Redline SS2 Demo"}
WORKSPACE_PROJECT_DESCRIPTION = "Starter review workspace for live Redline project, document, parser, and compare flows."
DEMO_DOCX_CONTENT_TYPE = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
PARSED_STATUSES = {"parsed", "parsed_with_warnings"}
DEMO_DOCUMENT_SPECS = [
    {
        "title": "Master Services Agreement",
        "document_type": "MSA",
        "description": "Primary contract for the live Redline review workflow.",
        "versions": [
            {
                "version_label": "v1.1",
                "file_name": "msa-v1.1.docx",
                "parse_status": "parsed",
                "notes": "Starter baseline contract draft for review and compare.",
            },
            {
                "version_label": "v2.0",
                "file_name": "msa-v2.0.docx",
                "parse_status": "parsed",
                "notes": "Starter revised contract draft for review and compare.",
            },
        ],
    },
    {
        "title": "Statement of Work",
        "document_type": "SOW",
        "description": "Implementation statement of work for live Redline review and compare workflows.",
        "versions": [
            {
                "version_label": "v1.0",
                "file_name": "sow-v1.0.docx",
                "parse_status": "parsed",
                "notes": "Starter baseline statement of work for review and compare.",
            },
            {
                "version_label": "v2.0",
                "file_name": "sow-v2.0.docx",
                "parse_status": "parsed",
                "notes": "Starter revised statement of work for review and compare.",
            },
        ],
    },
    {
        "title": "Security Addendum",
        "document_type": "DPA",
        "description": "Secondary contract artifact that keeps the project detail page populated with real data.",
        "versions": [
            {
                "version_label": "v1.0",
                "file_name": "security-addendum-v1.0.docx",
                "parse_status": "parsed",
                "notes": "Seeded secondary contract draft for the project inventory.",
            }
        ],
    },
]


def _ensure_demo_upload(document_id: int, file_name: str, version_label: str) -> str:
    stored_path = f"demo/document-{document_id}/{version_label}-{file_name}"
    if Path(file_name).suffix.lower() == ".docx":
        return store_bytes(
            stored_path,
            _build_demo_docx_bytes(file_name=file_name, version_label=version_label),
            content_type=DEMO_DOCX_CONTENT_TYPE,
        )

    payload = (
        f"Starter workspace upload for document {document_id} / {version_label}\n"
        f"Original file: {file_name}\n"
    ).encode("utf-8")
    return store_bytes(stored_path, payload, content_type="text/plain; charset=utf-8")


def _get_section_lines(section_id: str) -> list[str]:
    source_path = Path(__file__).resolve().parents[1] / "assets/source-contracts.md"
    if not source_path.exists():
        source_path = Path(__file__).resolve().parents[4] / "docs/demo/full-system-demo/source-contracts.md"
    if not source_path.exists():
        return [
            f"### {section_id}",
            f"Placeholder content for {section_id} because source-contracts.md was not found."
        ]

    current_id = None
    lines = []
    for line in source_path.read_text(encoding="utf-8").splitlines():
        if line.startswith("## ") and " - " in line:
            header = line.removeprefix("## ").strip()
            item_id = [part.strip() for part in header.split(" - ", 1)][0]
            if current_id == section_id:
                break
            if item_id == section_id:
                current_id = item_id
            continue
        if current_id == section_id:
            lines.append(line)

    # Trim empty lines from beginning/end
    start = 0
    end = len(lines)
    while start < end and not lines[start].strip():
        start += 1
    while end > start and not lines[end - 1].strip():
        end -= 1
    return lines[start:end]


def _add_markdown_table(document: DocxDocument, table_lines: list[str]) -> None:
    rows = [
        [cell.strip() for cell in line.strip("|").split("|")]
        for line in table_lines
        if not set(line.replace("|", "").strip()) <= {"-", " "}
    ]
    if not rows:
        return
    column_count = max(len(row) for row in rows)
    table = document.add_table(rows=len(rows), cols=column_count)
    table.style = "Table Grid"
    for row_index, row_values in enumerate(rows):
        for column_index in range(column_count):
            cell = table.cell(row_index, column_index)
            cell.text = row_values[column_index] if column_index < len(row_values) else ""
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = "Arial"
                    run.font.size = Pt(8.5)
                    run.bold = row_index == 0
    document.add_paragraph()


def _add_body_from_lines(document: DocxDocument, lines: list[str]) -> None:
    index = 0
    while index < len(lines):
        line = lines[index].rstrip()
        stripped = line.strip()
        if not stripped:
            index += 1
            continue
        if stripped.startswith("### "):
            document.add_heading(stripped.removeprefix("### ").strip(), level=1)
            index += 1
            continue
        if stripped.startswith("#### "):
            document.add_heading(stripped.removeprefix("#### ").strip(), level=2)
            index += 1
            continue
        if stripped.startswith("|"):
            table_lines: list[str] = []
            while index < len(lines) and lines[index].strip().startswith("|"):
                table_lines.append(lines[index].strip())
                index += 1
            _add_markdown_table(document, table_lines)
            continue
        if stripped.startswith("- "):
            paragraph = document.add_paragraph(style="List Bullet")
            paragraph.add_run(stripped.removeprefix("- ").strip())
            index += 1
            continue
        document.add_paragraph(stripped)
        index += 1


def _build_demo_docx(file_name: str, version_label: str) -> DocxDocument:
    section_id = None
    fn_lower = file_name.lower()
    vl_lower = version_label.lower()

    if "msa" in fn_lower:
        if "2.0" in vl_lower or "v2" in vl_lower:
            section_id = "MSA_V2"
        else:
            section_id = "MSA_V1"
    elif "sow" in fn_lower:
        if "2.0" in vl_lower or "v2" in vl_lower:
            section_id = "SOW_V2"
        else:
            section_id = "SOW_V1"
    elif "security-addendum" in fn_lower:
        section_id = "SECURITY_ADDENDUM"

    if not section_id:
        section_id = "MSA_V1"

    lines = _get_section_lines(section_id)

    document = DocxDocument()

    section = document.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10)
    normal.paragraph_format.space_after = Pt(6)

    # Use title of the section if found
    title_text = "Redline Seed Contract"
    if section_id == "MSA_V1":
        title_text = "Master Services Agreement (v1.1)"
    elif section_id == "MSA_V2":
        title_text = "Master Services Agreement (v2.0)"
    elif section_id == "SOW_V1":
        title_text = "Statement of Work (v1.0)"
    elif section_id == "SOW_V2":
        title_text = "Statement of Work (v2.0)"
    elif section_id == "SECURITY_ADDENDUM":
        title_text = "Security and Data Processing Addendum (v1.0)"

    title = document.add_paragraph()
    title_run = title.add_run(title_text)
    title_run.bold = True
    title_run.font.size = Pt(18)

    subtitle = document.add_paragraph()
    subtitle_run = subtitle.add_run("Redline full-system demo seed document")
    subtitle_run.italic = True
    subtitle_run.font.color.rgb = RGBColor(91, 104, 124)

    _add_body_from_lines(document, lines)
    return document


def _build_demo_docx_bytes(*, file_name: str, version_label: str) -> bytes:
    buffer = BytesIO()
    _build_demo_docx(file_name, version_label).save(buffer)
    return buffer.getvalue()


def _write_demo_docx(file_path: Path, *, file_name: str, version_label: str) -> None:
    document = _build_demo_docx(file_name, version_label)
    document.save(file_path)


def seed_demo_workspace(session: Session, current_user: User) -> dict[str, object]:
    from app.seed import seed_demo_users
    demo_users = seed_demo_users(session)

    matching_projects = list(
        session.scalars(select(Project).where(Project.name.in_(LEGACY_PROJECT_NAMES)))
    )
    project = next(
        (candidate for candidate in matching_projects if candidate.name == WORKSPACE_PROJECT_NAME),
        matching_projects[0] if matching_projects else None,
    )
    if project is None:
        project = Project(name=WORKSPACE_PROJECT_NAME, description=WORKSPACE_PROJECT_DESCRIPTION)
        session.add(project)
        session.commit()
        session.refresh(project)
    else:
        project.name = WORKSPACE_PROJECT_NAME
        project.description = WORKSPACE_PROJECT_DESCRIPTION
        session.add(project)
        session.commit()
        session.refresh(project)

    member_specs: dict[int, str | None] = {
        current_user.id: "owner",
    }
    for index, demo_user in enumerate(demo_users):
        member_specs.setdefault(demo_user.id, "reviewer" if index else "editor")

    existing_member_user_ids = {
        member.user_id
        for member in session.scalars(
            select(ProjectMember).where(ProjectMember.project_id == project.id)
        )
    }

    for user_id, role in member_specs.items():
        if user_id not in existing_member_user_ids:
            session.add(ProjectMember(project_id=project.id, user_id=user_id, role=role))
    session.commit()

    existing_documents = {
        document.title: document
        for document in session.scalars(select(Document).where(Document.project_id == project.id))
    }

    documents_seeded = 0
    versions_seeded = 0

    for document_spec in DEMO_DOCUMENT_SPECS:
        document = existing_documents.get(document_spec["title"])
        if document is None:
            document = Document(
                project_id=project.id,
                title=document_spec["title"],
                document_type=document_spec["document_type"],
                description=document_spec["description"],
            )
            session.add(document)
            session.commit()
            session.refresh(document)
            existing_documents[document.title] = document
            documents_seeded += 1
        else:
            document.document_type = document_spec["document_type"]
            document.description = document_spec["description"]
            session.add(document)
            session.commit()
            session.refresh(document)

        existing_versions = {
            version.version_label: version
            for version in session.scalars(
                select(DocumentVersion).where(DocumentVersion.document_id == document.id)
            )
        }

        for version_spec in document_spec["versions"]:
            file_path = _ensure_demo_upload(
                document.id,
                version_spec["file_name"],
                version_spec["version_label"],
            )
            version = existing_versions.get(version_spec["version_label"])
            if version is None:
                version = DocumentVersion(
                    document_id=document.id,
                    version_label=version_spec["version_label"],
                    file_name=version_spec["file_name"],
                    file_path=file_path,
                    uploaded_by_user_id=current_user.id,
                    parse_status=version_spec["parse_status"],
                    notes=version_spec["notes"],
                )
                session.add(version)
                session.commit()
                session.refresh(version)
                versions_seeded += 1
            else:
                version.file_name = version_spec["file_name"]
                version.file_path = file_path
                version.notes = version_spec["notes"]
                session.add(version)
                session.commit()
                session.refresh(version)

            # Trigger parsing immediately if the version status is parsed
            should_parse = (
                version_spec["parse_status"] == "parsed"
                and (
                    version.active_parse_run_id is None
                    or version.parse_status not in PARSED_STATUSES
                )
            )
            if should_parse:
                from app.services import document_parser
                try:
                    document_parser.parse_document_version(session, version)
                except Exception as exc:
                    raise RuntimeError(f"Failed parsing demo seed document {version.file_name}: {exc}") from exc

    session.commit()
    session.refresh(project)

    return {
        "project": project,
        "documents_seeded": documents_seeded,
        "versions_seeded": versions_seeded,
    }
