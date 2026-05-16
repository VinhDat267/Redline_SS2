import hashlib
import json
import re
from collections import Counter
from dataclasses import dataclass, field
from pathlib import Path
from pathlib import PurePosixPath
from typing import Iterator
from zipfile import ZipFile

from docx import Document as DocxDocument
from docx.blkcntnr import BlockItemContainer
from docx.document import Document as DocxDocumentObject
from docx.oxml import parse_xml
from docx.oxml.ns import qn
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from lxml import etree
from docx.table import Table
from docx.text.paragraph import Paragraph
from sqlalchemy.orm import Session

from app.core.config import BACKEND_ROOT
from app.models import (
    DocumentBlock,
    DocumentParseRun,
    DocumentSurface,
    DocumentTable,
    DocumentTableCell,
    DocumentTableColumn,
    DocumentTableRow,
    DocumentVersion,
)
from app.models.mixins import utcnow
from app.services.document_parser_quality import (
    analyze_docx_parser_quality,
)


PARSER_VERSION = "v1"
_WHITESPACE_RE = re.compile(r"\s+")
_NUMBERED_HEADING_RE = re.compile(r"^(?P<number>\d+(?:\.\d+){0,5})\.?\s+\S")
_NUMBERED_LIST_RE = re.compile(r"^(?:\d+[.)]|[A-Za-z][.)])\s+\S")
_PARENTHETICAL_LIST_RE = re.compile(r"^\((?:\d+|[A-Za-z]|[ivxlcdm]+)\)\s+\S", re.IGNORECASE)
_LEGAL_ARTICLE_HEADING_RE = re.compile(
    r"^(?:article|section|clause)\s+(?:\d+(?:\.\d+){0,5}|[ivxlcdm]+)\b\s+\S",
    re.IGNORECASE,
)
_CODE_LIKE_ROW_RE = re.compile(r"^(?:[A-Z]{2,}[-_ ]?\d+|\d)")
_NON_ALNUM_RE = re.compile(r"[^a-z0-9]+")
_SUPPORTED_AUTO_FIELD_PLACEHOLDERS = {
    "PAGE": "__FIELD_PAGE__",
    "NUMPAGES": "__FIELD_NUMPAGES__",
    "SECTIONPAGES": "__FIELD_SECTIONPAGES__",
    "DATE": "__FIELD_DATE__",
    "PRINTDATE": "__FIELD_PRINTDATE__",
    "SAVEDATE": "__FIELD_SAVEDATE__",
    "TIME": "__FIELD_TIME__",
}
_WORD_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
_VML_NS = "urn:schemas-microsoft-com:vml"
_DOCUMENT_RELS_NS = "http://schemas.openxmlformats.org/package/2006/relationships"
_DOCUMENT_RELS_NSMAP = {"rel": _DOCUMENT_RELS_NS}
_WORD_NSMAP = {"w": _WORD_NS}
_NOTE_RELATIONSHIP_TYPES = {
    "footnote": "http://schemas.openxmlformats.org/officeDocument/2006/relationships/footnotes",
    "endnote": "http://schemas.openxmlformats.org/officeDocument/2006/relationships/endnotes",
}
_IGNORED_NOTE_TYPES = {"separator", "continuationSeparator", "continuationNotice"}
_UNSUPPORTED_VISIBLE_TEXT_CONTAINER_TAGS = {
    f"{{{_WORD_NS}}}drawing",
    f"{{{_WORD_NS}}}pict",
    f"{{{_WORD_NS}}}txbxContent",
    f"{{{_VML_NS}}}textbox",
}


class DocumentParseError(Exception):
    """Raised when a document version cannot be parsed into valid blocks."""


@dataclass(slots=True)
class ParsedBlockDraft:
    block_key: str
    block_type: str
    section_title: str | None
    heading_level: int | None
    order_index: int
    surface_order_index: int
    raw_content: str
    normalized_content: str


@dataclass(slots=True)
class ParsedTableCellDraft:
    column_key: str
    column_index: int
    cell_key: str
    raw_content: str
    normalized_content: str
    merge_token: str | None = None
    row_span: int = 1
    col_span: int = 1
    merge_origin_key: str | None = None


@dataclass(slots=True)
class ParsedTableColumnDraft:
    column_key: str
    column_index: int
    header_text: str
    normalized_header_text: str
    source_kind: str


@dataclass(slots=True)
class ParsedTableRowDraft:
    row_key: str
    row_index: int
    is_header_row: bool
    structured_row_json: str
    block_draft: ParsedBlockDraft
    cells: list[ParsedTableCellDraft] = field(default_factory=list)


@dataclass(slots=True)
class ParsedTableDraft:
    table_key: str
    section_title: str | None
    table_order_index: int
    header_strategy: str
    caption_text: str | None
    normalized_caption_text: str | None
    columns: list[ParsedTableColumnDraft] = field(default_factory=list)
    rows: list[ParsedTableRowDraft] = field(default_factory=list)


@dataclass(slots=True)
class ParsedSurfaceDraft:
    surface_type: str
    surface_key: str
    logical_order_index: int
    section_ref: str | None
    notes: str | None = None
    blocks: list[ParsedBlockDraft] = field(default_factory=list)
    tables: list[ParsedTableDraft] = field(default_factory=list)


@dataclass(slots=True)
class ParsedDocumentDraft:
    surfaces: list[ParsedSurfaceDraft]
    warnings: list[str]
    quality_report: object | None = None
    pdf_summary: object | None = None


@dataclass(slots=True)
class _ResolvedTableCellDraft:
    column_index: int
    cell_key: str
    raw_content: str
    normalized_content: str
    merge_token: str | None = None
    row_span: int = 1
    col_span: int = 1
    merge_origin_key: str | None = None


class _OOXMLStoryContainer(BlockItemContainer):
    def __init__(self, element, parent):
        super().__init__(element, parent)


def parse_document_version(
    session: Session, version: DocumentVersion
) -> DocumentVersion:
    previous_active_parse_run_id = version.active_parse_run_id
    parse_run = DocumentParseRun(
        document_version_id=version.id,
        parser_version=PARSER_VERSION,
        status="running",
    )
    session.add(parse_run)
    session.flush()

    try:
        parsed_document = _build_document_draft(version)
    except DocumentParseError as exc:
        parse_run.status = "failed"
        parse_run.error_message = str(exc)
        parse_run.completed_at = utcnow()
        parse_run.warning_count = 0
        parse_run.summary_json = json.dumps(_build_parse_summary_payload([]))
        version.parse_status = "failed"
        version.active_parse_run_id = previous_active_parse_run_id
        version.parsed_snapshot = None
        session.add_all([parse_run, version])
        session.commit()
        session.refresh(version)
        raise exc

    quality_report = getattr(parsed_document, "quality_report", None)
    if quality_report is None:
        quality_report = analyze_docx_parser_quality(
            _resolve_file_path(version),
            canonical_texts=_collect_canonical_texts(parsed_document),
            canonical_block_count=sum(len(surface.blocks) for surface in parsed_document.surfaces),
        )
    combined_warnings = _deduplicate_warnings(
        [*parsed_document.warnings, *getattr(quality_report, "warnings", [])]
    )
    if getattr(quality_report, "policy_result", None) == "fail":
        failure_message = (
            getattr(quality_report, "error_message", None)
            or getattr(quality_report, "failure_message", None)
            or "Parser coverage policy failed for this document"
        )
        parse_run.status = "failed"
        parse_run.error_message = failure_message
        parse_run.completed_at = utcnow()
        parse_run.warning_count = len(combined_warnings)
        parse_run.summary_json = json.dumps(
            _build_parse_summary_payload(
                combined_warnings,
                quality_report,
                pdf_summary=getattr(parsed_document, "pdf_summary", None),
            )
        )
        version.parse_status = "failed"
        version.active_parse_run_id = previous_active_parse_run_id
        version.parsed_snapshot = None
        session.add_all([parse_run, version])
        session.commit()
        session.refresh(version)
        raise DocumentParseError(failure_message)

    persisted_surfaces_by_key: dict[str, DocumentSurface] = {}
    persisted_blocks_by_key: dict[str, DocumentBlock] = {}
    from app.services import rag_service

    block_embedding_keys: list[str] = []
    block_embedding_texts: list[str] = []
    for surface_draft in parsed_document.surfaces:
        for draft_block in surface_draft.blocks:
            block_embedding_keys.append(draft_block.block_key)
            block_embedding_texts.append(
                rag_service.build_embedding_text(
                    block_type=draft_block.block_type,
                    section_title=draft_block.section_title,
                    content=draft_block.normalized_content or draft_block.raw_content,
                )
            )
    try:
        embedding_payloads = rag_service.build_text_embedding_payloads(block_embedding_texts)
    except rag_service.EmbeddingProviderError as exc:
        embedding_payloads = [
            (None, None, None)
            for _block_key in block_embedding_keys
        ]
        combined_warnings = _deduplicate_warnings(
            [*combined_warnings, f"RAG embeddings skipped: {exc}"]
        )

    embedding_payloads_by_block_key = dict(
        zip(block_embedding_keys, embedding_payloads, strict=True)
    )

    for surface_draft in parsed_document.surfaces:
        surface = DocumentSurface(
            parse_run_id=parse_run.id,
            surface_type=surface_draft.surface_type,
            surface_key=surface_draft.surface_key,
            logical_order_index=surface_draft.logical_order_index,
            section_ref=surface_draft.section_ref,
            notes=surface_draft.notes,
        )
        session.add(surface)
        session.flush()
        persisted_surfaces_by_key[surface.surface_key] = surface

        for draft_block in surface_draft.blocks:
            (
                embedding_provider,
                embedding_vector,
                embedding_vector_json,
            ) = embedding_payloads_by_block_key[draft_block.block_key]
            block = DocumentBlock(
                document_version_id=version.id,
                parse_run_id=parse_run.id,
                surface_id=surface.id,
                block_key=draft_block.block_key,
                block_type=draft_block.block_type,
                section_title=draft_block.section_title,
                heading_level=draft_block.heading_level,
                order_index=draft_block.order_index,
                surface_order_index=draft_block.surface_order_index,
                raw_content=draft_block.raw_content,
                normalized_content=draft_block.normalized_content,
                embedding_provider=embedding_provider,
                embedding_vector=embedding_vector,
                embedding_vector_json=embedding_vector_json,
                embedding_generated_at=utcnow() if embedding_provider is not None else None,
            )
            session.add(block)
            persisted_blocks_by_key[draft_block.block_key] = block

    session.flush()

    for surface_draft in parsed_document.surfaces:
        surface = persisted_surfaces_by_key[surface_draft.surface_key]
        for table_draft in surface_draft.tables:
            table = DocumentTable(
                document_version_id=version.id,
                parse_run_id=parse_run.id,
                surface_id=surface.id,
                table_key=table_draft.table_key,
                section_title=table_draft.section_title,
                table_order_index=table_draft.table_order_index,
                header_strategy=table_draft.header_strategy,
                caption_text=table_draft.caption_text,
                normalized_caption_text=table_draft.normalized_caption_text,
            )
            session.add(table)
            session.flush()

            column_ids_by_key: dict[str, int | None] = {}
            for column_draft in table_draft.columns:
                column = DocumentTableColumn(
                    table_id=table.id,
                    column_key=column_draft.column_key,
                    column_index=column_draft.column_index,
                    header_text=column_draft.header_text,
                    normalized_header_text=column_draft.normalized_header_text,
                    source_kind=column_draft.source_kind,
                )
                session.add(column)
                session.flush()
                column_ids_by_key[column.column_key] = column.id

            for row_draft in table_draft.rows:
                row = DocumentTableRow(
                    table_id=table.id,
                    document_block_id=persisted_blocks_by_key[row_draft.block_draft.block_key].id,
                    row_key=row_draft.row_key,
                    row_index=row_draft.row_index,
                    is_header_row=row_draft.is_header_row,
                    structured_row_json=row_draft.structured_row_json,
                )
                session.add(row)
                session.flush()

                for cell_draft in row_draft.cells:
                    session.add(
                        DocumentTableCell(
                            row_id=row.id,
                            column_id=column_ids_by_key.get(cell_draft.column_key),
                            cell_key=cell_draft.cell_key,
                            column_index=cell_draft.column_index,
                            raw_content=cell_draft.raw_content,
                            normalized_content=cell_draft.normalized_content,
                            row_span=cell_draft.row_span,
                            col_span=cell_draft.col_span,
                            merge_origin_key=cell_draft.merge_origin_key,
                        )
                    )

    warning_count = len(combined_warnings)
    parse_status = (
        "parsed_with_warnings"
        if warning_count or getattr(quality_report, "policy_result", None) == "warn"
        else "parsed"
    )
    parse_run.status = parse_status
    parse_run.completed_at = utcnow()
    parse_run.warning_count = warning_count
    parse_run.summary_json = json.dumps(
        _build_parse_summary_payload(
            combined_warnings,
            quality_report,
            pdf_summary=getattr(parsed_document, "pdf_summary", None),
        )
    )
    version.parse_status = parse_status
    version.active_parse_run_id = parse_run.id
    version.parsed_snapshot = json.dumps(
        _build_parsed_snapshot(
            version.id,
            parse_run.id,
            parsed_document,
            warnings=combined_warnings,
            quality_report=quality_report,
        )
    )
    session.add_all([parse_run, version])
    session.commit()
    session.refresh(version)
    return version


def _build_document_draft(version: DocumentVersion) -> ParsedDocumentDraft:
    file_path = _resolve_file_path(version)
    if not file_path.exists():
        raise DocumentParseError("Document file not found")
    file_suffix = file_path.suffix.lower()
    if file_suffix == ".pdf":
        from app.services.document_pdf_parser import build_pdf_document_draft

        try:
            return build_pdf_document_draft(file_path)
        except ValueError as exc:
            raise DocumentParseError(str(exc)) from exc

    if file_suffix != ".docx":
        raise DocumentParseError("Only .docx and .pdf files are supported")

    try:
        document = DocxDocument(file_path)
    except Exception as exc:  # pragma: no cover - library-specific failure surface
        raise DocumentParseError("Unable to open .docx file") from exc

    surfaces: list[ParsedSurfaceDraft] = []
    warnings: list[str] = []
    next_global_order_index = 0
    next_table_order_index = 0
    next_surface_logical_order_index = 0

    body_surface_draft, next_global_order_index, next_table_order_index, body_warnings = _build_surface_draft(
        document,
        surface_type="body",
        surface_key="body-main",
        logical_order_index=next_surface_logical_order_index,
        section_ref=None,
        next_global_order_index=next_global_order_index,
        next_table_order_index=next_table_order_index,
    )
    if body_surface_draft is not None:
        surfaces.append(body_surface_draft)
        next_surface_logical_order_index += 1
        warnings.extend(body_warnings)

    for section_index, surface_type, surface_variant, container in _iter_header_footer_containers(document):
        surface_draft, next_global_order_index, next_table_order_index, surface_warnings = _build_surface_draft(
            container,
            surface_type=surface_type,
            surface_key=f"{surface_type}-section-{section_index}-{surface_variant}",
            logical_order_index=next_surface_logical_order_index,
            section_ref=f"section-{section_index}",
            next_global_order_index=next_global_order_index,
            next_table_order_index=next_table_order_index,
        )
        if surface_draft is None:
            continue

        surfaces.append(surface_draft)
        next_surface_logical_order_index += 1
        warnings.extend(surface_warnings)

    for surface_type, note_id, container in _iter_note_containers(file_path, document):
        surface_draft, next_global_order_index, next_table_order_index, surface_warnings = _build_surface_draft(
            container,
            surface_type=surface_type,
            surface_key=f"{surface_type}-{note_id}",
            logical_order_index=next_surface_logical_order_index,
            section_ref=None,
            next_global_order_index=next_global_order_index,
            next_table_order_index=next_table_order_index,
        )
        if surface_draft is None:
            continue

        surfaces.append(surface_draft)
        next_surface_logical_order_index += 1
        warnings.extend(surface_warnings)

    total_blocks = sum(len(surface.blocks) for surface in surfaces)
    if total_blocks == 0:
        raise DocumentParseError(
            "Parser produced no valid body blocks or supported document surfaces"
        )

    return ParsedDocumentDraft(surfaces=surfaces, warnings=warnings)


def normalize_content(raw_content: str) -> str:
    return _WHITESPACE_RE.sub(" ", raw_content).strip()


def classify_paragraph(
    paragraph: Paragraph, normalized_content: str
) -> tuple[str, int | None]:
    style_name = _get_style_name(paragraph)
    heading_level = _extract_heading_level_from_style(style_name)
    if heading_level is not None:
        return "heading", heading_level

    if _has_explicit_list_metadata(paragraph, style_name):
        return "list_item", None

    fallback_heading_level = _extract_heading_level_from_numbered_heading(
        normalized_content
    )
    if fallback_heading_level is not None:
        return "heading", fallback_heading_level

    fallback_heading_level = _extract_heading_level_from_legal_heading(
        normalized_content
    )
    if fallback_heading_level is not None:
        return "heading", fallback_heading_level

    if _has_marker_only_list(normalized_content):
        return "list_item", None

    return "paragraph", None


def build_block_key(order_index: int, block_type: str, normalized_content: str) -> str:
    content_hash = hashlib.sha1(normalized_content.encode("utf-8")).hexdigest()[:10]
    return f"blk-{order_index:04d}-{block_type}-{content_hash}"


def _collect_canonical_texts(parsed_document: ParsedDocumentDraft) -> list[str]:
    return [
        block.normalized_content or block.raw_content
        for surface in parsed_document.surfaces
        for block in surface.blocks
        if block.normalized_content or block.raw_content
    ]


def _deduplicate_warnings(warnings: list[str]) -> list[str]:
    return list(dict.fromkeys(warning for warning in warnings if warning))


def _build_parse_summary_payload(
    warnings: list[str],
    quality_report: object | None = None,
    *,
    pdf_summary: object | None = None,
) -> dict[str, object]:
    payload: dict[str, object] = {
        "warning_count": len(warnings),
        "warnings": warnings,
    }
    payload.update(_build_quality_summary_payload(quality_report))
    if pdf_summary is not None and hasattr(pdf_summary, "to_dict"):
        payload["pdf"] = pdf_summary.to_dict()
    return payload


def _build_parsed_snapshot(
    document_version_id: int,
    active_parse_run_id: int,
    parsed_document: ParsedDocumentDraft,
    *,
    warnings: list[str] | None = None,
    quality_report: object | None = None,
) -> dict[str, object]:
    all_blocks = [
        block
        for surface in parsed_document.surfaces
        for block in surface.blocks
    ]
    all_tables = [
        table
        for surface in parsed_document.surfaces
        for table in surface.tables
    ]
    block_counts = Counter(block.block_type for block in all_blocks)
    surface_counts = Counter(surface.surface_type for surface in parsed_document.surfaces)
    snapshot_warnings = parsed_document.warnings if warnings is None else warnings

    snapshot: dict[str, object] = {
        "parser_version": PARSER_VERSION,
        "document_version_id": document_version_id,
        "active_parse_run_id": active_parse_run_id,
        "total_surfaces": len(parsed_document.surfaces),
        "total_blocks": len(all_blocks),
        "counts_by_surface_type": dict(surface_counts),
        "counts_by_block_type": {
            "heading": block_counts.get("heading", 0),
            "paragraph": block_counts.get("paragraph", 0),
            "list_item": block_counts.get("list_item", 0),
            "table_row": block_counts.get("table_row", 0),
        },
        "table_count": len(all_tables),
        "row_count": sum(len(table.rows) for table in all_tables),
        "warning_count": len(snapshot_warnings),
        "warnings": snapshot_warnings,
    }
    snapshot.update(_build_quality_summary_payload(quality_report))
    pdf_summary = getattr(parsed_document, "pdf_summary", None)
    if pdf_summary is not None and hasattr(pdf_summary, "to_dict"):
        snapshot["pdf"] = pdf_summary.to_dict()
    return snapshot


def _build_quality_summary_payload(quality_report: object | None) -> dict[str, object]:
    if quality_report is None:
        return {}
    if hasattr(quality_report, "to_summary_payload"):
        return quality_report.to_summary_payload()

    payload: dict[str, object] = {
        "quality_policy_result": getattr(quality_report, "policy_result", None),
    }
    coverage = getattr(quality_report, "coverage", None)
    diagnostics = getattr(quality_report, "diagnostics", None)
    if coverage is not None and hasattr(coverage, "to_dict"):
        payload["coverage"] = coverage.to_dict()
    if diagnostics is not None:
        payload["diagnostics"] = [
            diagnostic.to_dict() if hasattr(diagnostic, "to_dict") else diagnostic
            for diagnostic in diagnostics
        ]
    return payload


def _build_table_draft(
    table: Table,
    *,
    current_section_title: str | None,
    table_order_index: int,
    next_order_index: int,
    next_surface_order_index: int,
) -> tuple[ParsedTableDraft | None, list[str]]:
    table_key = f"tbl-{table_order_index:04d}"
    resolved_rows, merge_warning = _resolve_table_rows(
        table,
        table_key=table_key,
    )
    raw_rows = [
        [cell.raw_content for cell in row_cells]
        for row_cells in resolved_rows
    ]
    normalized_rows = [
        [
            cell.normalized_content if cell.merge_token is None else ""
            for cell in row_cells
        ]
        for row_cells in resolved_rows
    ]
    populated_rows = [
        row_values
        for row_values in normalized_rows
        if any(cell_value for cell_value in row_values)
    ]
    if not populated_rows:
        return None, []

    warnings: list[str] = []
    if merge_warning is not None:
        warnings.append(merge_warning)
    header_is_explicit = _should_use_explicit_header_row(normalized_rows[0])
    header_strategy = "explicit_first_row" if header_is_explicit else "generated_columns"
    if not header_is_explicit:
        warnings.append(
            f"Generated columns for table {table_order_index + 1} because no explicit header row was detected"
        )

    max_column_count = max(len(row_cells) for row_cells in resolved_rows)
    columns = _build_table_columns(
        raw_rows=raw_rows,
        normalized_rows=normalized_rows,
        header_is_explicit=header_is_explicit,
        max_column_count=max_column_count,
    )

    rows: list[ParsedTableRowDraft] = []
    current_order_index = next_order_index
    for row_index, resolved_row in enumerate(resolved_rows):
        normalized_row = normalized_rows[row_index]
        if not any(normalized_row):
            continue

        row_draft = _build_table_row_draft(
            columns=columns,
            table_key=table_key,
            row_index=row_index,
            is_header_row=header_is_explicit and row_index == 0,
            resolved_row=resolved_row,
            section_title=current_section_title,
            order_index=current_order_index,
            surface_order_index=next_surface_order_index + len(rows),
        )
        rows.append(row_draft)
        current_order_index += 1

    return (
        ParsedTableDraft(
            table_key=table_key,
            section_title=current_section_title,
            table_order_index=table_order_index,
            header_strategy=header_strategy,
            caption_text=None,
            normalized_caption_text=None,
            columns=columns,
            rows=rows,
        ),
        warnings,
    )


def _build_table_columns(
    *,
    raw_rows: list[list[str]],
    normalized_rows: list[list[str]],
    header_is_explicit: bool,
    max_column_count: int,
) -> list[ParsedTableColumnDraft]:
    columns: list[ParsedTableColumnDraft] = []
    seen_column_keys: set[str] = set()
    for column_index in range(max_column_count):
        if header_is_explicit:
            raw_header_text = raw_rows[0][column_index] if column_index < len(raw_rows[0]) else ""
            normalized_header_text = (
                normalized_rows[0][column_index] if column_index < len(normalized_rows[0]) else ""
            )
            header_text = raw_header_text or f"col_{column_index + 1}"
            normalized_header_text = normalized_header_text or f"col_{column_index + 1}"
            column_key = _build_column_key(
                normalized_header_text,
                column_index,
                seen_column_keys=seen_column_keys,
            )
            source_kind = "explicit"
        else:
            header_text = f"col_{column_index + 1}"
            normalized_header_text = header_text
            column_key = header_text
            source_kind = "generated"

        columns.append(
            ParsedTableColumnDraft(
                column_key=column_key,
                column_index=column_index,
                header_text=header_text,
                normalized_header_text=normalized_header_text,
                source_kind=source_kind,
            )
        )

    return columns


def _build_table_row_draft(
    *,
    columns: list[ParsedTableColumnDraft],
    table_key: str,
    row_index: int,
    is_header_row: bool,
    resolved_row: list[_ResolvedTableCellDraft],
    section_title: str | None,
    order_index: int,
    surface_order_index: int,
) -> ParsedTableRowDraft:
    cell_drafts: list[ParsedTableCellDraft] = []
    serialized_cells: list[dict[str, object]] = []
    readable_segments: list[str] = []

    for column in columns:
        resolved_cell = (
            resolved_row[column.column_index]
            if column.column_index < len(resolved_row)
            else _ResolvedTableCellDraft(
                column_index=column.column_index,
                cell_key=f"row-{row_index:04d}-c{column.column_index:04d}",
                raw_content="",
                normalized_content="",
            )
        )
        normalized_value = resolved_cell.merge_token or resolved_cell.normalized_content or "__EMPTY__"
        cell_draft = ParsedTableCellDraft(
            column_key=column.column_key,
            column_index=column.column_index,
            cell_key=resolved_cell.cell_key,
            raw_content=resolved_cell.raw_content,
            normalized_content=resolved_cell.normalized_content,
            merge_token=resolved_cell.merge_token,
            row_span=resolved_cell.row_span,
            col_span=resolved_cell.col_span,
            merge_origin_key=resolved_cell.merge_origin_key,
        )
        cell_drafts.append(cell_draft)
        readable_segments.append(f"{column.normalized_header_text}: {resolved_cell.raw_content}")
        serialized_cells.append(
            {
                "column_key": column.column_key,
                "normalized_header_text": column.normalized_header_text,
                "normalized_value": normalized_value,
                "merge_token": resolved_cell.merge_token,
            }
        )

    raw_content = " || ".join(readable_segments)
    normalized_content = json.dumps(
        serialized_cells,
        separators=(",", ":"),
        ensure_ascii=True,
    )
    block_draft = ParsedBlockDraft(
        block_key=build_block_key(order_index, "table_row", normalized_content),
        block_type="table_row",
        section_title=section_title,
        heading_level=None,
        order_index=order_index,
        surface_order_index=surface_order_index,
        raw_content=raw_content,
        normalized_content=normalized_content,
    )
    structured_row_json = json.dumps(
        {
            "row_index": row_index,
            "is_header_row": is_header_row,
            "cells": [
                {
                    "column_key": column.column_key,
                    "column_index": column.column_index,
                    "header_text": column.header_text,
                    "normalized_header_text": column.normalized_header_text,
                    "raw_value": cell_draft.raw_content,
                    "normalized_value": cell_draft.merge_token
                    or cell_draft.normalized_content
                    or "__EMPTY__",
                    "merge_token": cell_draft.merge_token,
                    "row_span": cell_draft.row_span,
                    "col_span": cell_draft.col_span,
                    "merge_origin_key": cell_draft.merge_origin_key,
                }
                for cell_draft, column in zip(cell_drafts, columns, strict=True)
            ],
        },
        separators=(",", ":"),
        ensure_ascii=True,
    )

    return ParsedTableRowDraft(
        row_key=f"{table_key}-row-{row_index:04d}-{hashlib.sha1(normalized_content.encode('utf-8')).hexdigest()[:10]}",
        row_index=row_index,
        is_header_row=is_header_row,
        structured_row_json=structured_row_json,
        block_draft=block_draft,
        cells=cell_drafts,
    )


def _should_use_explicit_header_row(first_normalized_row: list[str]) -> bool:
    non_empty_cells = [cell for cell in first_normalized_row if cell]
    if not non_empty_cells or len(non_empty_cells) != len(first_normalized_row):
        return False
    if len(set(non_empty_cells)) != len(non_empty_cells):
        return False
    if any(_CODE_LIKE_ROW_RE.match(cell) for cell in non_empty_cells):
        return False
    return True


def _build_column_key(
    header_text: str,
    column_index: int,
    *,
    seen_column_keys: set[str] | None = None,
) -> str:
    normalized_header = normalize_content(header_text).lower()
    slug = _NON_ALNUM_RE.sub("_", normalized_header).strip("_")
    candidate = slug or f"col_{column_index + 1}"
    if seen_column_keys is None:
        return candidate

    if candidate not in seen_column_keys:
        seen_column_keys.add(candidate)
        return candidate

    suffix = 2
    deduplicated = f"{candidate}_{suffix}"
    while deduplicated in seen_column_keys:
        suffix += 1
        deduplicated = f"{candidate}_{suffix}"
    seen_column_keys.add(deduplicated)
    return deduplicated


def _resolve_table_rows(
    table: Table,
    *,
    table_key: str,
) -> tuple[list[list[_ResolvedTableCellDraft]], str | None]:
    grid_column_count = _get_table_grid_column_count(table)
    resolved_rows: list[list[_ResolvedTableCellDraft]] = []
    active_vertical_origins: dict[int, _ResolvedTableCellDraft] = {}
    merge_detected = False

    for row_index, row_element in enumerate(table._tbl.tr_lst):
        row_cells: list[_ResolvedTableCellDraft] = []
        next_vertical_origins: dict[int, _ResolvedTableCellDraft] = {}
        column_index = 0

        for cell_element in row_element.tc_lst:
            col_span = _get_table_cell_col_span(cell_element)
            vmerge_kind = _get_table_cell_vertical_merge_kind(cell_element)

            if vmerge_kind == "continue":
                merge_detected = True
                extended_origin_keys: set[str] = set()
                for offset in range(col_span):
                    effective_column_index = column_index + offset
                    origin_cell = active_vertical_origins.get(effective_column_index)
                    if origin_cell is not None and origin_cell.cell_key not in extended_origin_keys:
                        origin_cell.row_span += 1
                        extended_origin_keys.add(origin_cell.cell_key)
                    if origin_cell is not None:
                        next_vertical_origins[effective_column_index] = origin_cell
                    row_cells.append(
                        _ResolvedTableCellDraft(
                            column_index=effective_column_index,
                            cell_key=f"{table_key}-r{row_index:04d}-c{effective_column_index:04d}",
                            raw_content="",
                            normalized_content="",
                            merge_token="__MERGED__",
                            merge_origin_key=origin_cell.cell_key if origin_cell is not None else None,
                        )
                    )
                column_index += col_span
                continue

            raw_content = _collect_visible_text(cell_element)
            normalized_content = normalize_content(raw_content)
            origin_cell = _ResolvedTableCellDraft(
                column_index=column_index,
                cell_key=f"{table_key}-r{row_index:04d}-c{column_index:04d}",
                raw_content=raw_content,
                normalized_content=normalized_content,
                col_span=col_span,
            )
            row_cells.append(origin_cell)

            if vmerge_kind == "restart":
                merge_detected = True
                for offset in range(col_span):
                    next_vertical_origins[column_index + offset] = origin_cell

            if col_span > 1:
                merge_detected = True
                for offset in range(1, col_span):
                    effective_column_index = column_index + offset
                    row_cells.append(
                        _ResolvedTableCellDraft(
                            column_index=effective_column_index,
                            cell_key=f"{table_key}-r{row_index:04d}-c{effective_column_index:04d}",
                            raw_content="",
                            normalized_content="",
                            merge_token="__MERGED__",
                            merge_origin_key=origin_cell.cell_key,
                        )
                    )

            column_index += col_span

        while column_index < grid_column_count:
            row_cells.append(
                _ResolvedTableCellDraft(
                    column_index=column_index,
                    cell_key=f"{table_key}-r{row_index:04d}-c{column_index:04d}",
                    raw_content="",
                    normalized_content="",
                )
            )
            column_index += 1

        resolved_rows.append(sorted(row_cells, key=lambda cell: cell.column_index))
        active_vertical_origins = next_vertical_origins

    warning = None
    if merge_detected:
        warning = f"Applied merged cell normalization for table {table_key}"

    return resolved_rows, warning


def _get_table_grid_column_count(table: Table) -> int:
    table_grid = getattr(table._tbl, "tblGrid", None)
    grid_columns = getattr(table_grid, "gridCol_lst", None)
    if grid_columns:
        return len(grid_columns)

    max_column_count = 0
    for row_element in table._tbl.tr_lst:
        column_count = 0
        for cell_element in row_element.tc_lst:
            column_count += _get_table_cell_col_span(cell_element)
        max_column_count = max(max_column_count, column_count)
    return max_column_count


def _get_table_cell_col_span(cell_element) -> int:
    tc_properties = getattr(cell_element, "tcPr", None)
    grid_span = getattr(tc_properties, "gridSpan", None)
    if grid_span is None:
        return 1
    try:
        return max(int(grid_span.val), 1)
    except (TypeError, ValueError):
        return 1


def _get_table_cell_vertical_merge_kind(cell_element) -> str | None:
    tc_properties = getattr(cell_element, "tcPr", None)
    if tc_properties is None:
        return None

    vertical_merge = tc_properties.find(qn("w:vMerge"))
    if vertical_merge is None:
        return None

    return vertical_merge.get(qn("w:val")) or "continue"


def _build_surface_draft(
    container: object,
    *,
    surface_type: str,
    surface_key: str,
    logical_order_index: int,
    section_ref: str | None,
    next_global_order_index: int,
    next_table_order_index: int,
) -> tuple[ParsedSurfaceDraft | None, int, int, list[str]]:
    blocks: list[ParsedBlockDraft] = []
    tables: list[ParsedTableDraft] = []
    warnings: list[str] = []
    current_section_title: str | None = None

    for item in _iter_container_items(container):
        if isinstance(item, Paragraph):
            raw_content, paragraph_warnings = _extract_paragraph_raw_content(
                item,
                surface_type=surface_type,
            )
            warnings.extend(paragraph_warnings)
            normalized_content = normalize_content(raw_content)
            if not normalized_content:
                continue

            block_type, heading_level = classify_paragraph(item, normalized_content)
            section_title = current_section_title
            if block_type == "heading":
                section_title = normalized_content
                current_section_title = normalized_content

            blocks.append(
                ParsedBlockDraft(
                    block_key=build_block_key(next_global_order_index, block_type, normalized_content),
                    block_type=block_type,
                    section_title=section_title,
                    heading_level=heading_level,
                    order_index=next_global_order_index,
                    surface_order_index=len(blocks),
                    raw_content=raw_content,
                    normalized_content=normalized_content,
                )
            )
            next_global_order_index += 1
            continue

        table_draft, table_warnings = _build_table_draft(
            item,
            current_section_title=current_section_title,
            table_order_index=next_table_order_index,
            next_order_index=next_global_order_index,
            next_surface_order_index=len(blocks),
        )
        if table_draft is None:
            continue

        tables.append(table_draft)
        warnings.extend(table_warnings)
        blocks.extend(row_draft.block_draft for row_draft in table_draft.rows)
        next_global_order_index += len(table_draft.rows)
        next_table_order_index += 1

    if not blocks and not tables:
        return None, next_global_order_index, next_table_order_index, warnings

    return (
        ParsedSurfaceDraft(
            surface_type=surface_type,
            surface_key=surface_key,
            logical_order_index=logical_order_index,
            section_ref=section_ref,
            blocks=blocks,
            tables=tables,
        ),
        next_global_order_index,
        next_table_order_index,
        warnings,
    )


def _iter_header_footer_containers(
    document: DocxDocumentObject,
) -> Iterator[tuple[int, str, str, object]]:
    for section_index, section in enumerate(document.sections, start=1):
        for surface_type, surface_variant, container in (
            ("header", "default", section.header),
            ("header", "first", section.first_page_header),
            ("header", "even", section.even_page_header),
            ("footer", "default", section.footer),
            ("footer", "first", section.first_page_footer),
            ("footer", "even", section.even_page_footer),
        ):
            if container.is_linked_to_previous:
                continue
            yield section_index, surface_type, surface_variant, container


def _iter_note_containers(
    file_path: Path,
    document: DocxDocumentObject,
) -> Iterator[tuple[str, int, object]]:
    with ZipFile(file_path) as archive:
        referenced_note_ids = _collect_referenced_note_ids(archive)
        for surface_type, relationship_type in _NOTE_RELATIONSHIP_TYPES.items():
            part_xml = _read_related_story_part_xml(archive, relationship_type)
            if part_xml is None:
                continue

            for note_id, note_element in _iter_note_elements(part_xml, surface_type):
                if note_id not in referenced_note_ids[surface_type]:
                    continue
                yield surface_type, note_id, _OOXMLStoryContainer(note_element, document.part)


def _read_related_story_part_xml(archive: ZipFile, relationship_type: str) -> bytes | None:
    try:
        rels_xml = archive.read("word/_rels/document.xml.rels")
    except KeyError:
        return None

    rels_root = etree.fromstring(rels_xml)
    target = rels_root.xpath(
        f"./rel:Relationship[@Type='{relationship_type}']/@Target",
        namespaces=_DOCUMENT_RELS_NSMAP,
    )
    if not target:
        return None

    part_path = str(PurePosixPath("word") / PurePosixPath(target[0]))
    try:
        return archive.read(part_path)
    except KeyError as exc:
        raise DocumentParseError(f"Missing related parser part: {part_path}") from exc


def _collect_referenced_note_ids(archive: ZipFile) -> dict[str, set[int]]:
    try:
        document_xml = archive.read("word/document.xml")
    except KeyError:
        return {surface_type: set() for surface_type in _NOTE_RELATIONSHIP_TYPES}

    document_root = etree.fromstring(document_xml)
    referenced_note_ids: dict[str, set[int]] = {surface_type: set() for surface_type in _NOTE_RELATIONSHIP_TYPES}
    for surface_type in _NOTE_RELATIONSHIP_TYPES:
        for note_id_raw in document_root.xpath(
            f".//w:{surface_type}Reference/@w:id",
            namespaces=_WORD_NSMAP,
        ):
            try:
                note_id = int(note_id_raw)
            except ValueError:
                continue
            if note_id > 0:
                referenced_note_ids[surface_type].add(note_id)

    return referenced_note_ids


def _iter_note_elements(part_xml: bytes, surface_type: str) -> Iterator[tuple[int, object]]:
    root = parse_xml(part_xml)
    note_tag = qn(f"w:{surface_type}")

    for note_element in root.iterchildren():
        if note_element.tag != note_tag:
            continue

        note_type = note_element.get(qn("w:type"))
        if note_type in _IGNORED_NOTE_TYPES:
            continue

        note_id_raw = note_element.get(qn("w:id"))
        if note_id_raw is None:
            continue

        try:
            note_id = int(note_id_raw)
        except ValueError:
            continue

        if note_id <= 0:
            continue

        yield note_id, note_element


def _iter_container_items(container: object) -> Iterator[Paragraph | Table]:
    if isinstance(container, DocxDocumentObject):
        element = container.element.body
    else:
        element = container._element

    for child in element.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, container)
        elif isinstance(child, CT_Tbl):
            yield Table(child, container)


def _extract_paragraph_raw_content(
    paragraph: Paragraph,
    *,
    surface_type: str,
) -> tuple[str, list[str]]:
    raw_segments: list[str] = []
    warnings: list[str] = []
    normalize_auto_fields = surface_type in {"header", "footer"}
    complex_field_state = _ComplexFieldState()

    for child in paragraph._p.iterchildren():
        if normalize_auto_fields and child.tag == qn("w:fldSimple"):
            normalized_field, field_warning = _normalize_supported_auto_field(child, surface_type)
            if normalized_field is not None:
                raw_segments.append(normalized_field)
                warnings.append(field_warning)
                continue

        if normalize_auto_fields and child.tag == qn("w:r"):
            emitted_segment, emitted_warning = _consume_complex_field_run(child, surface_type, complex_field_state)
            if emitted_segment is not None:
                raw_segments.append(emitted_segment)
            if emitted_warning is not None:
                warnings.append(emitted_warning)
            if complex_field_state.inside_field:
                continue
            if _run_is_complex_field_related(child):
                continue

        raw_segments.append(_collect_visible_text(child))

    if complex_field_state.inside_field and complex_field_state.result_segments:
        raw_segments.append("".join(complex_field_state.result_segments))

    return "".join(raw_segments), warnings


def _normalize_supported_auto_field(element, surface_type: str) -> tuple[str | None, str | None]:
    instruction = (element.get(qn("w:instr")) or "").strip()
    return _normalize_supported_auto_field_instruction(instruction, surface_type)


@dataclass(slots=True)
class _ComplexFieldState:
    inside_field: bool = False
    instruction_parts: list[str] = field(default_factory=list)
    result_segments: list[str] = field(default_factory=list)


def _consume_complex_field_run(run_element, surface_type: str, state: _ComplexFieldState) -> tuple[str | None, str | None]:
    if not _run_is_complex_field_related(run_element) and not state.inside_field:
        return None, None

    field_char_types = [
        field_char.get(qn("w:fldCharType"))
        for field_char in run_element.iter(qn("w:fldChar"))
    ]
    instruction_text = "".join(
        instruction.text or ""
        for instruction in run_element.iter(qn("w:instrText"))
    )
    visible_text = _collect_visible_text(run_element)

    if "begin" in field_char_types:
        state.inside_field = True
        state.instruction_parts.clear()
        state.result_segments.clear()

    if state.inside_field and instruction_text:
        state.instruction_parts.append(instruction_text)

    if state.inside_field and visible_text and not instruction_text and "begin" not in field_char_types:
        state.result_segments.append(visible_text)

    if "end" not in field_char_types:
        return None, None

    instruction = "".join(state.instruction_parts).strip()
    normalized_field, field_warning = _normalize_supported_auto_field_instruction(instruction, surface_type)
    fallback_text = "".join(state.result_segments)
    state.inside_field = False
    state.instruction_parts.clear()
    state.result_segments.clear()

    if normalized_field is not None:
        return normalized_field, field_warning
    if fallback_text:
        return fallback_text, None
    return None, None


def _run_is_complex_field_related(run_element) -> bool:
    return any(
        descendant.tag in {qn("w:fldChar"), qn("w:instrText")}
        for descendant in run_element.iter()
    )


def _normalize_supported_auto_field_instruction(
    instruction: str,
    surface_type: str,
) -> tuple[str | None, str | None]:
    instruction = instruction.strip()
    if not instruction:
        return None, None

    field_name = instruction.split()[0].upper()
    placeholder = _SUPPORTED_AUTO_FIELD_PLACEHOLDERS.get(field_name)
    if placeholder is None:
        return None, None

    return placeholder, f"Normalized {surface_type} auto field {field_name}"


def _collect_visible_text(element) -> str:
    text_segments: list[str] = []
    for descendant in element.iter():
        if descendant.tag != qn("w:t") or not descendant.text:
            continue
        if any(
            ancestor.tag in _UNSUPPORTED_VISIBLE_TEXT_CONTAINER_TAGS
            for ancestor in descendant.iterancestors()
        ):
            continue
        text_segments.append(descendant.text)
    return "".join(text_segments)


def _resolve_file_path(version: DocumentVersion) -> Path:
    file_path = Path(version.file_path)
    if file_path.is_absolute():
        return file_path
    return BACKEND_ROOT / file_path


def _get_style_name(paragraph: Paragraph) -> str:
    if paragraph.style is None or paragraph.style.name is None:
        return ""
    return paragraph.style.name.strip()


def _extract_heading_level_from_style(style_name: str) -> int | None:
    if not style_name.lower().startswith("heading"):
        return None

    style_suffix = style_name[len("Heading") :].strip()
    if style_suffix.isdigit():
        return min(max(int(style_suffix), 1), 6)
    return 1


def _has_explicit_list_metadata(paragraph: Paragraph, style_name: str) -> bool:
    if "list" in style_name.lower():
        return True

    paragraph_properties = getattr(paragraph._p, "pPr", None)
    if (
        paragraph_properties is not None
        and getattr(paragraph_properties, "numPr", None) is not None
    ):
        return True

    return False


def _has_marker_only_list(normalized_content: str) -> bool:
    if normalized_content.startswith(("- ", "* ", "+ ")):
        return True

    return (
        _NUMBERED_LIST_RE.match(normalized_content) is not None
        or _PARENTHETICAL_LIST_RE.match(normalized_content) is not None
    )


def _extract_heading_level_from_numbered_heading(normalized_content: str) -> int | None:
    if len(normalized_content) > 120:
        return None

    match = _NUMBERED_HEADING_RE.match(normalized_content)
    if match is None:
        return None

    segments = [segment for segment in match.group("number").split(".") if segment]
    if not segments:
        return None

    return min(len(segments), 6)


def _extract_heading_level_from_legal_heading(normalized_content: str) -> int | None:
    if len(normalized_content) > 160:
        return None
    if _LEGAL_ARTICLE_HEADING_RE.match(normalized_content) is None:
        return None
    return 1
