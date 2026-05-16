"""Generate a professional DOCX review report for a compare run."""

from __future__ import annotations

import io
from datetime import datetime, timezone

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from sqlalchemy import select
from sqlalchemy.orm import Session, joinedload

from app.models import (
    ChangeItem,
    ChangeItemRequirementLink,
    CompareRun,
    Requirement,
    RequirementTestCaseMapping,
)
from app.services.compare import get_compare_run_detail


# Styling helpers.

_HEADER_BG = RGBColor(0x2D, 0x33, 0x3B)
_HEADER_FG = RGBColor(0xFF, 0xFF, 0xFF)
_ADDED_COLOR = RGBColor(0x16, 0x65, 0x34)
_REMOVED_COLOR = RGBColor(0x99, 0x1B, 0x1B)
_MODIFIED_COLOR = RGBColor(0x92, 0x40, 0x0E)
_MUTED_COLOR = RGBColor(0x6B, 0x72, 0x80)


def _set_cell_shading(cell, color_hex: str):
    """Apply background shading to a table cell."""
    from docx.oxml.ns import qn
    from lxml import etree

    shading = etree.SubElement(cell._tc.get_or_add_tcPr(), qn("w:shd"))
    shading.set(qn("w:fill"), color_hex)
    shading.set(qn("w:val"), "clear")


def _style_header_row(row):
    """Apply dark header styling to a table row."""
    for cell in row.cells:
        _set_cell_shading(cell, "2D333B")
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.color.rgb = _HEADER_FG
                run.font.bold = True
                run.font.size = Pt(9)


def _add_table_row(table, cells: list[str]):
    """Add a row with the given cell values."""
    row = table.add_row()
    for i, value in enumerate(cells):
        row.cells[i].text = str(value) if value else ""
        for paragraph in row.cells[i].paragraphs:
            paragraph.style.font.size = Pt(9)
    return row


def _change_type_label(change_type: str) -> str:
    return (change_type or "unknown").replace("_", " ").title()


def _review_status_label(status: str) -> str:
    return (status or "open").replace("_", " ").title()


# Main export function.

def generate_compare_run_docx(
    session: Session,
    compare_run_id: int,
    summary_text: str | None = None,
) -> io.BytesIO:
    """Build a DOCX report and return it as an in-memory BytesIO buffer."""

    # Load data.
    compare_detail = get_compare_run_detail(session, compare_run_id)

    change_items = (
        session.execute(
            select(ChangeItem)
            .where(ChangeItem.compare_run_id == compare_run_id)
            .options(
                joinedload(ChangeItem.assignee),
                joinedload(ChangeItem.ai_review_draft),
                joinedload(ChangeItem.review_comments),
                joinedload(ChangeItem.requirement_links)
                .joinedload(ChangeItemRequirementLink.requirement)
                .joinedload(Requirement.test_case_mappings)
                .joinedload(RequirementTestCaseMapping.test_case),
            )
        )
        .unique()
        .scalars()
        .all()
    )

    summary = compare_detail["summary"]
    doc_title = compare_detail["document"]["title"]
    source_label = compare_detail["source_version"]["version_label"]
    target_label = compare_detail["target_version"]["version_label"]
    now_str = datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M UTC")

    # Build document.
    doc = Document()

    # Default font
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)
    style.paragraph_format.space_after = Pt(4)

    # Title page.
    title = doc.add_heading("Redline Review Report", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(f"{doc_title}\n{source_label} → {target_label}")
    run.font.size = Pt(14)
    run.font.color.rgb = _MUTED_COLOR

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run(f"Generated: {now_str}")
    run.font.size = Pt(9)
    run.font.color.rgb = _MUTED_COLOR

    doc.add_paragraph()  # spacer

    # 1. Executive Summary.
    doc.add_heading("1. Executive Summary", level=1)

    if summary_text:
        for line in summary_text.strip().split("\n"):
            p = doc.add_paragraph(line.strip())
            p.paragraph_format.space_after = Pt(2)
    else:
        doc.add_paragraph(
            "No AI summary was generated for this compare run. "
            "You can generate one from the Summary / Export page.",
            style="Intense Quote",
        )

    # 2. Change Summary.
    doc.add_heading("2. Change Summary", level=1)

    summary_table = doc.add_table(rows=1, cols=4)
    summary_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    summary_table.style = "Table Grid"
    hdr = summary_table.rows[0]
    for i, label in enumerate(["Total Changes", "Added", "Removed", "Modified"]):
        hdr.cells[i].text = label
    _style_header_row(hdr)
    _add_table_row(summary_table, [
        str(summary["total_changes"]),
        str(summary["added"]),
        str(summary["removed"]),
        str(summary["modified"]),
    ])

    # Review counts
    review_counts = {"resolved": 0, "in_review": 0, "open": 0}
    for item in change_items:
        status = item.review_status
        if status in review_counts:
            review_counts[status] += 1

    doc.add_paragraph()
    review_table = doc.add_table(rows=1, cols=3)
    review_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    review_table.style = "Table Grid"
    hdr = review_table.rows[0]
    for i, label in enumerate(["Resolved", "In Review", "Open"]):
        hdr.cells[i].text = label
    _style_header_row(hdr)
    _add_table_row(review_table, [
        str(review_counts["resolved"]),
        str(review_counts["in_review"]),
        str(review_counts["open"]),
    ])

    # 3. Change Items Detail.
    doc.add_heading("3. Change Items", level=1)

    if not change_items:
        doc.add_paragraph("No change items were detected in this compare run.")
    else:
        detail_table = doc.add_table(rows=1, cols=6)
        detail_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        detail_table.style = "Table Grid"
        hdr = detail_table.rows[0]
        for i, label in enumerate([
            "#", "Section", "Type", "Status", "Old Content", "New Content"
        ]):
            hdr.cells[i].text = label
        _style_header_row(hdr)

        for idx, item in enumerate(change_items, 1):
            old_text = (item.old_content or "-")[:200]
            new_text = (item.new_content or "-")[:200]
            if item.old_content and len(item.old_content) > 200:
                old_text += "..."
            if item.new_content and len(item.new_content) > 200:
                new_text += "..."

            row = _add_table_row(detail_table, [
                str(idx),
                item.section_title or item.surface_key or "-",
                _change_type_label(item.change_type),
                _review_status_label(item.review_status),
                old_text,
                new_text,
            ])

            # Color the type cell
            type_cell = row.cells[2]
            for paragraph in type_cell.paragraphs:
                for run in paragraph.runs:
                    if item.change_type == "added":
                        run.font.color.rgb = _ADDED_COLOR
                    elif item.change_type == "removed":
                        run.font.color.rgb = _REMOVED_COLOR
                    elif item.change_type == "modified":
                        run.font.color.rgb = _MODIFIED_COLOR

    # 4. AI Review Insights.
    items_with_ai = [item for item in change_items if item.ai_review_draft]
    if items_with_ai:
        doc.add_heading("4. AI Review Insights", level=1)

        ai_table = doc.add_table(rows=1, cols=5)
        ai_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        ai_table.style = "Table Grid"
        hdr = ai_table.rows[0]
        for i, label in enumerate([
            "Section", "Risk", "Explanation", "Suggested Status", "Draft Comment"
        ]):
            hdr.cells[i].text = label
        _style_header_row(hdr)

        for item in items_with_ai:
            draft = item.ai_review_draft
            explanation = (draft.explanation or "-")[:300]
            if draft.explanation and len(draft.explanation) > 300:
                explanation += "..."
            _add_table_row(ai_table, [
                item.section_title or item.surface_key or "-",
                (draft.risk_level or "-").upper(),
                explanation,
                _review_status_label(draft.recommended_review_status) if draft.recommended_review_status else "-",
                (draft.draft_comment or "-")[:200],
            ])

    # 5. Traceability Matrix.
    all_links = []
    for item in change_items:
        for req_link in item.requirement_links:
            req = req_link.requirement
            tests = [
                f"{m.test_case.test_case_code}: {m.test_case.title}"
                for m in req.test_case_mappings
            ]
            all_links.append({
                "section": item.section_title or item.surface_key or "-",
                "change_type": _change_type_label(item.change_type),
                "requirement": f"{req.requirement_code}: {req.title}",
                "tests": tests,
            })

    if all_links:
        section_num = 5 if items_with_ai else 4
        doc.add_heading(f"{section_num}. Traceability Matrix", level=1)

        trace_table = doc.add_table(rows=1, cols=4)
        trace_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        trace_table.style = "Table Grid"
        hdr = trace_table.rows[0]
        for i, label in enumerate([
            "Section", "Change Type", "Requirement", "Impacted Tests"
        ]):
            hdr.cells[i].text = label
        _style_header_row(hdr)

        for link in all_links:
            _add_table_row(trace_table, [
                link["section"],
                link["change_type"],
                link["requirement"],
                "; ".join(link["tests"]) if link["tests"] else "No mapped tests",
            ])

    # 6. Review Comments.
    items_with_comments = [
        item for item in change_items
        if item.review_comments and len(item.review_comments) > 0
    ]
    if items_with_comments:
        section_num = (5 if items_with_ai else 4) + (1 if all_links else 0)
        doc.add_heading(f"{section_num}. Review Comments", level=1)

        for item in items_with_comments:
            section = item.section_title or item.surface_key or f"Change #{item.id}"
            doc.add_heading(section, level=3)
            for comment in item.review_comments:
                author = comment.author.display_name if comment.author else f"User {comment.author_user_id}"
                created = comment.created_at.strftime("%Y-%m-%d %H:%M") if comment.created_at else ""
                p = doc.add_paragraph()
                run = p.add_run(f"{author} ({created}): ")
                run.font.bold = True
                run.font.size = Pt(9)
                run = p.add_run(comment.content)
                run.font.size = Pt(9)

    # Footer.
    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("- Generated by Redline -")
    run.font.size = Pt(8)
    run.font.color.rgb = _MUTED_COLOR
    run.font.italic = True

    # Serialize.
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
