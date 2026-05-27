"""Generate a .docx review summary report for a compare run."""

from __future__ import annotations

import io
from datetime import datetime, timezone

from docx import Document as DocxDocument
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from sqlalchemy.orm import Session

from app.models.change_item import ChangeItem
from app.models.compare_run import CompareRun


def _rgb(hex_str: str) -> RGBColor:
    hex_str = hex_str.lstrip("#")
    return RGBColor(int(hex_str[0:2], 16), int(hex_str[2:4], 16), int(hex_str[4:6], 16))


_GREEN = _rgb("16714E")
_RED = _rgb("C03050")
_AMBER = _rgb("B07D0A")
_GREY = _rgb("848E9C")
_DARK = _rgb("1E2026")

_STATUS_LABELS = {
    "open": "Open",
    "in_review": "In Review",
    "resolved": "Resolved",
}

_CHANGE_TYPE_LABELS = {
    "added": "Added",
    "removed": "Removed",
    "modified": "Modified",
}


def _add_heading(doc: DocxDocument, text: str, level: int = 1) -> None:
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.color.rgb = _DARK


def _add_metadata_line(doc: DocxDocument, label: str, value: str) -> None:
    para = doc.add_paragraph()
    para.paragraph_format.space_after = Pt(2)
    run_label = para.add_run(f"{label}: ")
    run_label.bold = True
    run_label.font.size = Pt(10)
    run_label.font.color.rgb = _GREY
    run_value = para.add_run(value)
    run_value.font.size = Pt(10)
    run_value.font.color.rgb = _DARK


def _add_change_section(
    doc: DocxDocument,
    index: int,
    item: ChangeItem,
) -> None:
    # Section header
    title = item.section_title or f"Clause #{item.id}"
    change_type = _CHANGE_TYPE_LABELS.get(item.change_type, item.change_type or "Changed")
    review_label = _STATUS_LABELS.get(item.review_status, item.review_status or "Open")

    heading = doc.add_heading(f"{index}. {title}", level=2)
    for run in heading.runs:
        run.font.color.rgb = _DARK

    # Badges line
    badge_para = doc.add_paragraph()
    badge_para.paragraph_format.space_after = Pt(4)

    type_run = badge_para.add_run(f"[{change_type}]")
    type_run.bold = True
    type_run.font.size = Pt(9)
    ct = (item.change_type or "").lower()
    type_run.font.color.rgb = _GREEN if ct == "added" else _RED if ct == "removed" else _AMBER

    badge_para.add_run("  ")

    status_run = badge_para.add_run(f"[{review_label}]")
    status_run.bold = True
    status_run.font.size = Pt(9)
    rs = (item.review_status or "").lower()
    status_run.font.color.rgb = _GREEN if rs == "resolved" else _AMBER if rs == "in_review" else _RED

    # Summary
    if item.summary:
        summary_para = doc.add_paragraph()
        summary_run = summary_para.add_run(f"Summary: {item.summary}")
        summary_run.italic = True
        summary_run.font.size = Pt(10)
        summary_run.font.color.rgb = _GREY

    # Old content
    if item.old_content:
        doc.add_paragraph()
        label_para = doc.add_paragraph()
        label_run = label_para.add_run("Original Text:")
        label_run.bold = True
        label_run.font.size = Pt(10)
        label_run.font.color.rgb = _RED

        old_para = doc.add_paragraph()
        old_para.paragraph_format.left_indent = Inches(0.3)
        old_run = old_para.add_run(item.old_content.strip()[:2000])
        old_run.font.size = Pt(9)
        old_run.font.color.rgb = _rgb("474D57")

    # New content
    if item.new_content:
        label_para = doc.add_paragraph()
        label_run = label_para.add_run("Revised Text:")
        label_run.bold = True
        label_run.font.size = Pt(10)
        label_run.font.color.rgb = _GREEN

        new_para = doc.add_paragraph()
        new_para.paragraph_format.left_indent = Inches(0.3)
        new_run = new_para.add_run(item.new_content.strip()[:2000])
        new_run.font.size = Pt(9)
        new_run.font.color.rgb = _rgb("474D57")

    # AI Review Draft
    ai_draft = item.ai_review_draft
    if ai_draft and ai_draft.explanation:
        doc.add_paragraph()
        ai_label = doc.add_paragraph()
        ai_run = ai_label.add_run("✦ AI Analysis:")
        ai_run.bold = True
        ai_run.font.size = Pt(10)
        ai_run.font.color.rgb = _AMBER

        if ai_draft.risk_level:
            risk_run = ai_label.add_run(f"  [{ai_draft.risk_level.upper()} RISK]")
            risk_run.bold = True
            risk_run.font.size = Pt(9)
            risk = (ai_draft.risk_level or "").lower()
            risk_run.font.color.rgb = _RED if risk == "high" else _AMBER if risk == "medium" else _GREEN

        ai_para = doc.add_paragraph()
        ai_para.paragraph_format.left_indent = Inches(0.3)
        expl_run = ai_para.add_run(ai_draft.explanation.strip()[:2000])
        expl_run.font.size = Pt(9)
        expl_run.font.color.rgb = _rgb("474D57")

        if ai_draft.draft_comment:
            redline_label = doc.add_paragraph()
            redline_label.paragraph_format.left_indent = Inches(0.3)
            rl_run = redline_label.add_run("Suggested Redline: ")
            rl_run.bold = True
            rl_run.font.size = Pt(9)
            rl_run.font.color.rgb = _GREEN
            rl_text = redline_label.add_run(ai_draft.draft_comment.strip()[:1000])
            rl_text.font.size = Pt(9)
            rl_text.font.color.rgb = _rgb("474D57")

    # Comments
    if item.review_comments:
        doc.add_paragraph()
        comments_label = doc.add_paragraph()
        cl_run = comments_label.add_run(f"💬 Comments ({len(item.review_comments)}):")
        cl_run.bold = True
        cl_run.font.size = Pt(10)
        cl_run.font.color.rgb = _GREY

        for comment in item.review_comments:
            author = comment.author.display_name if comment.author else f"User {comment.author_user_id}"
            c_para = doc.add_paragraph()
            c_para.paragraph_format.left_indent = Inches(0.3)
            c_author = c_para.add_run(f"{author}: ")
            c_author.bold = True
            c_author.font.size = Pt(9)
            c_author.font.color.rgb = _DARK
            c_text = c_para.add_run(comment.content)
            c_text.font.size = Pt(9)
            c_text.font.color.rgb = _rgb("474D57")

    # Horizontal rule
    doc.add_paragraph("─" * 60)


def generate_review_report(
    session: Session,
    compare_run: CompareRun,
) -> tuple[io.BytesIO, str]:
    """Generate a .docx review report and return (bytes_io, filename)."""

    doc = DocxDocument()

    # Page title
    _add_heading(doc, "Review Summary Report", level=1)

    # Metadata
    source_label = (
        compare_run.source_version.version_label
        if compare_run.source_version
        else "Unknown"
    )
    target_label = (
        compare_run.target_version.version_label
        if compare_run.target_version
        else "Unknown"
    )
    contract_title = (
        compare_run.source_version.document.title
        if compare_run.source_version and compare_run.source_version.document
        else "Unknown Contract"
    )

    _add_metadata_line(doc, "Contract", contract_title)
    _add_metadata_line(doc, "Source Version", source_label)
    _add_metadata_line(doc, "Target Version", target_label)
    _add_metadata_line(doc, "Generated", datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M UTC"))

    # Statistics
    items: list[ChangeItem] = list(compare_run.change_items)
    total = len(items)
    resolved_count = sum(1 for i in items if i.review_status == "resolved")
    in_review_count = sum(1 for i in items if i.review_status == "in_review")
    open_count = sum(1 for i in items if i.review_status == "open")
    added_count = sum(1 for i in items if i.change_type == "added")
    modified_count = sum(1 for i in items if i.change_type == "modified")
    removed_count = sum(1 for i in items if i.change_type == "removed")

    _add_metadata_line(doc, "Total Changes", str(total))
    _add_metadata_line(
        doc,
        "Review Progress",
        f"{resolved_count} resolved · {in_review_count} in review · {open_count} open",
    )
    _add_metadata_line(
        doc,
        "Change Types",
        f"{added_count} added · {modified_count} modified · {removed_count} removed",
    )

    doc.add_paragraph()

    # Change items
    _add_heading(doc, "Detailed Changes", level=1)

    for index, item in enumerate(items, start=1):
        _add_change_section(doc, index, item)

    # Footer
    footer_para = doc.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer_para.add_run("Generated by Redline AI · Review Summary Export")
    footer_run.font.size = Pt(8)
    footer_run.font.color.rgb = _GREY
    footer_run.italic = True

    # Save to bytes
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    safe_title = "".join(c if c.isalnum() or c in " -_" else "_" for c in contract_title)[:50]
    filename = f"Review_Report_{safe_title}_{source_label}_vs_{target_label}.docx"
    filename = filename.replace(" ", "_")

    return buffer, filename
