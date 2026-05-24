from io import BytesIO

from docx import Document as DocxDocument
from sqlalchemy import select

from app.models import AIBatchJob, AIBatchJobItem, AIReviewDraft, ChangeItem
from app.services.llm_adapter import NormalizedAIReviewDraft


def _build_compare_docx(requirement_lines: list[str]) -> bytes:
    document = DocxDocument()
    document.add_paragraph("Requirements", style="Heading 1")
    for requirement_line in requirement_lines:
        document.add_paragraph(requirement_line)

    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _create_compare_run(client, auth_headers) -> int:
    project_response = client.post(
        "/api/v1/projects",
        json={"name": "AI Batch API Project", "description": "AI batch endpoint coverage"},
        headers=auth_headers,
    )
    assert project_response.status_code == 201
    project_id = project_response.json()["data"]["id"]

    document_response = client.post(
        f"/api/v1/projects/{project_id}/documents",
        json={
            "title": "AI Batch API Spec",
            "document_type": "SPEC",
            "description": "AI batch endpoint target document",
        },
        headers=auth_headers,
    )
    assert document_response.status_code == 201
    document_id = document_response.json()["data"]["id"]

    source_response = client.post(
        f"/api/v1/documents/{document_id}/versions",
        files={
            "file": (
                "ai-batch-api-source.docx",
                _build_compare_docx(
                    [
                        "The system shall support login.",
                        "The system shall write audit logs.",
                    ]
                ),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
        data={"version_label": "v1.0", "notes": "AI batch source"},
        headers=auth_headers,
    )
    assert source_response.status_code == 201
    source_version_id = source_response.json()["data"]["id"]

    target_response = client.post(
        f"/api/v1/documents/{document_id}/versions",
        files={
            "file": (
                "ai-batch-api-target.docx",
                _build_compare_docx(
                    [
                        "The system shall support secure login.",
                        "The system shall write tamper-proof audit logs.",
                    ]
                ),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
        data={"version_label": "v1.1", "notes": "AI batch target"},
        headers=auth_headers,
    )
    assert target_response.status_code == 201
    target_version_id = target_response.json()["data"]["id"]

    assert client.post(f"/api/v1/document-versions/{source_version_id}/parse", headers=auth_headers).status_code == 200
    assert client.post(f"/api/v1/document-versions/{target_version_id}/parse", headers=auth_headers).status_code == 200

    compare_response = client.post(
        f"/api/v1/documents/{document_id}/compare-runs",
        json={"source_version_id": source_version_id, "target_version_id": target_version_id},
        headers=auth_headers,
    )
    assert compare_response.status_code == 201
    return compare_response.json()["data"]["id"]


def _compare_run_setup(session_factory, compare_run_id: int) -> dict[str, int]:
    with session_factory() as session:
        change_item = session.scalar(
            select(ChangeItem).where(ChangeItem.compare_run_id == compare_run_id).order_by(ChangeItem.id)
        )
        assert change_item is not None
        compare_run = change_item.compare_run
        return {
            "document_id": compare_run.source_version.document_id,
            "source_version_id": compare_run.source_version_id,
            "target_version_id": compare_run.target_version_id,
        }


def _reparse_target_version(client, auth_headers, session_factory, compare_run_id: int) -> None:
    target_version_id = _compare_run_setup(session_factory, compare_run_id)["target_version_id"]
    response = client.post(f"/api/v1/document-versions/{target_version_id}/parse", headers=auth_headers)
    assert response.status_code == 200


def _supersede_compare_run(client, auth_headers, session_factory, compare_run_id: int) -> int:
    setup = _compare_run_setup(session_factory, compare_run_id)
    response = client.post(
        f"/api/v1/documents/{setup['document_id']}/compare-runs",
        json={
            "source_version_id": setup["source_version_id"],
            "target_version_id": setup["target_version_id"],
        },
        headers=auth_headers,
    )
    assert response.status_code == 201
    superseding_compare_run_id = response.json()["data"]["id"]
    assert superseding_compare_run_id != compare_run_id
    return superseding_compare_run_id


def test_batch_generate_endpoint_creates_job_immediately_without_generating_drafts(
    client,
    auth_headers,
    session_factory,
):
    compare_run_id = _create_compare_run(client, auth_headers)

    response = client.post(
        f"/api/v1/compare-runs/{compare_run_id}/ai-review-drafts/generate",
        json={"force_regenerate": False},
        headers=auth_headers,
    )

    assert response.status_code == 200
    payload = response.json()["data"]
    assert payload["compare_run_id"] == compare_run_id
    assert payload["status"] == "queued"
    assert payload["requested_count"] == 2
    assert payload["processed_count"] == 0
    assert payload["generated_count"] == 0
    assert payload["failed_count"] == 0
    assert payload["active"] is True

    second_response = client.post(
        f"/api/v1/compare-runs/{compare_run_id}/ai-review-drafts/generate",
        json={"force_regenerate": False},
        headers=auth_headers,
    )
    assert second_response.status_code == 200
    assert second_response.json()["data"]["job_id"] == payload["job_id"]

    with session_factory() as session:
        jobs = list(session.scalars(select(AIBatchJob).order_by(AIBatchJob.id)))
        items = list(
            session.scalars(
                select(AIBatchJobItem)
                .join(AIBatchJob, AIBatchJob.id == AIBatchJobItem.job_id)
                .where(AIBatchJob.compare_run_id == compare_run_id)
            )
        )
        drafts = list(
            session.scalars(
                select(AIReviewDraft)
                .join(ChangeItem, ChangeItem.id == AIReviewDraft.change_item_id)
                .where(ChangeItem.compare_run_id == compare_run_id)
            )
        )

    assert len(jobs) == 1
    assert len(items) == 2
    assert drafts == []


def test_ai_batch_generate_rejects_stale_compare_run(client, auth_headers, session_factory):
    compare_run_id = _create_compare_run(client, auth_headers)
    _reparse_target_version(client, auth_headers, session_factory, compare_run_id)

    response = client.post(
        f"/api/v1/compare-runs/{compare_run_id}/ai-review-drafts/generate",
        json={"force_regenerate": True},
        headers=auth_headers,
    )

    assert response.status_code == 422
    assert "stale" in response.json()["detail"].lower()


def test_ai_batch_generate_rejects_superseded_compare_run(client, auth_headers, session_factory):
    compare_run_id = _create_compare_run(client, auth_headers)
    _supersede_compare_run(client, auth_headers, session_factory, compare_run_id)

    response = client.post(
        f"/api/v1/compare-runs/{compare_run_id}/ai-review-drafts/generate",
        json={"force_regenerate": True},
        headers=auth_headers,
    )

    assert response.status_code == 422
    assert "superseded" in response.json()["detail"].lower()


def test_compare_run_detail_exposes_active_ai_batch_job(client, auth_headers):
    compare_run_id = _create_compare_run(client, auth_headers)

    create_response = client.post(
        f"/api/v1/compare-runs/{compare_run_id}/ai-review-drafts/generate",
        json={"force_regenerate": False},
        headers=auth_headers,
    )
    assert create_response.status_code == 200
    job_id = create_response.json()["data"]["job_id"]

    compare_run_response = client.get(
        f"/api/v1/compare-runs/{compare_run_id}",
        headers=auth_headers,
    )

    assert compare_run_response.status_code == 200
    payload = compare_run_response.json()["data"]
    assert payload["active_ai_batch_job"]["job_id"] == job_id
    assert payload["active_ai_batch_job"]["status"] == "queued"
    assert payload["ai_batch_summary"]["job_id"] == job_id


def test_ai_batch_job_endpoints_return_progress_and_item_results(
    client,
    auth_headers,
    session_factory,
    monkeypatch,
):
    compare_run_id = _create_compare_run(client, auth_headers)

    from app.services import ai_batch_jobs as ai_batch_job_service

    class StubAdapter:
        def __init__(self):
            self.calls = 0

        def generate_ai_review_draft(self, payload):
            self.calls += 1
            if self.calls == 1:
                return NormalizedAIReviewDraft(
                    suggested_assignee_user_id=1,
                    recommended_review_status="in_review",
                    explanation=f"Generated draft for item {payload['change_item_id']}",
                    risk_level="medium",
                    draft_comment="Verify the authentication impact.",
                    suggested_checks="Review impacted authentication tests.",
                    confidence=0.81,
                    generation_status="generated",
                    provider_used="gemini",
                    fallback_used=False,
                    error_message=None,
                )

            return NormalizedAIReviewDraft(
                suggested_assignee_user_id=None,
                recommended_review_status="open",
                explanation="AI draft generation failed.",
                risk_level=None,
                draft_comment=None,
                suggested_checks=None,
                confidence=None,
                generation_status="failed",
                provider_used="openai",
                fallback_used=True,
                error_message="fallback provider unavailable",
            )

    monkeypatch.setattr(ai_batch_job_service, "get_llm_adapter", lambda: StubAdapter())

    create_response = client.post(
        f"/api/v1/compare-runs/{compare_run_id}/ai-review-drafts/generate",
        json={"force_regenerate": False},
        headers=auth_headers,
    )
    assert create_response.status_code == 200
    job_id = create_response.json()["data"]["job_id"]

    processed = ai_batch_job_service.process_next_ai_batch_job(session_factory, concurrency=1)
    assert processed is True

    job_response = client.get(
        f"/api/v1/ai-batch-jobs/{job_id}",
        headers=auth_headers,
    )
    assert job_response.status_code == 200
    job_payload = job_response.json()["data"]
    assert job_payload["status"] == "completed_with_failures"
    assert job_payload["requested_count"] == 2
    assert job_payload["processed_count"] == 2
    assert job_payload["generated_count"] == 1
    assert job_payload["failed_count"] == 1

    items_response = client.get(
        f"/api/v1/ai-batch-jobs/{job_id}/items",
        headers=auth_headers,
    )
    assert items_response.status_code == 200
    item_payload = items_response.json()["data"]
    assert [item["status"] for item in item_payload] == ["generated", "failed"]
    assert item_payload[0]["provider_used"] == "gemini"
    assert item_payload[1]["error_message"] == "fallback provider unavailable"


def test_ai_batch_job_can_disable_rag_context(
    client,
    auth_headers,
    session_factory,
    monkeypatch,
):
    compare_run_id = _create_compare_run(client, auth_headers)

    from app.services import ai_batch_jobs as ai_batch_job_service

    captured_payloads = []

    class StubAdapter:
        def generate_ai_review_draft(self, payload):
            captured_payloads.append(payload)
            return NormalizedAIReviewDraft(
                suggested_assignee_user_id=None,
                recommended_review_status="open",
                explanation="Generated draft.",
                risk_level="medium",
                draft_comment="Review the clause update.",
                suggested_checks="Check liability alignment.",
                confidence=0.81,
                generation_status="generated",
                provider_used="stub",
                fallback_used=False,
                error_message=None,
            )

    monkeypatch.setattr(ai_batch_job_service, "get_llm_adapter", lambda: StubAdapter())

    create_response = client.post(
        f"/api/v1/compare-runs/{compare_run_id}/ai-review-drafts/generate",
        json={"force_regenerate": True, "use_rag": False},
        headers=auth_headers,
    )
    assert create_response.status_code == 200
    assert create_response.json()["data"]["use_rag"] is False

    processed = ai_batch_job_service.process_next_ai_batch_job(session_factory, concurrency=1)
    assert processed is True

    assert captured_payloads
    assert all(payload["rag_enabled"] is False for payload in captured_payloads)
    assert all(payload["rag_context"] == [] for payload in captured_payloads)
