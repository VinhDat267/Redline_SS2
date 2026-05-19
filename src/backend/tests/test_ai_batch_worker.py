from datetime import timedelta
from io import BytesIO
from threading import Event

from docx import Document as DocxDocument
from fastapi.testclient import TestClient
from sqlalchemy import select

from app.core.config import settings
from app.models import AIBatchJob, AIBatchJobItem, AIReviewDraft, ChangeItem
from app.main import create_app
from app.models.mixins import utcnow
from app.services.llm_adapter import NormalizedAIReviewDraft


def _build_compare_docx(requirement_lines: list[str]) -> bytes:
    document = DocxDocument()
    document.add_paragraph("Requirements", style="Heading 1")
    for requirement_line in requirement_lines:
        document.add_paragraph(requirement_line)

    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _create_compare_run(client, auth_headers) -> int:
    project_response = client.post(
        "/api/v1/projects",
        json={"name": "AI Batch Worker Project", "description": "Batch worker coverage"},
        headers=auth_headers,
    )
    assert project_response.status_code == 201
    project_id = project_response.json()["data"]["id"]

    document_response = client.post(
        f"/api/v1/projects/{project_id}/documents",
        json={
            "title": "AI Batch Worker Spec",
            "document_type": "SPEC",
            "description": "AI batch worker target document",
        },
        headers=auth_headers,
    )
    assert document_response.status_code == 201
    document_id = document_response.json()["data"]["id"]

    source_response = client.post(
        f"/api/v1/documents/{document_id}/versions",
        files={
            "file": (
                "ai-batch-source.docx",
                _build_compare_docx(
                    [
                        "The system shall support login.",
                        "The system shall write audit logs.",
                    ]
                ),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
        data={"version_label": "v1.0", "notes": "AI batch source"},
        headers=auth_headers,
    )
    assert source_response.status_code == 201
    source_version_id = source_response.json()["data"]["id"]

    target_response = client.post(
        f"/api/v1/documents/{document_id}/versions",
        files={
            "file": (
                "ai-batch-target.docx",
                _build_compare_docx(
                    [
                        "The system shall support secure login.",
                        "The system shall write tamper-proof audit logs.",
                    ]
                ),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
        data={"version_label": "v1.1", "notes": "AI batch target"},
        headers=auth_headers,
    )
    assert target_response.status_code == 201
    target_version_id = target_response.json()["data"]["id"]

    assert client.post(f"/api/v1/document-versions/{source_version_id}/parse", headers=auth_headers).status_code == 200
    assert client.post(f"/api/v1/document-versions/{target_version_id}/parse", headers=auth_headers).status_code == 200

    compare_response = client.post(
        f"/api/v1/documents/{document_id}/compare-runs",
        json={"source_version_id": source_version_id, "target_version_id": target_version_id},
        headers=auth_headers,
    )
    assert compare_response.status_code == 201
    return compare_response.json()["data"]["id"]


def test_create_batch_job_reuses_active_job(client, auth_headers, session_factory):
    compare_run_id = _create_compare_run(client, auth_headers)

    from app.services import ai_batch_jobs as ai_batch_job_service

    with session_factory() as session:
        first_job = ai_batch_job_service.create_compare_run_ai_batch_job(
            session,
            compare_run_id=compare_run_id,
            actor_user_id=1,
            force_regenerate=False,
        )
        session.commit()

    with session_factory() as session:
        second_job = ai_batch_job_service.create_compare_run_ai_batch_job(
            session,
            compare_run_id=compare_run_id,
            actor_user_id=1,
            force_regenerate=False,
        )
        session.commit()

        jobs = list(session.scalars(select(AIBatchJob).order_by(AIBatchJob.id)))
        items = list(
            session.scalars(
                select(AIBatchJobItem)
                .join(AIBatchJob, AIBatchJob.id == AIBatchJobItem.job_id)
                .where(AIBatchJob.compare_run_id == compare_run_id)
                .order_by(AIBatchJobItem.id)
            )
        )

    assert first_job["job_id"] == second_job["job_id"]
    assert len(jobs) == 1
    assert len(items) == 2
    assert all(item.status == "queued" for item in items)


def test_process_next_batch_job_updates_drafts_and_job_counts(
    client,
    auth_headers,
    session_factory,
    monkeypatch,
):
    compare_run_id = _create_compare_run(client, auth_headers)

    from app.services import ai_batch_jobs as ai_batch_job_service

    class StubAdapter:
        def __init__(self):
            self.calls = 0

        def generate_ai_review_draft(self, payload):
            self.calls += 1
            if self.calls == 1:
                return NormalizedAIReviewDraft(
                    suggested_assignee_user_id=1,
                    recommended_review_status="in_review",
                    explanation=f"Generated draft for item {payload['change_item_id']}",
                    risk_level="medium",
                    draft_comment="Verify the authentication impact.",
                    suggested_checks="Review impacted authentication tests.",
                    confidence=0.81,
                    generation_status="generated",
                    provider_used="gemini",
                    fallback_used=False,
                    error_message=None,
                )

            return NormalizedAIReviewDraft(
                suggested_assignee_user_id=None,
                recommended_review_status="open",
                explanation="AI draft generation failed.",
                risk_level=None,
                draft_comment=None,
                suggested_checks=None,
                confidence=None,
                generation_status="failed",
                provider_used="openai",
                fallback_used=True,
                error_message="fallback provider unavailable",
            )

    monkeypatch.setattr(ai_batch_job_service, "get_llm_adapter", lambda: StubAdapter())

    with session_factory() as session:
        created_job = ai_batch_job_service.create_compare_run_ai_batch_job(
            session,
            compare_run_id=compare_run_id,
            actor_user_id=1,
            force_regenerate=False,
        )
        session.commit()

    processed = ai_batch_job_service.process_next_ai_batch_job(
        session_factory,
        concurrency=1,
    )

    assert processed is True

    with session_factory() as session:
        job = session.get(AIBatchJob, created_job["job_id"])
        assert job is not None
        assert job.status == "completed_with_failures"
        assert job.processed_count == 2
        assert job.generated_count == 1
        assert job.failed_count == 1

        items = list(
            session.scalars(
                select(AIBatchJobItem)
                .where(AIBatchJobItem.job_id == job.id)
                .order_by(AIBatchJobItem.id)
            )
        )
        assert [item.status for item in items] == ["generated", "failed"]

        drafts = list(
            session.scalars(
                select(AIReviewDraft)
                .join(ChangeItem, ChangeItem.id == AIReviewDraft.change_item_id)
                .where(ChangeItem.compare_run_id == compare_run_id)
                .order_by(AIReviewDraft.change_item_id)
            )
        )
        assert len(drafts) == 2
        assert drafts[0].generation_status == "generated"
        assert drafts[1].generation_status == "failed"


def test_process_next_batch_job_groups_ai_review_provider_calls(
    client,
    auth_headers,
    session_factory,
    monkeypatch,
):
    compare_run_id = _create_compare_run(client, auth_headers)

    from app.services import ai_batch_jobs as ai_batch_job_service

    class StubAdapter:
        def __init__(self):
            self.batch_calls: list[list[int]] = []
            self.single_calls = 0

        def generate_ai_review_drafts_batch(self, payloads):
            change_item_ids = [int(payload["change_item_id"]) for payload in payloads]
            self.batch_calls.append(change_item_ids)
            return [
                NormalizedAIReviewDraft(
                    suggested_assignee_user_id=1,
                    recommended_review_status="in_review",
                    explanation=f"Generated batch draft for item {change_item_id}",
                    risk_level="medium",
                    draft_comment="Verify the batch-generated review.",
                    suggested_checks="Review impacted legal clauses.",
                    confidence=0.82,
                    generation_status="generated",
                    provider_used="gemini",
                    fallback_used=False,
                    error_message=None,
                )
                for change_item_id in change_item_ids
            ]

        def generate_ai_review_draft(self, payload):
            self.single_calls += 1
            return NormalizedAIReviewDraft(
                suggested_assignee_user_id=None,
                recommended_review_status="open",
                explanation="Single generation should not be used for a batch-capable adapter.",
                risk_level=None,
                draft_comment=None,
                suggested_checks=None,
                confidence=None,
                generation_status="failed",
                provider_used="stub",
                fallback_used=False,
                error_message="single path used",
            )

    adapter = StubAdapter()
    monkeypatch.setattr(ai_batch_job_service, "get_llm_adapter", lambda: adapter)
    monkeypatch.setattr(settings, "ai_review_batch_size", 2)
    monkeypatch.setattr(settings, "ai_batch_inter_item_delay", 0.0)

    with session_factory() as session:
        created_job = ai_batch_job_service.create_compare_run_ai_batch_job(
            session,
            compare_run_id=compare_run_id,
            actor_user_id=1,
            force_regenerate=False,
        )
        session.commit()

    processed = ai_batch_job_service.process_next_ai_batch_job(
        session_factory,
        concurrency=1,
    )

    assert processed is True
    assert len(adapter.batch_calls) == 1
    assert len(adapter.batch_calls[0]) == 2
    assert adapter.single_calls == 0

    with session_factory() as session:
        job = session.get(AIBatchJob, created_job["job_id"])
        assert job is not None
        assert job.status == "completed"
        assert job.processed_count == 2
        assert job.generated_count == 2
        assert job.failed_count == 0


def test_recover_stale_running_job_requeues_job_and_item(client, auth_headers, session_factory):
    compare_run_id = _create_compare_run(client, auth_headers)

    from app.services import ai_batch_jobs as ai_batch_job_service

    with session_factory() as session:
        created_job = ai_batch_job_service.create_compare_run_ai_batch_job(
            session,
            compare_run_id=compare_run_id,
            actor_user_id=1,
            force_regenerate=False,
        )
        session.commit()

        job = session.get(AIBatchJob, created_job["job_id"])
        assert job is not None
        first_item = session.scalar(
            select(AIBatchJobItem)
            .where(AIBatchJobItem.job_id == job.id)
            .order_by(AIBatchJobItem.id)
        )
        assert first_item is not None

        stale_time = utcnow() - timedelta(minutes=20)
        job.status = "running"
        job.last_heartbeat_at = stale_time
        first_item.status = "running"
        first_item.last_heartbeat_at = stale_time
        session.add(job)
        session.add(first_item)
        session.commit()

    with session_factory() as session:
        recovered_count = ai_batch_job_service.requeue_stale_ai_batch_jobs(
            session,
            item_stale_after=timedelta(minutes=5),
            job_stale_after=timedelta(minutes=10),
        )
        session.commit()

        job = session.get(AIBatchJob, created_job["job_id"])
        first_item = session.scalar(
            select(AIBatchJobItem)
            .where(AIBatchJobItem.job_id == created_job["job_id"])
            .order_by(AIBatchJobItem.id)
        )

    assert recovered_count == 1
    assert job is not None and job.status == "queued"
    assert first_item is not None and first_item.status == "queued"


def test_ai_batch_worker_wake_triggers_processing(monkeypatch, session_factory):
    from app.services.ai_batch_worker import AIBatchWorker

    processed = Event()

    def fake_process_next_job(next_session_factory, *, concurrency):
        assert next_session_factory is session_factory
        assert concurrency == 2
        processed.set()
        return False

    worker = AIBatchWorker(
        session_factory,
        concurrency=2,
        poll_interval_seconds=60,
        process_next_job=fake_process_next_job,
    )

    worker.start()
    worker.wake()

    assert processed.wait(timeout=2), "worker did not process after wake()"

    worker.stop()


def test_create_app_can_disable_ai_batch_worker(session_factory):
    application = create_app(session_factory=session_factory, start_ai_worker=False)

    with TestClient(application):
        assert getattr(application.state, "ai_batch_worker", None) is None
