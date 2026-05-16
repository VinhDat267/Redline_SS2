from io import BytesIO

from docx import Document as DocxDocument
from sqlalchemy import select

from app.models import AIReviewDraft, ChangeItem
from app.services.llm_adapter import NormalizedAIReviewDraft


def _build_compare_docx(requirement_lines: list[str]) -> bytes:
    document = DocxDocument()
    document.add_paragraph("Requirements", style="Heading 1")
    for requirement_line in requirement_lines:
        document.add_paragraph(requirement_line)

    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _create_compare_run(client, auth_headers) -> int:
    project_response = client.post(
        "/api/v1/projects",
        json={"name": "AI Draft API Project", "description": "AI draft endpoint coverage"},
        headers=auth_headers,
    )
    assert project_response.status_code == 201
    project_id = project_response.json()["data"]["id"]

    document_response = client.post(
        f"/api/v1/projects/{project_id}/documents",
        json={
            "title": "AI Draft API Spec",
            "document_type": "SPEC",
            "description": "AI draft target document",
        },
        headers=auth_headers,
    )
    assert document_response.status_code == 201
    document_id = document_response.json()["data"]["id"]

    source_response = client.post(
        f"/api/v1/documents/{document_id}/versions",
        files={
            "file": (
                "ai-draft-source.docx",
                _build_compare_docx(
                    [
                        "The system shall support login.",
                        "The system shall write audit logs.",
                    ]
                ),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
        data={"version_label": "v1.0", "notes": "AI draft source"},
        headers=auth_headers,
    )
    assert source_response.status_code == 201
    source_version_id = source_response.json()["data"]["id"]

    target_response = client.post(
        f"/api/v1/documents/{document_id}/versions",
        files={
            "file": (
                "ai-draft-target.docx",
                _build_compare_docx(
                    [
                        "The system shall support secure login.",
                        "The system shall write tamper-proof audit logs.",
                    ]
                ),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
        data={"version_label": "v1.1", "notes": "AI draft target"},
        headers=auth_headers,
    )
    assert target_response.status_code == 201
    target_version_id = target_response.json()["data"]["id"]

    assert client.post(f"/api/v1/document-versions/{source_version_id}/parse", headers=auth_headers).status_code == 200
    assert client.post(f"/api/v1/document-versions/{target_version_id}/parse", headers=auth_headers).status_code == 200

    compare_response = client.post(
        f"/api/v1/documents/{document_id}/compare-runs",
        json={"source_version_id": source_version_id, "target_version_id": target_version_id},
        headers=auth_headers,
    )
    assert compare_response.status_code == 201
    return compare_response.json()["data"]["id"]


def test_regenerate_keeps_previous_draft_when_new_attempt_fails(
    client,
    auth_headers,
    session_factory,
    monkeypatch,
):
    compare_run_id = _create_compare_run(client, auth_headers)

    with session_factory() as session:
        change_item = session.scalar(
            select(ChangeItem)
            .where(ChangeItem.compare_run_id == compare_run_id)
            .order_by(ChangeItem.id)
        )
        assert change_item is not None
        session.add(
            AIReviewDraft(
                change_item_id=change_item.id,
                suggested_assignee_user_id=1,
                recommended_review_status="in_review",
                explanation="Existing saved explanation",
                risk_level="medium",
                draft_comment="Existing saved comment",
                suggested_checks="Existing saved checks",
                confidence=0.77,
                generation_status="generated",
                provider_used="gemini",
                fallback_used=False,
            )
        )
        session.commit()
        change_item_id = change_item.id

    class StubAdapter:
        def generate_ai_review_draft(self, payload):
            return NormalizedAIReviewDraft(
                suggested_assignee_user_id=None,
                recommended_review_status="open",
                explanation="AI draft generation failed.",
                risk_level=None,
                draft_comment=None,
                suggested_checks=None,
                confidence=None,
                generation_status="failed",
                provider_used="openai",
                fallback_used=True,
                error_message="fallback provider unavailable",
            )

    from app.services import ai_review_drafts as ai_review_draft_service

    monkeypatch.setattr(ai_review_draft_service, "get_llm_adapter", lambda: StubAdapter())

    response = client.post(
        f"/api/v1/change-items/{change_item_id}/ai-review-draft/generate",
        json={"force_regenerate": True},
        headers=auth_headers,
    )

    assert response.status_code == 200
    payload = response.json()["data"]
    assert payload["change_item_id"] == change_item_id
    assert payload["ai_review_draft"]["generation_status"] == "failed"
    assert payload["ai_review_draft"]["explanation"] == "Existing saved explanation"
    assert payload["ai_review_draft"]["provider_used"] == "openai"
    assert payload["ai_review_draft"]["fallback_used"] is True
    assert payload["ai_review_draft"]["error_message"] == "fallback provider unavailable"


def test_ai_review_generation_includes_rag_context(
    client,
    auth_headers,
    monkeypatch,
):
    compare_run_id = _create_compare_run(client, auth_headers)

    queue_response = client.get(
        f"/api/v1/compare-runs/{compare_run_id}/change-items",
        headers=auth_headers,
    )
    assert queue_response.status_code == 200
    change_item_id = queue_response.json()["data"][0]["id"]

    captured_payload = {}

    class StubAdapter:
        def generate_ai_review_draft(self, payload):
            captured_payload.update(payload)
            return NormalizedAIReviewDraft(
                suggested_assignee_user_id=None,
                recommended_review_status="open",
                explanation="Legal review summary.",
                risk_level="medium",
                draft_comment="Review the clause update.",
                suggested_checks="Check liability alignment.",
                confidence=0.81,
                generation_status="generated",
                provider_used="stub",
                fallback_used=False,
                error_message=None,
            )

    from app.services import ai_review_drafts as ai_review_draft_service

    monkeypatch.setattr(ai_review_draft_service, "get_llm_adapter", lambda: StubAdapter())

    response = client.post(
        f"/api/v1/change-items/{change_item_id}/ai-review-draft/generate",
        json={"force_regenerate": True},
        headers=auth_headers,
    )

    assert response.status_code == 200
    assert "rag_context" in captured_payload
    assert captured_payload["rag_context"]
    first_rag_item = captured_payload["rag_context"][0]
    assert "block_id" in first_rag_item
    assert "content" in first_rag_item


def test_ai_review_generation_can_disable_rag_context(
    client,
    auth_headers,
    monkeypatch,
):
    compare_run_id = _create_compare_run(client, auth_headers)

    queue_response = client.get(
        f"/api/v1/compare-runs/{compare_run_id}/change-items",
        headers=auth_headers,
    )
    assert queue_response.status_code == 200
    change_item_id = queue_response.json()["data"][0]["id"]

    captured_payload = {}

    class StubAdapter:
        def generate_ai_review_draft(self, payload):
            captured_payload.update(payload)
            return NormalizedAIReviewDraft(
                suggested_assignee_user_id=None,
                recommended_review_status="open",
                explanation="Legal review summary.",
                risk_level="medium",
                draft_comment="Review the clause update.",
                suggested_checks="Check liability alignment.",
                confidence=0.81,
                generation_status="generated",
                provider_used="stub",
                fallback_used=False,
                error_message=None,
            )

    from app.services import ai_review_drafts as ai_review_draft_service

    monkeypatch.setattr(ai_review_draft_service, "get_llm_adapter", lambda: StubAdapter())

    response = client.post(
        f"/api/v1/change-items/{change_item_id}/ai-review-draft/generate",
        json={"force_regenerate": True, "use_rag": False},
        headers=auth_headers,
    )

    assert response.status_code == 200
    assert captured_payload["rag_enabled"] is False
    assert captured_payload["rag_context"] == []
