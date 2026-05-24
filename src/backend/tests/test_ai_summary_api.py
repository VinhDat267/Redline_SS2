import pytest
from io import BytesIO

from docx import Document as DocxDocument
from sqlalchemy import select

from app.models import AIReviewDraft, ChangeItem
from app.services.llm_adapter import NormalizedAIReviewDraft

def _build_compare_docx(requirement_lines: list[str]) -> bytes:
    document = DocxDocument()
    document.add_paragraph("Requirements", style="Heading 1")
    for requirement_line in requirement_lines:
        document.add_paragraph(requirement_line)

    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _create_compare_run(client, auth_headers) -> int:
    project_response = client.post(
        "/api/v1/projects",
        json={"name": "AI Summary API Project", "description": "AI summary endpoint coverage"},
        headers=auth_headers,
    )
    assert project_response.status_code == 201
    project_id = project_response.json()["data"]["id"]

    document_response = client.post(
        f"/api/v1/projects/{project_id}/documents",
        json={
            "title": "AI Summary API Spec",
            "document_type": "SPEC",
            "description": "AI summary target document",
        },
        headers=auth_headers,
    )
    assert document_response.status_code == 201
    document_id = document_response.json()["data"]["id"]

    source_response = client.post(
        f"/api/v1/documents/{document_id}/versions",
        files={
            "file": (
                "ai-summary-source.docx",
                _build_compare_docx(
                    [
                        "The system shall support login.",
                        "The system shall write audit logs.",
                    ]
                ),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
        data={"version_label": "v1.0", "notes": "AI summary source"},
        headers=auth_headers,
    )
    assert source_response.status_code == 201
    source_version_id = source_response.json()["data"]["id"]

    target_response = client.post(
        f"/api/v1/documents/{document_id}/versions",
        files={
            "file": (
                "ai-summary-target.docx",
                _build_compare_docx(
                    [
                        "The system shall support secure login.",
                        "The system shall write tamper-proof audit logs.",
                    ]
                ),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
        data={"version_label": "v1.1", "notes": "AI summary target"},
        headers=auth_headers,
    )
    assert target_response.status_code == 201
    target_version_id = target_response.json()["data"]["id"]

    assert client.post(f"/api/v1/document-versions/{source_version_id}/parse", headers=auth_headers).status_code == 200
    assert client.post(f"/api/v1/document-versions/{target_version_id}/parse", headers=auth_headers).status_code == 200

    compare_response = client.post(
        f"/api/v1/documents/{document_id}/compare-runs",
        json={"source_version_id": source_version_id, "target_version_id": target_version_id},
        headers=auth_headers,
    )
    assert compare_response.status_code == 201
    return compare_response.json()["data"]["id"]


def _target_version_id_for_compare_run(session_factory, compare_run_id: int) -> int:
    with session_factory() as session:
        change_item = session.scalar(
            select(ChangeItem).where(ChangeItem.compare_run_id == compare_run_id).order_by(ChangeItem.id)
        )
        assert change_item is not None
        return change_item.compare_run.target_version_id


def test_generate_compare_run_ai_summary(client, auth_headers):
    compare_run_id = _create_compare_run(client, auth_headers)

    response = client.post(
        f"/api/v1/compare-runs/{compare_run_id}/ai-summary-drafts/generate",
        headers=auth_headers,
    )

    assert response.status_code == 200
    payload = response.json()
    assert "summary_text" in payload
    assert isinstance(payload["summary_text"], str)
    assert "provider_used" in payload
    assert "fallback_used" in payload
    assert "error_message" in payload


def test_generate_compare_run_ai_summary_rejects_stale_compare_run(client, auth_headers, session_factory):
    compare_run_id = _create_compare_run(client, auth_headers)
    target_version_id = _target_version_id_for_compare_run(session_factory, compare_run_id)
    assert client.post(f"/api/v1/document-versions/{target_version_id}/parse", headers=auth_headers).status_code == 200

    response = client.post(
        f"/api/v1/compare-runs/{compare_run_id}/ai-summary-drafts/generate",
        headers=auth_headers,
    )

    assert response.status_code == 422
    assert "stale" in response.json()["detail"].lower()
