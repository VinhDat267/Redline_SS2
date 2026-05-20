from io import BytesIO

from docx import Document as DocxDocument
from sqlalchemy import select

from app.models import (
    AIReviewDraft,
    ChangeItem,
    ChangeItemRequirementLink,
    Requirement,
    RequirementTestCaseMapping,
    TestCase as CaseModel,
)
from app.services.llm_adapter import (
    NormalizedTraceabilitySuggestion,
    NormalizedTraceabilitySuggestionResult,
)


def _build_compare_docx(requirement_line: str) -> bytes:
    document = DocxDocument()
    document.add_paragraph("Requirements", style="Heading 1")
    document.add_paragraph(requirement_line)

    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _create_parsed_compare_versions(client, auth_headers) -> tuple[int, int, int]:
    project_response = client.post(
        "/api/v1/projects",
        json={"name": "Change Item API Project", "description": "Change item endpoint coverage"},
        headers=auth_headers,
    )
    assert project_response.status_code == 201
    project_id = project_response.json()["data"]["id"]

    document_response = client.post(
        f"/api/v1/projects/{project_id}/documents",
        json={
            "title": "Change Item API Spec",
            "document_type": "SPEC",
            "description": "Change item target document",
        },
        headers=auth_headers,
    )
    assert document_response.status_code == 201
    document_id = document_response.json()["data"]["id"]

    source_response = client.post(
        f"/api/v1/documents/{document_id}/versions",
        files={
            "file": (
                "change-item-source.docx",
                _build_compare_docx("The system shall support login."),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
        data={"version_label": "v1.0", "notes": "Change item source"},
        headers=auth_headers,
    )
    assert source_response.status_code == 201
    source_version_id = source_response.json()["data"]["id"]

    target_response = client.post(
        f"/api/v1/documents/{document_id}/versions",
        files={
            "file": (
                "change-item-target.docx",
                _build_compare_docx("The system shall support secure login."),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
        data={"version_label": "v1.1", "notes": "Change item target"},
        headers=auth_headers,
    )
    assert target_response.status_code == 201
    target_version_id = target_response.json()["data"]["id"]

    parse_source_response = client.post(
        f"/api/v1/document-versions/{source_version_id}/parse",
        headers=auth_headers,
    )
    parse_target_response = client.post(
        f"/api/v1/document-versions/{target_version_id}/parse",
        headers=auth_headers,
    )
    assert parse_source_response.status_code == 200
    assert parse_target_response.status_code == 200

    return document_id, source_version_id, target_version_id


def _create_compare_run(client, auth_headers) -> int:
    document_id, source_version_id, target_version_id = _create_parsed_compare_versions(client, auth_headers)
    response = client.post(
        f"/api/v1/documents/{document_id}/compare-runs",
        json={"source_version_id": source_version_id, "target_version_id": target_version_id},
        headers=auth_headers,
    )
    assert response.status_code == 201
    return response.json()["data"]["id"]


def _create_compare_run_with_requirement_links(client, auth_headers, session_factory) -> int:
    compare_run_id = _create_compare_run(client, auth_headers)

    with session_factory() as session:
        change_item = session.scalar(
            select(ChangeItem).where(ChangeItem.compare_run_id == compare_run_id).order_by(ChangeItem.id)
        )
        assert change_item is not None

        requirement = Requirement(
            document_id=change_item.source_version.document_id,
            requirement_code="REQ-LOGIN-001",
            title="Secure login",
            description="The system shall support secure login.",
            source_section=change_item.section_title,
            status="active",
        )
        test_case = CaseModel(
            project_id=change_item.source_version.document.project_id,
            test_case_code="TC-LOGIN-001",
            title="Verify secure login",
            description="Regression coverage for secure login.",
            priority="high",
            status="ready",
        )
        session.add_all([requirement, test_case])
        session.flush()

        session.add(
            RequirementTestCaseMapping(
                requirement_id=requirement.id,
                test_case_id=test_case.id,
                notes="Security traceability",
            )
        )
        session.add(
            ChangeItemRequirementLink(
                change_item_id=change_item.id,
                requirement_id=requirement.id,
                notes="Compare item linked to secure login requirement",
            )
        )
        session.commit()

        return change_item.id


def _seed_active_requirement(session_factory, change_item_id: int, *, code: str = "REQ-LOGIN-001") -> int:
    with session_factory() as session:
        change_item = session.get(ChangeItem, change_item_id)
        assert change_item is not None
        requirement = Requirement(
            document_id=change_item.source_version.document_id,
            requirement_code=code,
            title="Secure login",
            description="The system shall support secure login.",
            source_section=change_item.section_title,
            status="active",
        )
        session.add(requirement)
        session.commit()
        return requirement.id


def test_get_change_item_returns_review_context_and_impact_aggregate(
    client,
    auth_headers,
    session_factory,
):
    change_item_id = _create_compare_run_with_requirement_links(client, auth_headers, session_factory)

    response = client.get(f"/api/v1/change-items/{change_item_id}", headers=auth_headers)

    assert response.status_code == 200
    payload = response.json()["data"]
    assert payload["id"] == change_item_id
    assert payload["review_status"] == "open"
    assert payload["linked_requirements"][0]["requirement_code"] == "REQ-LOGIN-001"
    assert payload["impacted_tests"][0]["test_case_code"] == "TC-LOGIN-001"
    assert payload["comments"] == []


def test_patch_change_item_updates_review_status_and_assignee(client, auth_headers):
    change_item_id = _create_compare_run(client, auth_headers)

    response = client.patch(
        f"/api/v1/change-items/{change_item_id}",
        json={"review_status": "resolved", "assignee_user_id": 1},
        headers=auth_headers,
    )

    assert response.status_code == 200
    payload = response.json()["data"]
    assert payload["review_status"] == "resolved"
    assert payload["assignee_user_id"] == 1


def test_patch_change_item_can_clear_assignee_and_summary(client, auth_headers):
    change_item_id = _create_compare_run(client, auth_headers)

    seed_response = client.patch(
        f"/api/v1/change-items/{change_item_id}",
        json={
            "review_status": "in_review",
            "assignee_user_id": 1,
            "summary": "Follow up on the secure login wording.",
        },
        headers=auth_headers,
    )
    assert seed_response.status_code == 200

    clear_response = client.patch(
        f"/api/v1/change-items/{change_item_id}",
        json={"assignee_user_id": None, "summary": None},
        headers=auth_headers,
    )

    assert clear_response.status_code == 200
    payload = clear_response.json()["data"]
    assert payload["assignee_user_id"] is None
    assert payload["summary"] is None


def test_post_change_item_comment_persists_review_discussion(client, auth_headers):
    change_item_id = _create_compare_run(client, auth_headers)

    response = client.post(
        f"/api/v1/change-items/{change_item_id}/comments",
        json={"content": "Need a security regression check before resolving this item."},
        headers=auth_headers,
    )

    assert response.status_code == 201
    payload = response.json()["data"]
    assert payload["content"] == "Need a security regression check before resolving this item."

    detail_response = client.get(f"/api/v1/change-items/{change_item_id}", headers=auth_headers)
    assert detail_response.status_code == 200
    detail_payload = detail_response.json()["data"]
    assert len(detail_payload["comments"]) == 1
    assert detail_payload["comments"][0]["content"] == payload["content"]


def test_change_item_detail_includes_provider_and_fallback_metadata(
    client,
    auth_headers,
    session_factory,
):
    change_item_id = _create_compare_run(client, auth_headers)

    with session_factory() as session:
        session.add(
            AIReviewDraft(
                change_item_id=change_item_id,
                suggested_assignee_user_id=1,
                recommended_review_status="in_review",
                explanation="AI-generated explanation",
                risk_level="medium",
                draft_comment="Review the secure login impact.",
                suggested_checks="Run authentication regression tests.",
                confidence=0.8,
                generation_status="generated",
                provider_used="openai",
                fallback_used=True,
                error_message=None,
            )
        )
        session.commit()

    response = client.get(f"/api/v1/change-items/{change_item_id}", headers=auth_headers)

    assert response.status_code == 200
    payload = response.json()["data"]
    assert payload["ai_review_draft"]["provider_used"] == "openai"
    assert payload["ai_review_draft"]["fallback_used"] is True
    assert payload["ai_review_draft"]["error_message"] is None


def test_change_item_detail_returns_requirement_specific_test_mappings(
    client,
    auth_headers,
    session_factory,
):
    change_item_id = _create_compare_run(client, auth_headers)

    with session_factory() as session:
        change_item = session.get(ChangeItem, change_item_id)
        assert change_item is not None

        requirement_login = Requirement(
            document_id=change_item.source_version.document_id,
            requirement_code="REQ-LOGIN-001",
            title="Secure login",
            description="The system shall support secure login.",
            source_section=change_item.section_title,
            status="active",
        )
        requirement_mfa = Requirement(
            document_id=change_item.source_version.document_id,
            requirement_code="REQ-MFA-002",
            title="Admin MFA",
            description="The system shall enforce MFA for admin users.",
            source_section=change_item.section_title,
            status="active",
        )
        test_login = CaseModel(
            project_id=change_item.source_version.document.project_id,
            test_case_code="TC-LOGIN-001",
            title="Verify secure login",
            description="Regression coverage for secure login.",
            priority="high",
            status="ready",
        )
        test_mfa = CaseModel(
            project_id=change_item.source_version.document.project_id,
            test_case_code="TC-MFA-002",
            title="Verify admin MFA",
            description="Regression coverage for MFA.",
            priority="high",
            status="ready",
        )
        session.add_all([requirement_login, requirement_mfa, test_login, test_mfa])
        session.flush()

        session.add_all(
            [
                RequirementTestCaseMapping(
                    requirement_id=requirement_login.id,
                    test_case_id=test_login.id,
                    notes="Login traceability",
                ),
                RequirementTestCaseMapping(
                    requirement_id=requirement_mfa.id,
                    test_case_id=test_mfa.id,
                    notes="MFA traceability",
                ),
                ChangeItemRequirementLink(
                    change_item_id=change_item.id,
                    requirement_id=requirement_login.id,
                    notes="Login requirement link",
                ),
                ChangeItemRequirementLink(
                    change_item_id=change_item.id,
                    requirement_id=requirement_mfa.id,
                    notes="MFA requirement link",
                ),
            ]
        )
        session.commit()

    response = client.get(f"/api/v1/change-items/{change_item_id}", headers=auth_headers)

    assert response.status_code == 200
    payload = response.json()["data"]
    linked_requirements = {item["requirement_code"]: item for item in payload["linked_requirements"]}
    assert [item["test_case_code"] for item in linked_requirements["REQ-LOGIN-001"]["mapped_test_cases"]] == [
        "TC-LOGIN-001"
    ]
    assert [item["test_case_code"] for item in linked_requirements["REQ-MFA-002"]["mapped_test_cases"]] == [
        "TC-MFA-002"
    ]
    assert sorted(item["test_case_code"] for item in payload["impacted_tests"]) == [
        "TC-LOGIN-001",
        "TC-MFA-002",
    ]


def test_manual_requirement_link_ignores_client_supplied_ai_link_type(
    client,
    auth_headers,
    session_factory,
):
    change_item_id = _create_compare_run(client, auth_headers)
    requirement_id = _seed_active_requirement(session_factory, change_item_id)

    response = client.post(
        f"/api/v1/change-items/{change_item_id}/requirement-links",
        json={
            "requirement_id": requirement_id,
            "notes": "Manual traceability link",
            "link_type": "ai_suggested",
        },
        headers=auth_headers,
    )

    assert response.status_code == 201
    payload = response.json()["data"]
    assert payload["linked_requirements"][0]["link_type"] == "manual"


def test_requirement_link_with_requirement_from_inaccessible_project_returns_404(
    client,
    auth_headers,
    register_user,
):
    change_item_id = _create_compare_run(client, auth_headers)
    outsider = register_user(email="traceability-outsider@example.com", display_name="Traceability Outsider")

    project_response = client.post(
        "/api/v1/projects",
        json={"name": "Outsider Traceability Project", "description": "Different tenant"},
        headers=outsider["headers"],
    )
    assert project_response.status_code == 201
    outsider_project_id = project_response.json()["data"]["id"]
    document_response = client.post(
        f"/api/v1/projects/{outsider_project_id}/documents",
        json={"title": "Outsider requirements", "document_type": "SPEC"},
        headers=outsider["headers"],
    )
    assert document_response.status_code == 201
    outsider_document_id = document_response.json()["data"]["id"]
    requirement_response = client.post(
        f"/api/v1/projects/{outsider_project_id}/requirements",
        json={
            "document_id": outsider_document_id,
            "requirement_code": "REQ-OUTSIDER-001",
            "title": "Outsider-only requirement",
        },
        headers=outsider["headers"],
    )
    assert requirement_response.status_code == 201
    outsider_requirement_id = requirement_response.json()["data"]["id"]

    response = client.post(
        f"/api/v1/change-items/{change_item_id}/requirement-links",
        json={"requirement_id": outsider_requirement_id},
        headers=auth_headers,
    )

    assert response.status_code == 404


def test_ai_traceability_suggestion_acceptance_requires_server_token(
    client,
    auth_headers,
    session_factory,
    monkeypatch,
):
    from app.services import ai_traceability as ai_traceability_service

    change_item_id = _create_compare_run(client, auth_headers)
    requirement_id = _seed_active_requirement(session_factory, change_item_id)

    class FakeTraceabilityAdapter:
        def generate_traceability_suggestions(self, payload):
            return NormalizedTraceabilitySuggestionResult(
                suggestions=[
                    NormalizedTraceabilitySuggestion(
                        requirement_code="REQ-LOGIN-001",
                        title="Secure login",
                        confidence=0.86,
                        rationale="The change strengthens login security.",
                        relevance_type="directly_affected",
                    ),
                    NormalizedTraceabilitySuggestion(
                        requirement_code="REQ-HALLUCINATED-999",
                        title="Hallucinated obligation",
                        confidence=0.99,
                        rationale="Should be discarded because it was not in the prompt.",
                        relevance_type="related",
                    ),
                ],
                provider_used="test-provider",
                fallback_used=False,
                error_message=None,
            )

    monkeypatch.setattr(ai_traceability_service, "get_llm_adapter", lambda: FakeTraceabilityAdapter())

    suggest_response = client.post(
        f"/api/v1/change-items/{change_item_id}/suggest-links",
        headers=auth_headers,
    )

    assert suggest_response.status_code == 200
    suggestion_payload = suggest_response.json()["data"]
    assert suggestion_payload["provider_used"] == "test-provider"
    assert [item["requirement_id"] for item in suggestion_payload["suggestions"]] == [requirement_id]
    suggestion_token = suggestion_payload["suggestions"][0]["suggestion_token"]
    assert suggestion_token

    spoof_response = client.post(
        f"/api/v1/change-items/{change_item_id}/requirement-links/ai-suggested",
        json={"requirement_id": requirement_id, "suggestion_token": "0" * 64},
        headers=auth_headers,
    )
    assert spoof_response.status_code == 400

    accept_response = client.post(
        f"/api/v1/change-items/{change_item_id}/requirement-links/ai-suggested",
        json={"requirement_id": requirement_id, "suggestion_token": suggestion_token},
        headers=auth_headers,
    )

    assert accept_response.status_code == 201
    payload = accept_response.json()["data"]
    assert payload["linked_requirements"][0]["link_type"] == "ai_suggested"
