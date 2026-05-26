import json
from io import BytesIO

import pytest
from docx import Document as DocxDocument
from sqlalchemy import select

from app.models import AIReviewDraft, ChangeItem, CompareRun
from app.services import compare as compare_service


def _build_compare_docx(
    *,
    body_paragraphs: list[tuple[str, str | None]],
    header_paragraphs: list[str] | None = None,
    footer_paragraphs: list[str] | None = None,
    body_tables: list[list[list[str]]] | None = None,
) -> bytes:
    document = DocxDocument()
    for text, style in body_paragraphs:
        paragraph = document.add_paragraph(text)
        if style is not None:
            paragraph.style = style

    for table_rows in body_tables or []:
        table = document.add_table(rows=len(table_rows), cols=len(table_rows[0]))
        for row_index, row_values in enumerate(table_rows):
            for column_index, value in enumerate(row_values):
                table.cell(row_index, column_index).text = value

    section = document.sections[0]
    for text in header_paragraphs or []:
        section.header.add_paragraph(text)
    for text in footer_paragraphs or []:
        section.footer.add_paragraph(text)

    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _create_project_and_document(client, auth_headers, *, project_name: str, document_title: str) -> int:
    project_response = client.post(
        "/api/v1/projects",
        json={"name": project_name, "description": f"{project_name} description"},
        headers=auth_headers,
    )
    assert project_response.status_code == 201
    project_id = project_response.json()["data"]["id"]

    document_response = client.post(
        f"/api/v1/projects/{project_id}/documents",
        json={
            "title": document_title,
            "document_type": "SPEC",
            "description": f"{document_title} document",
        },
        headers=auth_headers,
    )
    assert document_response.status_code == 201
    return document_response.json()["data"]["id"]


def _upload_version(
    client,
    auth_headers,
    *,
    document_id: int,
    version_label: str,
    file_name: str,
    payload: bytes,
) -> int:
    response = client.post(
        f"/api/v1/documents/{document_id}/versions",
        files={
            "file": (
                file_name,
                payload,
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
        data={"version_label": version_label, "notes": version_label},
        headers=auth_headers,
    )
    assert response.status_code == 201
    return response.json()["data"]["id"]


def _parse_version(client, auth_headers, version_id: int) -> None:
    response = client.post(f"/api/v1/document-versions/{version_id}/parse", headers=auth_headers)
    assert response.status_code == 200


def _create_parsed_compare_versions(
    client,
    auth_headers,
    *,
    source_payload: bytes,
    target_payload: bytes,
    project_name: str = "Compare API Project",
    document_title: str = "Compare API Spec",
) -> tuple[int, int, int]:
    document_id = _create_project_and_document(
        client,
        auth_headers,
        project_name=project_name,
        document_title=document_title,
    )
    source_version_id = _upload_version(
        client,
        auth_headers,
        document_id=document_id,
        version_label="v1.0",
        file_name="compare-source.docx",
        payload=source_payload,
    )
    target_version_id = _upload_version(
        client,
        auth_headers,
        document_id=document_id,
        version_label="v1.1",
        file_name="compare-target.docx",
        payload=target_payload,
    )
    _parse_version(client, auth_headers, source_version_id)
    _parse_version(client, auth_headers, target_version_id)
    return document_id, source_version_id, target_version_id


def _create_compare_run(
    client,
    auth_headers,
    *,
    source_payload: bytes,
    target_payload: bytes,
    project_name: str = "Compare API Project",
    document_title: str = "Compare API Spec",
) -> tuple[dict[str, object], list[dict[str, object]]]:
    document_id, source_version_id, target_version_id = _create_parsed_compare_versions(
        client,
        auth_headers,
        source_payload=source_payload,
        target_payload=target_payload,
        project_name=project_name,
        document_title=document_title,
    )
    response = client.post(
        f"/api/v1/documents/{document_id}/compare-runs",
        json={"source_version_id": source_version_id, "target_version_id": target_version_id},
        headers=auth_headers,
    )
    assert response.status_code == 201
    compare_run = response.json()["data"]
    queue_response = client.get(
        f"/api/v1/compare-runs/{compare_run['id']}/change-items",
        headers=auth_headers,
    )
    assert queue_response.status_code == 200
    return compare_run, queue_response.json()["data"]


def test_create_compare_run_builds_change_items_from_two_parsed_versions(client, auth_headers):
    compare_run, queue = _create_compare_run(
        client,
        auth_headers,
        source_payload=_build_compare_docx(
            body_paragraphs=[
                ("Requirements", "Heading 1"),
                ("The system shall support login.", None),
            ]
        ),
        target_payload=_build_compare_docx(
            body_paragraphs=[
                ("Requirements", "Heading 1"),
                ("The system shall support secure login.", None),
            ]
        ),
    )

    assert compare_run["compare_status"] == "completed"
    assert compare_run["summary"]["total_changes"] == 1
    assert queue[0]["change_type"] == "modified"


def test_compare_preserves_later_exact_matches_when_target_inserts_midstream_paragraph(client, auth_headers):
    compare_run, queue = _create_compare_run(
        client,
        auth_headers,
        source_payload=_build_compare_docx(
            body_paragraphs=[
                ("Requirements", "Heading 1"),
                ("The system shall support login.", None),
                ("The system shall write audit logs.", None),
            ]
        ),
        target_payload=_build_compare_docx(
            body_paragraphs=[
                ("Requirements", "Heading 1"),
                ("The system shall support login.", None),
                ("The system shall support MFA.", None),
                ("The system shall write audit logs.", None),
            ]
        ),
        project_name="Compare Text Alignment Project",
        document_title="Compare Text Alignment Document",
    )

    assert compare_run["summary"] == {"total_changes": 1, "added": 1, "removed": 0, "modified": 0}
    assert len(queue) == 1
    assert queue[0]["change_type"] == "added"
    assert queue[0]["surface_type"] == "body"
    assert queue[0]["new_content"] == "The system shall support MFA."


def test_compare_keeps_header_changes_partitioned_away_from_body_content(client, auth_headers):
    compare_run, queue = _create_compare_run(
        client,
        auth_headers,
        source_payload=_build_compare_docx(
            body_paragraphs=[
                ("Requirements", "Heading 1"),
                ("The body content stays unchanged.", None),
            ],
            header_paragraphs=["Release 1.0"],
        ),
        target_payload=_build_compare_docx(
            body_paragraphs=[
                ("Requirements", "Heading 1"),
                ("The body content stays unchanged.", None),
            ],
            header_paragraphs=["Release 1.1"],
        ),
        project_name="Compare Header Partition Project",
        document_title="Compare Header Partition Document",
    )

    assert compare_run["summary"] == {"total_changes": 1, "added": 0, "removed": 0, "modified": 1}
    assert len(queue) == 1
    assert queue[0]["change_type"] == "modified"
    assert queue[0]["surface_type"] == "header"
    assert queue[0]["old_content"] == "Release 1.0"
    assert queue[0]["new_content"] == "Release 1.1"


def test_compare_aligns_table_rows_by_requirement_key_and_exposes_structured_diff(
    client,
    auth_headers,
):
    compare_run, queue = _create_compare_run(
        client,
        auth_headers,
        source_payload=_build_compare_docx(
            body_paragraphs=[("Requirements", "Heading 1")],
            body_tables=[
                [
                    ["Requirement ID", "Title"],
                    ["REQ-001", "Login"],
                    ["REQ-002", "Register"],
                ]
            ],
        ),
        target_payload=_build_compare_docx(
            body_paragraphs=[("Requirements", "Heading 1")],
            body_tables=[
                [
                    ["Requirement ID", "Title"],
                    ["REQ-000", "Scope"],
                    ["REQ-001", "Secure Login"],
                    ["REQ-002", "Register"],
                ]
            ],
        ),
        project_name="Compare Table Alignment Project",
        document_title="Compare Table Alignment Document",
    )

    assert compare_run["summary"] == {"total_changes": 2, "added": 1, "removed": 0, "modified": 1}
    assert len(queue) == 2

    added_item = next(item for item in queue if item["change_type"] == "added")
    modified_item = next(item for item in queue if item["change_type"] == "modified")

    assert added_item["table_key"] is not None
    assert added_item["new_content"] == "Requirement ID: REQ-000 || Title: Scope"
    assert modified_item["old_content"] == "Requirement ID: REQ-001 || Title: Login"
    assert modified_item["new_content"] == "Requirement ID: REQ-001 || Title: Secure Login"

    detail_response = client.get(
        f"/api/v1/change-items/{modified_item['id']}",
        headers=auth_headers,
    )
    assert detail_response.status_code == 200
    detail_payload = detail_response.json()["data"]
    assert detail_payload["structured_diff_json"] is not None

    structured_diff = json.loads(detail_payload["structured_diff_json"])
    assert structured_diff["changed_columns"] == [
        {
            "column_key": "title",
            "header_text": "Title",
            "old_value": "Login",
            "new_value": "Secure Login",
        }
    ]


def test_compare_marks_warning_status_when_table_alignment_has_to_fallback(client, auth_headers):
    compare_run, queue = _create_compare_run(
        client,
        auth_headers,
        source_payload=_build_compare_docx(
            body_paragraphs=[("Execution Matrix", "Heading 1")],
            body_tables=[
                [
                    ["Bucket", "State"],
                    ["A", "open"],
                    ["A", "open"],
                ]
            ],
        ),
        target_payload=_build_compare_docx(
            body_paragraphs=[("Execution Matrix", "Heading 1")],
            body_tables=[
                [
                    ["Bucket", "State"],
                    ["A", "open"],
                    ["A", "open"],
                    ["A", "open"],
                ]
            ],
        ),
        project_name="Compare Warning Project",
        document_title="Compare Warning Document",
    )

    assert compare_run["compare_status"] == "completed_with_warnings"
    assert compare_run["warning_count"] >= 1
    assert len(compare_run["warnings"]) >= 1
    assert len(queue) == 5


def test_compare_queue_sorts_by_target_anchor_then_source_anchor(client, auth_headers):
    compare_run, queue = _create_compare_run(
        client,
        auth_headers,
        source_payload=_build_compare_docx(
            body_paragraphs=[
                ("Requirements", "Heading 1"),
                ("The system shall support login.", None),
            ]
        ),
        target_payload=_build_compare_docx(
            body_paragraphs=[
                ("Requirements", "Heading 1"),
                ("The system shall support secure login.", None),
            ]
        ),
        project_name="Compare Sort Project",
        document_title="Compare Sort Document",
    )

    detail_response = client.get(f"/api/v1/compare-runs/{compare_run['id']}", headers=auth_headers)
    assert detail_response.status_code == 200

    sort_keys = [item["sort_key"] for item in queue]
    assert sort_keys == sorted(sort_keys)


def test_compare_queue_supports_paginated_filtered_response(client, auth_headers):
    compare_run, _queue = _create_compare_run(
        client,
        auth_headers,
        source_payload=_build_compare_docx(
            body_paragraphs=[
                ("I like apples.", None),
                ("I like bananas.", None),
            ]
        ),
        target_payload=_build_compare_docx(
            body_paragraphs=[
                ("I like apples and pears.", None),
                ("I like bananas.", None),
                ("I like grapes.", None),
                ("I like oranges.", None),
            ]
        ),
        project_name="Compare Pagination Project",
        document_title="Compare Pagination Document",
    )

    response = client.get(
        f"/api/v1/compare-runs/{compare_run['id']}/change-items",
        params={"limit": 1, "offset": 1, "change_type": "added"},
        headers=auth_headers,
    )

    assert response.status_code == 200
    payload = response.json()["data"]
    assert payload["total_count"] == 2
    assert payload["limit"] == 1
    assert payload["offset"] == 1
    assert payload["review_counts"] == {
        "total": 2,
        "open": 2,
        "in_review": 0,
        "resolved": 0,
    }
    assert len(payload["items"]) == 1
    assert payload["items"][0]["change_type"] == "added"
    assert payload["items"][0]["new_content"] == "I like oranges."

    search_response = client.get(
        f"/api/v1/compare-runs/{compare_run['id']}/change-items",
        params={"limit": 10, "search": "pears"},
        headers=auth_headers,
    )

    assert search_response.status_code == 200
    search_payload = search_response.json()["data"]
    assert search_payload["total_count"] == 1
    assert search_payload["items"][0]["change_type"] == "modified"


def test_compare_run_detail_does_not_hydrate_full_change_queue(client, auth_headers, monkeypatch):
    compare_run, _queue = _create_compare_run(
        client,
        auth_headers,
        source_payload=_build_compare_docx(
            body_paragraphs=[
                ("I like apples.", None),
                ("I like bananas.", None),
            ]
        ),
        target_payload=_build_compare_docx(
            body_paragraphs=[
                ("I like apples and pears.", None),
                ("I like bananas.", None),
                ("I like grapes.", None),
            ]
        ),
        project_name="Compare Detail Pagination Project",
        document_title="Compare Detail Pagination Document",
    )

    def fail_if_full_queue_loaded(*_args, **_kwargs):
        raise AssertionError("Compare detail should not hydrate the full change queue")

    monkeypatch.setattr(compare_service, "list_compare_run_change_items", fail_if_full_queue_loaded)

    response = client.get(f"/api/v1/compare-runs/{compare_run['id']}", headers=auth_headers)

    assert response.status_code == 200
    payload = response.json()["data"]
    assert payload["selected_change_item_id"] is not None
    assert payload["summary"]["total_changes"] == 2


def test_compare_queue_includes_ai_generation_status(client, auth_headers, session_factory):
    compare_run, _queue = _create_compare_run(
        client,
        auth_headers,
        source_payload=_build_compare_docx(
            body_paragraphs=[
                ("Requirements", "Heading 1"),
                ("The system shall support login.", None),
            ]
        ),
        target_payload=_build_compare_docx(
            body_paragraphs=[
                ("Requirements", "Heading 1"),
                ("The system shall support secure login.", None),
            ]
        ),
        project_name="Compare AI Queue Project",
        document_title="Compare AI Queue Document",
    )

    with session_factory() as session:
        change_item = session.scalar(
            select(ChangeItem).where(ChangeItem.compare_run_id == compare_run["id"]).order_by(ChangeItem.id)
        )
        assert change_item is not None
        session.add(
            AIReviewDraft(
                change_item_id=change_item.id,
                suggested_assignee_user_id=1,
                recommended_review_status="in_review",
                explanation="AI draft already generated.",
                risk_level="medium",
                draft_comment="Review the authentication change.",
                suggested_checks="Run authentication regression tests.",
                confidence=0.8,
                generation_status="generated",
                provider_used="gemini",
                fallback_used=False,
            )
        )
        session.commit()

    queue_response = client.get(
        f"/api/v1/compare-runs/{compare_run['id']}/change-items",
        headers=auth_headers,
    )
    assert queue_response.status_code == 200
    payload = queue_response.json()["data"]
    assert payload[0]["ai_generation_status"] == "generated"
    assert payload[0]["has_ai_review_draft"] is True


def test_compare_run_persists_failed_status_when_engine_raises_after_run_creation(
    client,
    auth_headers,
    session_factory,
    monkeypatch,
):
    document_id, source_version_id, target_version_id = _create_parsed_compare_versions(
        client,
        auth_headers,
        source_payload=_build_compare_docx(
            body_paragraphs=[
                ("Requirements", "Heading 1"),
                ("The system shall support login.", None),
            ]
        ),
        target_payload=_build_compare_docx(
            body_paragraphs=[
                ("Requirements", "Heading 1"),
                ("The system shall support secure login.", None),
            ]
        ),
        project_name="Compare Failure Project",
        document_title="Compare Failure Document",
    )

    def _explode(*args, **kwargs):
        raise RuntimeError("compare engine exploded")

    monkeypatch.setattr(compare_service, "_build_change_items", _explode)

    with session_factory() as session:
        with pytest.raises(RuntimeError, match="compare engine exploded"):
            compare_service.create_compare_run(
                session,
                document_id=document_id,
                source_version_id=source_version_id,
                target_version_id=target_version_id,
                actor_user_id=1,
            )

    with session_factory() as session:
        compare_run = session.scalar(select(CompareRun).order_by(CompareRun.id.desc()))
        assert compare_run is not None
        assert compare_run.compare_status == "failed"
        assert compare_run.error_message == "compare engine exploded"
        assert compare_run.completed_at is not None


def test_compare_routes_require_project_membership(client, auth_headers, register_user):
    source_payload = _build_compare_docx(
        body_paragraphs=[
            ("Requirements", "Heading 1"),
            ("The system shall support login.", None),
        ]
    )
    target_payload = _build_compare_docx(
        body_paragraphs=[
            ("Requirements", "Heading 1"),
            ("The system shall support secure login.", None),
        ]
    )
    document_id, source_version_id, target_version_id = _create_parsed_compare_versions(
        client,
        auth_headers,
        source_payload=source_payload,
        target_payload=target_payload,
        project_name="Protected Compare Project",
        document_title="Protected Compare Document",
    )

    compare_response = client.post(
        f"/api/v1/documents/{document_id}/compare-runs",
        json={"source_version_id": source_version_id, "target_version_id": target_version_id},
        headers=auth_headers,
    )
    assert compare_response.status_code == 201
    compare_run_id = compare_response.json()["data"]["id"]

    outsider = register_user(email="compare-outsider@example.com", display_name="Compare Outsider")

    outsider_create_response = client.post(
        f"/api/v1/documents/{document_id}/compare-runs",
        json={"source_version_id": source_version_id, "target_version_id": target_version_id},
        headers=outsider["headers"],
    )
    assert outsider_create_response.status_code == 404

    outsider_detail_response = client.get(
        f"/api/v1/compare-runs/{compare_run_id}",
        headers=outsider["headers"],
    )
    assert outsider_detail_response.status_code == 404

    outsider_queue_response = client.get(
        f"/api/v1/compare-runs/{compare_run_id}/change-items",
        headers=outsider["headers"],
    )
    assert outsider_queue_response.status_code == 404


def test_compare_preserves_all_change_items_even_when_ai_review_cap_is_lower(
    client,
    auth_headers,
    monkeypatch,
):
    from app.core.config import settings

    monkeypatch.setattr(settings, "ai_review_max_items_per_job", 2)

    compare_run, queue = _create_compare_run(
        client,
        auth_headers,
        source_payload=_build_compare_docx(
            body_paragraphs=[
                ("I like apples.", None),
                ("I like bananas.", None),
            ]
        ),
        target_payload=_build_compare_docx(
            body_paragraphs=[
                ("I like apples and pears.", None), # modified
                ("I like bananas.", None),
                ("I like grapes.", None), # added
                ("I like oranges.", None), # added
            ]
        ),
        project_name="Compare Capping Project",
        document_title="Compare Capping Document",
    )

    assert compare_run["compare_status"] == "completed"
    assert compare_run["warning_count"] == 0
    assert compare_run["summary"] == {"total_changes": 3, "added": 2, "removed": 0, "modified": 1}
    assert len(queue) == 3
    change_types = [item["change_type"] for item in queue]
    assert change_types == ["modified", "added", "added"]
