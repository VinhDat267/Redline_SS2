import hashlib
import json
from pathlib import Path
from zipfile import ZipFile

import pytest
import fitz
from docx import Document as DocxDocument
from sqlalchemy import select

from app.core.config import settings
from app.models import (
    Document,
    DocumentBlock,
    DocumentParseRun,
    DocumentSurface,
    DocumentVersion,
    Project,
)
from app.services import document_parser
from app.services import document_pdf_parser


def _save_docx(path: Path, paragraphs: list[tuple[str, str | None]]) -> None:
    document = DocxDocument()
    for text, style in paragraphs:
        paragraph = document.add_paragraph(text)
        if style is not None:
            paragraph.style = style

    document.save(path)


def _save_text_pdf(path: Path, pages: list[str]) -> None:
    document = fitz.open()
    for page_text in pages:
        page = document.new_page(width=612, height=792)
        y = 72
        for line in page_text.splitlines():
            page.insert_text((72, y), line, fontsize=11)
            y += 18
    document.save(path)
    document.close()


def _save_blank_pdf(path: Path) -> None:
    document = fitz.open()
    page = document.new_page(width=612, height=792)
    page.draw_rect(fitz.Rect(72, 72, 220, 220), color=(0, 0, 0), fill=(0, 0, 0))
    document.save(path)
    document.close()


def _create_document_version(session_factory, file_path: Path) -> int:
    with session_factory() as session:
        project = Project(name="Parser Project", description="Parser runtime tests")
        document = Document(
            project=project,
            title="Parser Document",
            document_type="SPEC",
            description="Parser Document",
        )
        version = DocumentVersion(
            document=document,
            version_label="v1.0",
            file_name=file_path.name,
            file_path=str(file_path),
            parse_status="pending",
        )
        session.add_all([project, document, version])
        session.commit()
        session.refresh(version)
        return version.id


def _read_zip_text(path: Path, part_name: str) -> str:
    with ZipFile(path) as archive:
        return archive.read(part_name).decode("utf-8")


def _replace_zip_part(path: Path, part_name: str, content: str) -> None:
    replacement_path = path.with_suffix(".replacement.docx")
    with ZipFile(path, "r") as source, ZipFile(replacement_path, "w") as target:
        for item in source.infolist():
            data = content.encode("utf-8") if item.filename == part_name else source.read(item.filename)
            target.writestr(item, data)
    replacement_path.replace(path)


def _inject_textbox(path: Path, text: str) -> None:
    document_xml = _read_zip_text(path, "word/document.xml")
    textbox_xml = (
        "<w:p><w:r><w:pict><v:shape xmlns:v=\"urn:schemas-microsoft-com:vml\">"
        "<v:textbox><w:txbxContent><w:p><w:r>"
        f"<w:t>{text}</w:t>"
        "</w:r></w:p></w:txbxContent></v:textbox></v:shape></w:pict></w:r></w:p>"
    )
    _replace_zip_part(path, "word/document.xml", document_xml.replace("<w:sectPr", f"{textbox_xml}<w:sectPr", 1))


def _inject_content_control(path: Path, text: str) -> None:
    document_xml = _read_zip_text(path, "word/document.xml")
    content_control_xml = (
        "<w:sdt><w:sdtContent><w:p><w:r>"
        f"<w:t>{text}</w:t>"
        "</w:r></w:p></w:sdtContent></w:sdt>"
    )
    _replace_zip_part(
        path,
        "word/document.xml",
        document_xml.replace("<w:sectPr", f"{content_control_xml}<w:sectPr", 1),
    )


def test_parse_document_version_persists_blocks_with_section_tracking(session_factory, tmp_path: Path):
    file_path = tmp_path / "parser-sample.docx"
    _save_docx(
        file_path,
        [
            ("Introduction", "Heading 1"),
            ("System overview", None),
            ("1. Create account", "List Number"),
            ("   ", None),
            ("2.1 API", None),
            ("The API accepts JSON payloads.", None),
        ],
    )
    version_id = _create_document_version(session_factory, file_path)

    with session_factory() as session:
        version = session.get(DocumentVersion, version_id)
        assert version is not None

        parsed_version = document_parser.parse_document_version(session, version)
        parse_runs = list(
            session.scalars(
                select(DocumentParseRun)
                .where(DocumentParseRun.document_version_id == version_id)
                .order_by(DocumentParseRun.id)
            )
        )
        surfaces = list(
            session.scalars(
                select(DocumentSurface)
                .where(DocumentSurface.parse_run_id == parsed_version.active_parse_run_id)
                .order_by(DocumentSurface.logical_order_index)
            )
        )
        blocks = list(
            session.scalars(
                select(DocumentBlock)
                .where(DocumentBlock.document_version_id == version_id)
                .order_by(DocumentBlock.order_index)
            )
        )

    assert parsed_version.parse_status == "parsed"
    assert parsed_version.active_parse_run_id is not None
    assert [parse_run.status for parse_run in parse_runs] == ["parsed"]
    assert len(surfaces) == 1
    assert surfaces[0].surface_type == "body"
    assert surfaces[0].surface_key == "body-main"
    assert [block.block_type for block in blocks] == [
        "heading",
        "paragraph",
        "list_item",
        "heading",
        "paragraph",
    ]
    assert [block.order_index for block in blocks] == [0, 1, 2, 3, 4]
    assert [block.surface_order_index for block in blocks] == [0, 1, 2, 3, 4]
    assert {block.parse_run_id for block in blocks} == {parsed_version.active_parse_run_id}
    assert {block.surface_id for block in blocks} == {surfaces[0].id}
    assert blocks[0].section_title == "Introduction"
    assert blocks[0].heading_level == 1
    assert blocks[1].section_title == "Introduction"
    assert blocks[2].section_title == "Introduction"
    assert blocks[2].heading_level is None
    assert blocks[3].section_title == "2.1 API"
    assert blocks[3].heading_level == 2
    assert blocks[4].section_title == "2.1 API"
    assert blocks[0].block_key == (
        f"blk-0000-heading-{hashlib.sha1('Introduction'.encode('utf-8')).hexdigest()[:10]}"
    )
    assert blocks[2].block_key == (
        f"blk-0002-list_item-{hashlib.sha1('1. Create account'.encode('utf-8')).hexdigest()[:10]}"
    )

    parsed_snapshot = json.loads(parsed_version.parsed_snapshot or "{}")
    assert parsed_snapshot["parser_version"] == "v1"
    assert parsed_snapshot["document_version_id"] == version_id
    assert parsed_snapshot["active_parse_run_id"] == parsed_version.active_parse_run_id
    assert parsed_snapshot["total_surfaces"] == 1
    assert parsed_snapshot["total_blocks"] == 5
    assert parsed_snapshot["counts_by_surface_type"] == {"body": 1}
    assert parsed_snapshot["counts_by_block_type"] == {
        "heading": 2,
        "paragraph": 2,
        "list_item": 1,
        "table_row": 0,
    }
    assert parsed_snapshot["table_count"] == 0
    assert parsed_snapshot["row_count"] == 0
    assert parsed_snapshot["warning_count"] == 0
    assert parsed_snapshot["warnings"] == []


def test_parse_document_version_persists_parser_truth_when_embeddings_fail(
    monkeypatch,
    session_factory,
    tmp_path: Path,
):
    monkeypatch.setattr(settings, "rag_embedding_provider", "openai_compatible")
    monkeypatch.setattr(settings, "rag_embedding_api_key", None)
    monkeypatch.setattr(settings, "ai_openai_api_key", None)
    monkeypatch.setattr(settings, "rag_embedding_fallback_to_local_hash", False)
    file_path = tmp_path / "embedding-unavailable.docx"
    _save_docx(
        file_path,
        [
            ("Services", "Heading 1"),
            ("The supplier must deliver the services by the milestone date.", None),
        ],
    )
    version_id = _create_document_version(session_factory, file_path)

    with session_factory() as session:
        version = session.get(DocumentVersion, version_id)
        assert version is not None

        parsed_version = document_parser.parse_document_version(session, version)
        parse_run = session.get(DocumentParseRun, parsed_version.active_parse_run_id)
        blocks = list(
            session.scalars(
                select(DocumentBlock)
                .where(DocumentBlock.parse_run_id == parsed_version.active_parse_run_id)
                .order_by(DocumentBlock.order_index)
            )
        )

    assert parse_run is not None
    assert parsed_version.parse_status == "parsed_with_warnings"
    assert parse_run.status == "parsed_with_warnings"
    assert parse_run.warning_count == 1
    assert [block.normalized_content for block in blocks] == [
        "Services",
        "The supplier must deliver the services by the milestone date.",
    ]
    assert all(block.embedding_provider is None for block in blocks)
    assert all(block.embedding_vector is None for block in blocks)
    assert all(block.embedding_vector_json is None for block in blocks)
    summary = json.loads(parse_run.summary_json or "{}")
    assert summary["warnings"] == [
        "RAG embeddings skipped: OpenAI-compatible embedding provider is not configured."
    ]


def test_parse_pdf_document_version_persists_page_surfaces(session_factory, tmp_path: Path):
    file_path = tmp_path / "contract.pdf"
    _save_text_pdf(
        file_path,
        [
            "\n".join(
                [
                    "1. Definitions",
                    "Agreement means this services agreement.",
                    "(a) Confidential Information includes business terms.",
                    "The receiving party must protect it.",
                ]
            )
        ],
    )
    version_id = _create_document_version(session_factory, file_path)

    with session_factory() as session:
        version = session.get(DocumentVersion, version_id)
        assert version is not None

        parsed_version = document_parser.parse_document_version(session, version)
        parse_run = session.get(DocumentParseRun, parsed_version.active_parse_run_id)
        surfaces = list(
            session.scalars(
                select(DocumentSurface)
                .where(DocumentSurface.parse_run_id == parsed_version.active_parse_run_id)
                .order_by(DocumentSurface.logical_order_index)
            )
        )
        blocks = list(
            session.scalars(
                select(DocumentBlock)
                .where(DocumentBlock.document_version_id == version_id)
                .order_by(DocumentBlock.order_index)
            )
        )

    assert parse_run is not None
    assert parsed_version.parse_status == "parsed"
    assert parse_run.status == "parsed"
    assert surfaces[0].surface_type == "page"
    assert surfaces[0].surface_key == "pdf-page-1"
    assert [block.block_type for block in blocks] == [
        "heading",
        "paragraph",
        "list_item",
        "paragraph",
    ]

    parsed_snapshot = json.loads(parsed_version.parsed_snapshot or "{}")
    assert parsed_snapshot["counts_by_surface_type"] == {"page": 1}
    assert parsed_snapshot["pdf"]["page_count"] == 1
    assert parsed_snapshot["coverage"]["policy_result"] == "pass"


def test_parse_pdf_ocr_warning_sets_active_parse_run(
    monkeypatch,
    session_factory,
    tmp_path: Path,
):
    file_path = tmp_path / "scan.pdf"
    _save_blank_pdf(file_path)
    version_id = _create_document_version(session_factory, file_path)

    def fake_ocr(page, page_index, settings):
        return document_pdf_parser.OcrPageResult(
            text="1. Payment Terms\nInvoices are due within thirty days.",
            average_confidence=91.0,
            retained_token_count=12,
            low_confidence_token_ratio=0.05,
            languages="eng+vie",
            dpi=200,
        )

    monkeypatch.setattr(document_pdf_parser, "run_ocr_for_page", fake_ocr)

    with session_factory() as session:
        version = session.get(DocumentVersion, version_id)
        assert version is not None

        parsed_version = document_parser.parse_document_version(session, version)
        parse_run = session.get(DocumentParseRun, parsed_version.active_parse_run_id)

    assert parse_run is not None
    assert parsed_version.parse_status == "parsed_with_warnings"
    assert parse_run.status == "parsed_with_warnings"
    assert parse_run.warning_count == 1
    summary = json.loads(parse_run.summary_json or "{}")
    assert summary["pdf"]["extraction_modes_by_page"] == {"1": "ocr"}
    assert summary["diagnostics"][0]["metadata"]["confidence"] == 91.0


def test_pdf_quality_failure_keeps_previous_active_parse_run(
    monkeypatch,
    session_factory,
    tmp_path: Path,
):
    success_file_path = tmp_path / "success.docx"
    _save_docx(
        success_file_path,
        [
            ("Overview", "Heading 1"),
            ("The parser should create an active parse run.", None),
        ],
    )
    failure_file_path = tmp_path / "low-confidence.pdf"
    _save_blank_pdf(failure_file_path)
    version_id = _create_document_version(session_factory, success_file_path)

    def fake_ocr(page, page_index, settings):
        return document_pdf_parser.OcrPageResult(
            text="1. Indemnity\nThe supplier shall indemnify the customer.",
            average_confidence=50.0,
            retained_token_count=7,
            low_confidence_token_ratio=0.6,
            languages="eng+vie",
            dpi=200,
        )

    monkeypatch.setattr(document_pdf_parser, "run_ocr_for_page", fake_ocr)

    with session_factory() as session:
        version = session.get(DocumentVersion, version_id)
        assert version is not None

        parsed_version = document_parser.parse_document_version(session, version)
        first_active_parse_run_id = parsed_version.active_parse_run_id
        assert first_active_parse_run_id is not None

        version.file_path = str(failure_file_path)
        version.file_name = failure_file_path.name
        session.add(version)
        session.commit()
        session.refresh(version)

        with pytest.raises(document_parser.DocumentParseError, match="quality policy failed"):
            document_parser.parse_document_version(session, version)

        refreshed_version = session.get(DocumentVersion, version_id)
        parse_runs = list(
            session.scalars(
                select(DocumentParseRun)
                .where(DocumentParseRun.document_version_id == version_id)
                .order_by(DocumentParseRun.id)
            )
        )

    assert refreshed_version is not None
    assert refreshed_version.parse_status == "failed"
    assert refreshed_version.active_parse_run_id == first_active_parse_run_id
    assert [parse_run.status for parse_run in parse_runs] == ["parsed", "failed"]
    failed_summary = json.loads(parse_runs[-1].summary_json or "{}")
    assert failed_summary["coverage"]["policy_result"] == "fail"
    assert failed_summary["diagnostics"][0]["code"] == "pdf_ocr_quality_failed"


def test_parse_document_version_marks_failed_when_no_valid_body_blocks(
    session_factory,
    tmp_path: Path,
):
    file_path = tmp_path / "blank.docx"
    DocxDocument().save(file_path)
    version_id = _create_document_version(session_factory, file_path)

    with session_factory() as session:
        version = session.get(DocumentVersion, version_id)
        assert version is not None

        with pytest.raises(document_parser.DocumentParseError, match="no valid body blocks"):
            document_parser.parse_document_version(session, version)

        refreshed_version = session.get(DocumentVersion, version_id)
        parse_runs = list(
            session.scalars(
                select(DocumentParseRun)
                .where(DocumentParseRun.document_version_id == version_id)
                .order_by(DocumentParseRun.id)
            )
        )
        blocks = list(
            session.scalars(
                select(DocumentBlock).where(DocumentBlock.document_version_id == version_id)
            )
        )

    assert refreshed_version is not None
    assert refreshed_version.parse_status == "failed"
    assert refreshed_version.active_parse_run_id is None
    assert [parse_run.status for parse_run in parse_runs] == ["failed"]
    assert refreshed_version.parsed_snapshot is None
    assert blocks == []


def test_failed_reparse_keeps_previous_active_parse_run(session_factory, tmp_path: Path):
    success_file_path = tmp_path / "success.docx"
    _save_docx(
        success_file_path,
        [
            ("Overview", "Heading 1"),
            ("The parser should create an active parse run.", None),
        ],
    )
    failure_file_path = tmp_path / "failure.docx"
    DocxDocument().save(failure_file_path)
    version_id = _create_document_version(session_factory, success_file_path)

    with session_factory() as session:
        version = session.get(DocumentVersion, version_id)
        assert version is not None

        parsed_version = document_parser.parse_document_version(session, version)
        first_active_parse_run_id = parsed_version.active_parse_run_id
        assert first_active_parse_run_id is not None

        version.file_path = str(failure_file_path)
        session.add(version)
        session.commit()
        session.refresh(version)

        with pytest.raises(document_parser.DocumentParseError, match="no valid body blocks"):
            document_parser.parse_document_version(session, version)

        refreshed_version = session.get(DocumentVersion, version_id)
        parse_runs = list(
            session.scalars(
                select(DocumentParseRun)
                .where(DocumentParseRun.document_version_id == version_id)
                .order_by(DocumentParseRun.id)
            )
        )

    assert refreshed_version is not None
    assert refreshed_version.parse_status == "failed"
    assert refreshed_version.active_parse_run_id == first_active_parse_run_id
    assert [parse_run.status for parse_run in parse_runs] == ["parsed", "failed"]


def test_successful_reparse_creates_new_active_parse_run(session_factory, tmp_path: Path):
    file_path = tmp_path / "reparse.docx"
    _save_docx(
        file_path,
        [
            ("Overview", "Heading 1"),
            ("The first parser truth should remain attached to its parse run.", None),
        ],
    )
    version_id = _create_document_version(session_factory, file_path)

    with session_factory() as session:
        version = session.get(DocumentVersion, version_id)
        assert version is not None

        first_version = document_parser.parse_document_version(session, version)
        first_active_parse_run_id = first_version.active_parse_run_id
        assert first_active_parse_run_id is not None

        _save_docx(
            file_path,
            [
                ("Overview", "Heading 1"),
                ("The second parser truth should become the active parse run.", None),
            ],
        )

        second_version = document_parser.parse_document_version(session, version)
        second_active_parse_run_id = second_version.active_parse_run_id
        assert second_active_parse_run_id is not None

        parse_runs = list(
            session.scalars(
                select(DocumentParseRun)
                .where(DocumentParseRun.document_version_id == version_id)
                .order_by(DocumentParseRun.id)
            )
        )
        first_blocks = list(
            session.scalars(
                select(DocumentBlock)
                .where(DocumentBlock.parse_run_id == first_active_parse_run_id)
                .order_by(DocumentBlock.order_index)
            )
        )
        second_blocks = list(
            session.scalars(
                select(DocumentBlock)
                .where(DocumentBlock.parse_run_id == second_active_parse_run_id)
                .order_by(DocumentBlock.order_index)
            )
        )

    assert second_active_parse_run_id != first_active_parse_run_id
    assert [parse_run.status for parse_run in parse_runs] == ["parsed", "parsed"]
    assert [block.normalized_content for block in first_blocks] == [
        "Overview",
        "The first parser truth should remain attached to its parse run.",
    ]
    assert [block.normalized_content for block in second_blocks] == [
        "Overview",
        "The second parser truth should become the active parse run.",
    ]


def test_parse_document_version_records_quality_diagnostics_as_warnings(
    session_factory,
    tmp_path: Path,
):
    file_path = tmp_path / "quality-warning.docx"
    _save_docx(
        file_path,
        [
            ("Services Agreement", "Heading 1"),
            ("The parties agree to payment and confidentiality obligations.", None),
        ],
    )
    _inject_textbox(file_path, "Text box side letter excluded from canonical parser truth.")
    version_id = _create_document_version(session_factory, file_path)

    with session_factory() as session:
        version = session.get(DocumentVersion, version_id)
        assert version is not None

        parsed_version = document_parser.parse_document_version(session, version)
        parse_run = session.get(DocumentParseRun, parsed_version.active_parse_run_id)
        blocks = list(
            session.scalars(
                select(DocumentBlock)
                .where(DocumentBlock.parse_run_id == parsed_version.active_parse_run_id)
                .order_by(DocumentBlock.order_index)
            )
        )

    assert parse_run is not None
    assert parsed_version.parse_status == "parsed_with_warnings"
    assert parse_run.status == "parsed_with_warnings"
    assert parse_run.warning_count >= 1
    assert [block.block_type for block in blocks] == ["heading", "paragraph", "paragraph"]
    assert "Text box side letter" in blocks[2].normalized_content

    summary = json.loads(parse_run.summary_json or "{}")
    assert summary["coverage"]["policy_result"] == "warn"
    assert summary["diagnostics"][0]["code"] == "unsupported_textbox"
    assert summary["diagnostics"][0]["source_part"] == "word/document.xml"
    assert summary["diagnostics"][0]["occurrence_key"]


def test_quality_policy_failure_does_not_replace_active_parse_run(
    session_factory,
    tmp_path: Path,
):
    success_file_path = tmp_path / "quality-success.docx"
    _save_docx(
        success_file_path,
        [
            ("Master Services Agreement", "Heading 1"),
            ("The canonical body text is complete.", None),
        ],
    )
    failure_file_path = tmp_path / "quality-failure.docx"
    _save_docx(
        failure_file_path,
        [
            ("Master Services Agreement", "Heading 1"),
            ("Visible canonical body text.", None),
        ],
    )
    _inject_content_control(
        failure_file_path,
        (
            "Hidden payment obligation confidentiality covenant indemnity duty "
            "audit right termination right renewal condition service credit "
            "data security insurance notice governing law venue limitation "
            "liability survival assignment subcontractor compliance warranty."
        ),
    )
    version_id = _create_document_version(session_factory, success_file_path)

    with session_factory() as session:
        version = session.get(DocumentVersion, version_id)
        assert version is not None

        parsed_version = document_parser.parse_document_version(session, version)
        first_active_parse_run_id = parsed_version.active_parse_run_id
        assert first_active_parse_run_id is not None

        version.file_path = str(failure_file_path)
        session.add(version)
        session.commit()
        session.refresh(version)

        with pytest.raises(document_parser.DocumentParseError, match="coverage is too low"):
            document_parser.parse_document_version(session, version)

        refreshed_version = session.get(DocumentVersion, version_id)
        parse_runs = list(
            session.scalars(
                select(DocumentParseRun)
                .where(DocumentParseRun.document_version_id == version_id)
                .order_by(DocumentParseRun.id)
            )
        )
        failed_parse_run = parse_runs[-1]
        failed_blocks = list(
            session.scalars(
                select(DocumentBlock).where(DocumentBlock.parse_run_id == failed_parse_run.id)
            )
        )

    assert refreshed_version is not None
    assert refreshed_version.parse_status == "failed"
    assert refreshed_version.active_parse_run_id == first_active_parse_run_id
    assert [parse_run.status for parse_run in parse_runs] == ["parsed", "failed"]
    assert failed_parse_run.warning_count >= 1
    assert json.loads(failed_parse_run.summary_json or "{}")["coverage"]["policy_result"] == "fail"
    assert failed_blocks == []


def test_legal_contract_numbering_classification(session_factory, tmp_path: Path):
    file_path = tmp_path / "legal-numbering.docx"
    _save_docx(
        file_path,
        [
            ("1. Definitions", None),
            ("Agreement means this contract.", None),
            ("1.1 Payment Terms", None),
            ("Invoices are due within 30 days.", None),
            ("Article 5 Confidentiality", None),
            ("(a) The recipient must protect confidential information.", None),
            ("(i) The recipient must limit internal access.", None),
        ],
    )
    version_id = _create_document_version(session_factory, file_path)

    with session_factory() as session:
        version = session.get(DocumentVersion, version_id)
        assert version is not None

        parsed_version = document_parser.parse_document_version(session, version)
        blocks = list(
            session.scalars(
                select(DocumentBlock)
                .where(DocumentBlock.parse_run_id == parsed_version.active_parse_run_id)
                .order_by(DocumentBlock.order_index)
            )
        )

    assert [(block.normalized_content, block.block_type, block.heading_level) for block in blocks] == [
        ("1. Definitions", "heading", 1),
        ("Agreement means this contract.", "paragraph", None),
        ("1.1 Payment Terms", "heading", 2),
        ("Invoices are due within 30 days.", "paragraph", None),
        ("Article 5 Confidentiality", "heading", 1),
        ("(a) The recipient must protect confidential information.", "list_item", None),
        ("(i) The recipient must limit internal access.", "list_item", None),
    ]
