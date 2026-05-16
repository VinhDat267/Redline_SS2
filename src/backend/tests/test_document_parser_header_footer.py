import hashlib
import json
from pathlib import Path

from docx import Document as DocxDocument
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches
from sqlalchemy import select

from app.models import Document, DocumentBlock, DocumentParseRun, DocumentSurface, DocumentTable, DocumentVersion, Project
from app.services import document_parser


def _save_docx_with_header_footer(
    path: Path,
    *,
    body_paragraphs: list[tuple[str, str | None]],
    header_paragraphs: list[str] | None = None,
    header_table_rows: list[list[str]] | None = None,
    footer_paragraphs: list[str] | None = None,
    header_fields: list[tuple[str, str]] | None = None,
    header_complex_fields: list[tuple[str, str]] | None = None,
) -> None:
    document = DocxDocument()
    for text, style in body_paragraphs:
        paragraph = document.add_paragraph(text)
        if style is not None:
            paragraph.style = style

    section = document.sections[0]

    if header_paragraphs:
        for text in header_paragraphs:
            section.header.add_paragraph(text)

    if header_fields:
        for prefix_text, field_name in header_fields:
            paragraph = section.header.add_paragraph(prefix_text)
            _append_simple_field(paragraph, field_name)

    if header_complex_fields:
        for prefix_text, field_name in header_complex_fields:
            paragraph = section.header.add_paragraph(prefix_text)
            _append_complex_field(paragraph, field_name)

    if header_table_rows:
        table = section.header.add_table(rows=len(header_table_rows), cols=len(header_table_rows[0]), width=Inches(6))
        for row_index, row_values in enumerate(header_table_rows):
            for column_index, value in enumerate(row_values):
                table.cell(row_index, column_index).text = value

    if footer_paragraphs:
        for text in footer_paragraphs:
            section.footer.add_paragraph(text)

    document.save(path)


def _append_simple_field(paragraph, field_name: str) -> None:
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), field_name)
    run = OxmlElement("w:r")
    text = OxmlElement("w:t")
    text.text = "1"
    run.append(text)
    field.append(run)
    paragraph._p.append(field)


def _append_complex_field(paragraph, field_name: str) -> None:
    begin_run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    begin_run._r.append(begin)

    instruction_run = paragraph.add_run()
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = f" {field_name} "
    instruction_run._r.append(instruction)

    separate_run = paragraph.add_run()
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    separate_run._r.append(separate)

    result_run = paragraph.add_run("99")

    end_run = paragraph.add_run()
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    end_run._r.append(end)


def _create_document_version(session_factory, file_path: Path) -> int:
    with session_factory() as session:
        project = Project(name="Parser Header Footer Project", description="Parser header/footer tests")
        document = Document(
            project=project,
            title="Parser Header Footer Document",
            document_type="SPEC",
            description="Parser header/footer document",
        )
        version = DocumentVersion(
            document=document,
            version_label="v1.0",
            file_name=file_path.name,
            file_path=str(file_path),
            parse_status="pending",
        )
        session.add_all([project, document, version])
        session.commit()
        session.refresh(version)
        return version.id


def test_parse_document_version_persists_header_and_footer_surfaces(session_factory, tmp_path: Path):
    file_path = tmp_path / "header-footer-surfaces.docx"
    _save_docx_with_header_footer(
        file_path,
        body_paragraphs=[
            ("Scope", "Heading 1"),
            ("Body paragraph.", None),
        ],
        header_paragraphs=["Release Notes"],
        header_table_rows=[
            ["Key", "Value"],
            ["Revision", "1.0"],
        ],
        footer_paragraphs=["Confidential"],
    )
    version_id = _create_document_version(session_factory, file_path)

    with session_factory() as session:
        version = session.get(DocumentVersion, version_id)
        assert version is not None

        parsed_version = document_parser.parse_document_version(session, version)
        active_parse_run_id = parsed_version.active_parse_run_id
        surfaces = list(
            session.scalars(
                select(DocumentSurface)
                .where(DocumentSurface.parse_run_id == active_parse_run_id)
                .order_by(DocumentSurface.logical_order_index)
            )
        )
        blocks = list(
            session.scalars(
                select(DocumentBlock)
                .where(DocumentBlock.parse_run_id == active_parse_run_id)
                .order_by(DocumentBlock.order_index)
            )
        )
        tables = list(
            session.scalars(
                select(DocumentTable)
                .where(DocumentTable.parse_run_id == active_parse_run_id)
                .order_by(DocumentTable.table_order_index)
            )
        )

    assert parsed_version.parse_status == "parsed"
    assert [surface.surface_type for surface in surfaces] == ["body", "header", "footer"]
    assert [surface.surface_key for surface in surfaces] == [
        "body-main",
        "header-section-1-default",
        "footer-section-1-default",
    ]
    assert [surface.logical_order_index for surface in surfaces] == [0, 1, 2]
    assert [surface.section_ref for surface in surfaces] == [None, "section-1", "section-1"]

    header_surface = surfaces[1]
    footer_surface = surfaces[2]
    header_blocks = [block for block in blocks if block.surface_id == header_surface.id]
    footer_blocks = [block for block in blocks if block.surface_id == footer_surface.id]

    assert [block.block_type for block in blocks] == [
        "heading",
        "paragraph",
        "paragraph",
        "table_row",
        "table_row",
        "paragraph",
    ]
    assert [block.surface_order_index for block in header_blocks] == [0, 1, 2]
    assert [block.raw_content for block in header_blocks] == [
        "Release Notes",
        "Key: Key || Value: Value",
        "Key: Revision || Value: 1.0",
    ]
    assert [block.raw_content for block in footer_blocks] == ["Confidential"]

    assert len(tables) == 1
    assert tables[0].surface_id == header_surface.id
    assert tables[0].table_order_index == 0

    parsed_snapshot = json.loads(parsed_version.parsed_snapshot or "{}")
    assert parsed_snapshot["total_surfaces"] == 3
    assert parsed_snapshot["counts_by_surface_type"] == {
        "body": 1,
        "header": 1,
        "footer": 1,
    }
    assert parsed_snapshot["counts_by_block_type"] == {
        "heading": 1,
        "paragraph": 3,
        "list_item": 0,
        "table_row": 2,
    }
    assert parsed_snapshot["table_count"] == 1
    assert parsed_snapshot["row_count"] == 2


def test_parse_document_version_normalizes_supported_header_auto_fields(session_factory, tmp_path: Path):
    file_path = tmp_path / "header-field-normalization.docx"
    _save_docx_with_header_footer(
        file_path,
        body_paragraphs=[
            ("Overview", "Heading 1"),
            ("Body paragraph.", None),
        ],
        header_fields=[("Page ", "PAGE")],
    )
    version_id = _create_document_version(session_factory, file_path)

    with session_factory() as session:
        version = session.get(DocumentVersion, version_id)
        assert version is not None

        parsed_version = document_parser.parse_document_version(session, version)
        active_parse_run_id = parsed_version.active_parse_run_id
        parse_run = session.scalar(
            select(DocumentParseRun).where(DocumentParseRun.id == active_parse_run_id)
        )
        header_surface = session.scalar(
            select(DocumentSurface)
            .where(DocumentSurface.parse_run_id == active_parse_run_id)
            .where(DocumentSurface.surface_type == "header")
        )
        header_blocks = list(
            session.scalars(
                select(DocumentBlock)
                .where(DocumentBlock.surface_id == header_surface.id)
                .order_by(DocumentBlock.surface_order_index)
            )
        )

    assert parse_run is not None
    assert parsed_version.parse_status == "parsed_with_warnings"
    assert parse_run.status == "parsed_with_warnings"
    assert parse_run.warning_count == 1
    assert header_blocks[0].raw_content == "Page __FIELD_PAGE__"
    assert header_blocks[0].normalized_content == "Page __FIELD_PAGE__"
    assert header_blocks[0].block_key == (
        f"blk-0002-paragraph-{hashlib.sha1('Page __FIELD_PAGE__'.encode('utf-8')).hexdigest()[:10]}"
    )

    summary = json.loads(parse_run.summary_json or "{}")
    assert summary["warning_count"] == 1
    assert any("PAGE" in warning for warning in summary["warnings"])

    parsed_snapshot = json.loads(parsed_version.parsed_snapshot or "{}")
    assert parsed_snapshot["counts_by_surface_type"] == {
        "body": 1,
        "header": 1,
    }
    assert parsed_snapshot["warning_count"] == 1
    assert any("PAGE" in warning for warning in parsed_snapshot["warnings"])


def test_parse_document_version_normalizes_supported_complex_header_auto_fields(session_factory, tmp_path: Path):
    file_path = tmp_path / "header-complex-field-normalization.docx"
    _save_docx_with_header_footer(
        file_path,
        body_paragraphs=[
            ("Overview", "Heading 1"),
            ("Body paragraph.", None),
        ],
        header_complex_fields=[("Page ", "PAGE")],
    )
    version_id = _create_document_version(session_factory, file_path)

    with session_factory() as session:
        version = session.get(DocumentVersion, version_id)
        assert version is not None

        parsed_version = document_parser.parse_document_version(session, version)
        active_parse_run_id = parsed_version.active_parse_run_id
        parse_run = session.scalar(
            select(DocumentParseRun).where(DocumentParseRun.id == active_parse_run_id)
        )
        header_surface = session.scalar(
            select(DocumentSurface)
            .where(DocumentSurface.parse_run_id == active_parse_run_id)
            .where(DocumentSurface.surface_type == "header")
        )
        header_blocks = list(
            session.scalars(
                select(DocumentBlock)
                .where(DocumentBlock.surface_id == header_surface.id)
                .order_by(DocumentBlock.surface_order_index)
            )
        )

    assert parse_run is not None
    assert parsed_version.parse_status == "parsed_with_warnings"
    assert parse_run.status == "parsed_with_warnings"
    assert parse_run.warning_count == 1
    assert header_blocks[0].raw_content == "Page __FIELD_PAGE__"
    assert header_blocks[0].normalized_content == "Page __FIELD_PAGE__"
    assert header_blocks[0].block_key == (
        f"blk-0002-paragraph-{hashlib.sha1('Page __FIELD_PAGE__'.encode('utf-8')).hexdigest()[:10]}"
    )

    summary = json.loads(parse_run.summary_json or "{}")
    assert summary["warning_count"] == 1
    assert any("PAGE" in warning for warning in summary["warnings"])
