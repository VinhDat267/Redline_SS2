import json
from copy import deepcopy
from pathlib import Path
from zipfile import ZipFile

from docx import Document as DocxDocument
from docx.oxml.ns import qn
from lxml import etree
from sqlalchemy import select

from app.models import Document, DocumentBlock, DocumentSurface, DocumentTable, DocumentVersion, Project
from app.services import document_parser


_REL_NS = "http://schemas.openxmlformats.org/package/2006/relationships"
_CONTENT_TYPE_NS = "http://schemas.openxmlformats.org/package/2006/content-types"
_WORD_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
_OFFICE_REL_NS = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
_NSMAP = {"w": _WORD_NS, "rel": _REL_NS, "ct": _CONTENT_TYPE_NS}


def _save_docx_with_notes(
    path: Path,
    *,
    body_paragraphs: list[tuple[str, str | None]],
    footnotes: list[dict[str, object]] | None = None,
    endnotes: list[dict[str, object]] | None = None,
) -> None:
    document = DocxDocument()
    for text, style in body_paragraphs:
        paragraph = document.add_paragraph(text)
        if style is not None:
            paragraph.style = style

    document.save(path)
    _inject_note_parts(path, footnotes=footnotes or [], endnotes=endnotes or [])


def _inject_note_parts(
    path: Path,
    *,
    footnotes: list[dict[str, object]],
    endnotes: list[dict[str, object]],
) -> None:
    if not footnotes and not endnotes:
        return

    with ZipFile(path, "r") as archive:
        entries = {info.filename: archive.read(info.filename) for info in archive.infolist()}

    document_xml = etree.fromstring(entries["word/document.xml"])
    body_paragraphs = document_xml.xpath("./w:body/w:p", namespaces=_NSMAP)

    if footnotes:
        for paragraph, note in zip(body_paragraphs, footnotes, strict=False):
            _append_note_reference(paragraph, "footnoteReference", int(note["id"]))
        entries["word/footnotes.xml"] = _build_note_part_xml("footnotes", "footnote", footnotes)

    if endnotes:
        for paragraph, note in zip(body_paragraphs[len(footnotes) :], endnotes, strict=False):
            _append_note_reference(paragraph, "endnoteReference", int(note["id"]))
        entries["word/endnotes.xml"] = _build_note_part_xml("endnotes", "endnote", endnotes)

    entries["word/document.xml"] = etree.tostring(
        document_xml,
        xml_declaration=True,
        encoding="UTF-8",
        standalone=True,
    )
    entries["word/_rels/document.xml.rels"] = _update_document_relationships(
        entries["word/_rels/document.xml.rels"],
        include_footnotes=bool(footnotes),
        include_endnotes=bool(endnotes),
    )
    entries["[Content_Types].xml"] = _update_content_types(
        entries["[Content_Types].xml"],
        include_footnotes=bool(footnotes),
        include_endnotes=bool(endnotes),
    )

    with ZipFile(path, "w") as archive:
        for filename, payload in entries.items():
            archive.writestr(filename, payload)


def _append_note_reference(paragraph_element, reference_tag: str, note_id: int) -> None:
    run = etree.SubElement(paragraph_element, qn("w:r"))
    reference = etree.SubElement(run, qn(f"w:{reference_tag}"))
    reference.set(qn("w:id"), str(note_id))


def _build_note_part_xml(root_tag: str, item_tag: str, notes: list[dict[str, object]]) -> bytes:
    root = etree.Element(qn(f"w:{root_tag}"), nsmap={"w": _WORD_NS})
    for note in notes:
        note_element = etree.SubElement(root, qn(f"w:{item_tag}"))
        note_element.set(qn("w:id"), str(int(note["id"])))
        for item in note["items"]:
            note_element.append(_build_item_element(item))

    return etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone=True)


def _build_item_element(item: dict[str, object]):
    item_type = item["type"]
    if item_type == "paragraph":
        helper_document = DocxDocument()
        paragraph = helper_document.add_paragraph(str(item["text"]))
        style = item.get("style")
        if style is not None:
            paragraph.style = str(style)
        return deepcopy(paragraph._p)

    if item_type == "table":
        rows = item["rows"]
        assert isinstance(rows, list)
        helper_document = DocxDocument()
        table = helper_document.add_table(rows=len(rows), cols=len(rows[0]))
        for row_index, row_values in enumerate(rows):
            for column_index, value in enumerate(row_values):
                table.cell(row_index, column_index).text = str(value)
        return deepcopy(table._tbl)

    raise ValueError(f"Unsupported note item type: {item_type}")


def _update_document_relationships(
    content: bytes,
    *,
    include_footnotes: bool,
    include_endnotes: bool,
) -> bytes:
    root = etree.fromstring(content)
    existing_ids = [
        int(value[3:])
        for value in root.xpath("./rel:Relationship/@Id", namespaces=_NSMAP)
        if value.startswith("rId") and value[3:].isdigit()
    ]
    next_id = max(existing_ids, default=0) + 1

    def add_relationship(rel_type: str, target: str) -> None:
        nonlocal next_id
        if root.xpath(f"./rel:Relationship[@Type='{rel_type}']", namespaces=_NSMAP):
            return
        relationship = etree.SubElement(root, f"{{{_REL_NS}}}Relationship")
        relationship.set("Id", f"rId{next_id}")
        relationship.set("Type", rel_type)
        relationship.set("Target", target)
        next_id += 1

    if include_footnotes:
        add_relationship(f"{_OFFICE_REL_NS}/footnotes", "footnotes.xml")
    if include_endnotes:
        add_relationship(f"{_OFFICE_REL_NS}/endnotes", "endnotes.xml")

    return etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone=True)


def _update_content_types(
    content: bytes,
    *,
    include_footnotes: bool,
    include_endnotes: bool,
) -> bytes:
    root = etree.fromstring(content)

    def add_override(part_name: str, content_type: str) -> None:
        if root.xpath(
            f"./ct:Override[@PartName='{part_name}']",
            namespaces=_NSMAP,
        ):
            return
        override = etree.SubElement(root, f"{{{_CONTENT_TYPE_NS}}}Override")
        override.set("PartName", part_name)
        override.set("ContentType", content_type)

    if include_footnotes:
        add_override(
            "/word/footnotes.xml",
            "application/vnd.openxmlformats-officedocument.wordprocessingml.footnotes+xml",
        )
    if include_endnotes:
        add_override(
            "/word/endnotes.xml",
            "application/vnd.openxmlformats-officedocument.wordprocessingml.endnotes+xml",
        )

    return etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone=True)


def _create_document_version(session_factory, file_path: Path) -> int:
    with session_factory() as session:
        project = Project(name="Parser Notes Project", description="Parser note tests")
        document = Document(
            project=project,
            title="Parser Notes Document",
            document_type="SPEC",
            description="Parser notes document",
        )
        version = DocumentVersion(
            document=document,
            version_label="v1.0",
            file_name=file_path.name,
            file_path=str(file_path),
            parse_status="pending",
        )
        session.add_all([project, document, version])
        session.commit()
        session.refresh(version)
        return version.id


def test_parse_document_version_persists_footnote_and_endnote_surfaces(session_factory, tmp_path: Path):
    file_path = tmp_path / "notes-surfaces.docx"
    _save_docx_with_notes(
        file_path,
        body_paragraphs=[
            ("Overview", "Heading 1"),
            ("Body paragraph with footnote.", None),
            ("Body paragraph with endnote.", None),
        ],
        footnotes=[
            {
                "id": 2,
                "items": [
                    {"type": "paragraph", "text": "Footnote details"},
                ],
            }
        ],
        endnotes=[
            {
                "id": 3,
                "items": [
                    {"type": "paragraph", "text": "Endnote details"},
                ],
            }
        ],
    )
    version_id = _create_document_version(session_factory, file_path)

    with session_factory() as session:
        version = session.get(DocumentVersion, version_id)
        assert version is not None

        parsed_version = document_parser.parse_document_version(session, version)
        active_parse_run_id = parsed_version.active_parse_run_id
        surfaces = list(
            session.scalars(
                select(DocumentSurface)
                .where(DocumentSurface.parse_run_id == active_parse_run_id)
                .order_by(DocumentSurface.logical_order_index)
            )
        )
        blocks = list(
            session.scalars(
                select(DocumentBlock)
                .where(DocumentBlock.parse_run_id == active_parse_run_id)
                .order_by(DocumentBlock.order_index)
            )
        )

    assert parsed_version.parse_status == "parsed"
    assert [surface.surface_type for surface in surfaces] == ["body", "footnote", "endnote"]
    assert [surface.surface_key for surface in surfaces] == ["body-main", "footnote-2", "endnote-3"]
    assert [surface.logical_order_index for surface in surfaces] == [0, 1, 2]
    assert [block.raw_content for block in blocks] == [
        "Overview",
        "Body paragraph with footnote.",
        "Body paragraph with endnote.",
        "Footnote details",
        "Endnote details",
    ]

    parsed_snapshot = json.loads(parsed_version.parsed_snapshot or "{}")
    assert parsed_snapshot["counts_by_surface_type"] == {
        "body": 1,
        "footnote": 1,
        "endnote": 1,
    }
    assert parsed_snapshot["counts_by_block_type"] == {
        "heading": 1,
        "paragraph": 4,
        "list_item": 0,
        "table_row": 0,
    }


def test_parse_document_version_ignores_orphan_notes_without_body_references(session_factory, tmp_path: Path):
    file_path = tmp_path / "notes-ignore-orphans.docx"
    _save_docx_with_notes(
        file_path,
        body_paragraphs=[
            ("Overview with referenced footnote.", "Heading 1"),
        ],
        footnotes=[
            {
                "id": 2,
                "items": [
                    {"type": "paragraph", "text": "Referenced footnote details"},
                ],
            },
            {
                "id": 3,
                "items": [
                    {"type": "paragraph", "text": "Orphan footnote details"},
                ],
            },
        ],
    )
    version_id = _create_document_version(session_factory, file_path)

    with session_factory() as session:
        version = session.get(DocumentVersion, version_id)
        assert version is not None

        parsed_version = document_parser.parse_document_version(session, version)
        active_parse_run_id = parsed_version.active_parse_run_id
        surfaces = list(
            session.scalars(
                select(DocumentSurface)
                .where(DocumentSurface.parse_run_id == active_parse_run_id)
                .order_by(DocumentSurface.logical_order_index)
            )
        )
        blocks = list(
            session.scalars(
                select(DocumentBlock)
                .where(DocumentBlock.parse_run_id == active_parse_run_id)
                .order_by(DocumentBlock.order_index)
            )
        )

    assert parsed_version.parse_status == "parsed"
    assert [surface.surface_key for surface in surfaces] == ["body-main", "footnote-2"]
    assert [block.raw_content for block in blocks] == [
        "Overview with referenced footnote.",
        "Referenced footnote details",
    ]


def test_parse_document_version_persists_table_truth_inside_footnote_surface(session_factory, tmp_path: Path):
    file_path = tmp_path / "notes-table.docx"
    _save_docx_with_notes(
        file_path,
        body_paragraphs=[
            ("Appendix", "Heading 1"),
            ("Body paragraph.", None),
        ],
        footnotes=[
            {
                "id": 2,
                "items": [
                    {"type": "paragraph", "text": "Footnote matrix"},
                    {
                        "type": "table",
                        "rows": [
                            ["Field", "Value"],
                            ["Status", "Draft"],
                        ],
                    },
                ],
            }
        ],
    )
    version_id = _create_document_version(session_factory, file_path)

    with session_factory() as session:
        version = session.get(DocumentVersion, version_id)
        assert version is not None

        parsed_version = document_parser.parse_document_version(session, version)
        active_parse_run_id = parsed_version.active_parse_run_id
        footnote_surface = session.scalar(
            select(DocumentSurface)
            .where(DocumentSurface.parse_run_id == active_parse_run_id)
            .where(DocumentSurface.surface_type == "footnote")
        )
        assert footnote_surface is not None
        tables = list(
            session.scalars(
                select(DocumentTable)
                .where(DocumentTable.parse_run_id == active_parse_run_id)
                .where(DocumentTable.surface_id == footnote_surface.id)
                .order_by(DocumentTable.table_order_index)
            )
        )
        footnote_blocks = list(
            session.scalars(
                select(DocumentBlock)
                .where(DocumentBlock.surface_id == footnote_surface.id)
                .order_by(DocumentBlock.surface_order_index)
            )
        )

    assert parsed_version.parse_status == "parsed"
    assert len(tables) == 1
    assert tables[0].table_key == "tbl-0000"
    assert [block.raw_content for block in footnote_blocks] == [
        "Footnote matrix",
        "Field: Field || Value: Value",
        "Field: Status || Value: Draft",
    ]

    parsed_snapshot = json.loads(parsed_version.parsed_snapshot or "{}")
    assert parsed_snapshot["counts_by_surface_type"] == {
        "body": 1,
        "footnote": 1,
    }
    assert parsed_snapshot["table_count"] == 1
    assert parsed_snapshot["row_count"] == 2
