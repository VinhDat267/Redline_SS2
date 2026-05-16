import json
from pathlib import Path

from docx import Document as DocxDocument
from sqlalchemy import select

from app.models import (
    Document,
    DocumentBlock,
    DocumentParseRun,
    DocumentTable,
    DocumentTableCell,
    DocumentTableColumn,
    DocumentTableRow,
    DocumentVersion,
    Project,
)
from app.services import document_parser


def _save_docx_with_items(path: Path, items: list[dict[str, object]]) -> None:
    document = DocxDocument()
    for item in items:
        if item["type"] == "paragraph":
            paragraph = document.add_paragraph(str(item["text"]))
            style = item.get("style")
            if style is not None:
                paragraph.style = str(style)
        elif item["type"] == "table":
            rows = item["rows"]
            assert isinstance(rows, list)
            table = document.add_table(rows=len(rows), cols=len(rows[0]))
            for row_index, row_values in enumerate(rows):
                for column_index, cell_value in enumerate(row_values):
                    table.cell(row_index, column_index).text = str(cell_value)
        else:  # pragma: no cover - test helper guard
            raise ValueError(f"Unsupported item type: {item['type']}")

    document.save(path)


def _save_docx_with_merged_table(path: Path) -> None:
    document = DocxDocument()
    document.add_paragraph("Matrices").style = "Heading 1"
    table = document.add_table(rows=3, cols=3)
    values = [
        ["Requirement ID", "Title", "Notes"],
        ["REQ-001", "Login", "Shared note"],
        ["REQ-002", "Register", ""],
    ]
    for row_index, row_values in enumerate(values):
        for column_index, cell_value in enumerate(row_values):
            table.cell(row_index, column_index).text = cell_value

    table.cell(2, 0).merge(table.cell(2, 1)).text = "REQ-002 / Register"
    table.cell(1, 2).merge(table.cell(2, 2)).text = "Shared note"
    document.save(path)


def _create_document_version(session_factory, file_path: Path) -> int:
    with session_factory() as session:
        project = Project(name="Parser Table Project", description="Parser table tests")
        document = Document(
            project=project,
            title="Parser Table Document",
            document_type="SPEC",
            description="Parser table document",
        )
        version = DocumentVersion(
            document=document,
            version_label="v1.0",
            file_name=file_path.name,
            file_path=str(file_path),
            parse_status="pending",
        )
        session.add_all([project, document, version])
        session.commit()
        session.refresh(version)
        return version.id


def test_parse_document_version_persists_structured_table_truth(session_factory, tmp_path: Path):
    file_path = tmp_path / "table-structured.docx"
    _save_docx_with_items(
        file_path,
        [
            {"type": "paragraph", "text": "Requirements", "style": "Heading 1"},
            {"type": "paragraph", "text": "The table below defines the requirement list."},
            {
                "type": "table",
                "rows": [
                    ["Requirement ID", "Title"],
                    ["REQ-001", "Login"],
                    ["REQ-002", "Register"],
                ],
            },
            {"type": "paragraph", "text": "End of table section."},
        ],
    )
    version_id = _create_document_version(session_factory, file_path)

    with session_factory() as session:
        version = session.get(DocumentVersion, version_id)
        assert version is not None

        parsed_version = document_parser.parse_document_version(session, version)
        active_parse_run_id = parsed_version.active_parse_run_id
        tables = list(
            session.scalars(
                select(DocumentTable)
                .where(DocumentTable.parse_run_id == active_parse_run_id)
                .order_by(DocumentTable.table_order_index)
            )
        )
        columns = list(
            session.scalars(
                select(DocumentTableColumn)
                .where(DocumentTableColumn.table_id == tables[0].id)
                .order_by(DocumentTableColumn.column_index)
            )
        )
        rows = list(
            session.scalars(
                select(DocumentTableRow)
                .where(DocumentTableRow.table_id == tables[0].id)
                .order_by(DocumentTableRow.row_index)
            )
        )
        cells = list(
            session.scalars(
                select(DocumentTableCell)
                .join(DocumentTableRow, DocumentTableCell.row_id == DocumentTableRow.id)
                .where(DocumentTableRow.table_id == tables[0].id)
                .order_by(DocumentTableRow.row_index, DocumentTableCell.column_index)
            )
        )
        blocks = list(
            session.scalars(
                select(DocumentBlock)
                .where(DocumentBlock.parse_run_id == active_parse_run_id)
                .order_by(DocumentBlock.order_index)
            )
        )

    assert len(tables) == 1
    assert tables[0].header_strategy == "explicit_first_row"
    assert tables[0].section_title == "Requirements"
    assert tables[0].table_order_index == 0

    assert [column.column_key for column in columns] == ["requirement_id", "title"]
    assert [column.header_text for column in columns] == ["Requirement ID", "Title"]
    assert [column.normalized_header_text for column in columns] == ["Requirement ID", "Title"]
    assert [column.source_kind for column in columns] == ["explicit", "explicit"]

    assert len(rows) == 3
    assert [row.row_index for row in rows] == [0, 1, 2]
    assert [row.is_header_row for row in rows] == [True, False, False]
    assert all(row.document_block_id is not None for row in rows)

    assert [block.block_type for block in blocks] == [
        "heading",
        "paragraph",
        "table_row",
        "table_row",
        "table_row",
        "paragraph",
    ]
    assert blocks[2].raw_content == "Requirement ID: Requirement ID || Title: Title"
    assert blocks[3].raw_content == "Requirement ID: REQ-001 || Title: Login"
    assert json.loads(blocks[3].normalized_content) == [
        {
            "column_key": "requirement_id",
            "normalized_header_text": "Requirement ID",
            "normalized_value": "REQ-001",
            "merge_token": None,
        },
        {
            "column_key": "title",
            "normalized_header_text": "Title",
            "normalized_value": "Login",
            "merge_token": None,
        },
    ]
    assert json.loads(rows[1].structured_row_json)["cells"][0]["column_key"] == "requirement_id"
    assert len(cells) == 6
    assert all(cell.row_span == 1 for cell in cells)
    assert all(cell.col_span == 1 for cell in cells)

    parsed_snapshot = json.loads(parsed_version.parsed_snapshot or "{}")
    assert parsed_snapshot["table_count"] == 1
    assert parsed_snapshot["row_count"] == 3
    assert parsed_snapshot["total_blocks"] == 6
    assert parsed_snapshot["counts_by_block_type"]["table_row"] == 3


def test_parse_document_version_normalizes_merged_cells_into_row_truth(session_factory, tmp_path: Path):
    file_path = tmp_path / "table-merged-cells.docx"
    _save_docx_with_merged_table(file_path)
    version_id = _create_document_version(session_factory, file_path)

    with session_factory() as session:
        version = session.get(DocumentVersion, version_id)
        assert version is not None

        parsed_version = document_parser.parse_document_version(session, version)
        active_parse_run_id = parsed_version.active_parse_run_id
        parse_run = session.get(DocumentParseRun, active_parse_run_id)
        table = session.scalar(
            select(DocumentTable).where(DocumentTable.parse_run_id == active_parse_run_id)
        )
        assert table is not None
        rows = list(
            session.scalars(
                select(DocumentTableRow)
                .where(DocumentTableRow.table_id == table.id)
                .order_by(DocumentTableRow.row_index)
            )
        )
        cells = list(
            session.scalars(
                select(DocumentTableCell)
                .join(DocumentTableRow, DocumentTableCell.row_id == DocumentTableRow.id)
                .where(DocumentTableRow.table_id == table.id)
                .order_by(DocumentTableRow.row_index, DocumentTableCell.column_index)
            )
        )
        blocks = list(
            session.scalars(
                select(DocumentBlock)
                .where(DocumentBlock.parse_run_id == active_parse_run_id)
                .order_by(DocumentBlock.order_index)
            )
        )

    assert parse_run is not None
    assert parsed_version.parse_status == "parsed_with_warnings"
    assert parse_run.status == "parsed_with_warnings"
    assert any("merged" in warning.lower() for warning in json.loads(parse_run.summary_json or "{}")["warnings"])

    merged_row_payload = json.loads(rows[2].structured_row_json)
    assert merged_row_payload["cells"] == [
        {
            "column_key": "requirement_id",
            "column_index": 0,
            "header_text": "Requirement ID",
            "normalized_header_text": "Requirement ID",
            "raw_value": "REQ-002 / Register",
            "normalized_value": "REQ-002 / Register",
            "merge_token": None,
            "row_span": 1,
            "col_span": 2,
            "merge_origin_key": None,
        },
        {
            "column_key": "title",
            "column_index": 1,
            "header_text": "Title",
            "normalized_header_text": "Title",
            "raw_value": "",
            "normalized_value": "__MERGED__",
            "merge_token": "__MERGED__",
            "row_span": 1,
            "col_span": 1,
            "merge_origin_key": "tbl-0000-r0002-c0000",
        },
        {
            "column_key": "notes",
            "column_index": 2,
            "header_text": "Notes",
            "normalized_header_text": "Notes",
            "raw_value": "",
            "normalized_value": "__MERGED__",
            "merge_token": "__MERGED__",
            "row_span": 1,
            "col_span": 1,
            "merge_origin_key": "tbl-0000-r0001-c0002",
        },
    ]
    assert json.loads(blocks[3].normalized_content) == [
        {
            "column_key": "requirement_id",
            "normalized_header_text": "Requirement ID",
            "normalized_value": "REQ-002 / Register",
            "merge_token": None,
        },
        {
            "column_key": "title",
            "normalized_header_text": "Title",
            "normalized_value": "__MERGED__",
            "merge_token": "__MERGED__",
        },
        {
            "column_key": "notes",
            "normalized_header_text": "Notes",
            "normalized_value": "__MERGED__",
            "merge_token": "__MERGED__",
        },
    ]
    assert [(cell.row_span, cell.col_span, cell.merge_origin_key) for cell in cells] == [
        (1, 1, None),
        (1, 1, None),
        (1, 1, None),
        (1, 1, None),
        (1, 1, None),
        (2, 1, None),
        (1, 2, None),
        (1, 1, "tbl-0000-r0002-c0000"),
        (1, 1, "tbl-0000-r0001-c0002"),
    ]


def test_parse_document_version_generates_columns_when_table_has_no_header(session_factory, tmp_path: Path):
    file_path = tmp_path / "table-generated-columns.docx"
    _save_docx_with_items(
        file_path,
        [
            {"type": "paragraph", "text": "Examples", "style": "Heading 1"},
            {
                "type": "table",
                "rows": [
                    ["REQ-001", "Login"],
                    ["REQ-002", "Register"],
                ],
            },
        ],
    )
    version_id = _create_document_version(session_factory, file_path)

    with session_factory() as session:
        version = session.get(DocumentVersion, version_id)
        assert version is not None

        parsed_version = document_parser.parse_document_version(session, version)
        active_parse_run_id = parsed_version.active_parse_run_id
        table = session.scalar(
            select(DocumentTable).where(DocumentTable.parse_run_id == active_parse_run_id)
        )
        assert table is not None
        columns = list(
            session.scalars(
                select(DocumentTableColumn)
                .where(DocumentTableColumn.table_id == table.id)
                .order_by(DocumentTableColumn.column_index)
            )
        )
        rows = list(
            session.scalars(
                select(DocumentTableRow)
                .where(DocumentTableRow.table_id == table.id)
                .order_by(DocumentTableRow.row_index)
            )
        )
        blocks = list(
            session.scalars(
                select(DocumentBlock)
                .where(DocumentBlock.parse_run_id == active_parse_run_id)
                .order_by(DocumentBlock.order_index)
            )
        )

    assert table.header_strategy == "generated_columns"
    assert [column.column_key for column in columns] == ["col_1", "col_2"]
    assert [column.header_text for column in columns] == ["col_1", "col_2"]
    assert [column.normalized_header_text for column in columns] == ["col_1", "col_2"]
    assert [column.source_kind for column in columns] == ["generated", "generated"]
    assert [row.is_header_row for row in rows] == [False, False]
    assert [block.block_type for block in blocks] == ["heading", "table_row", "table_row"]
    assert blocks[1].raw_content == "col_1: REQ-001 || col_2: Login"
    assert json.loads(blocks[1].normalized_content) == [
        {
            "column_key": "col_1",
            "normalized_header_text": "col_1",
            "normalized_value": "REQ-001",
            "merge_token": None,
        },
        {
            "column_key": "col_2",
            "normalized_header_text": "col_2",
            "normalized_value": "Login",
            "merge_token": None,
        },
    ]


def test_parse_document_version_deduplicates_colliding_column_keys(session_factory, tmp_path: Path):
    file_path = tmp_path / "table-duplicate-column-keys.docx"
    _save_docx_with_items(
        file_path,
        [
            {"type": "paragraph", "text": "Mappings", "style": "Heading 1"},
            {
                "type": "table",
                "rows": [
                    ["User ID", "User-ID", "User_ID"],
                    ["A1", "B1", "C1"],
                ],
            },
        ],
    )
    version_id = _create_document_version(session_factory, file_path)

    with session_factory() as session:
        version = session.get(DocumentVersion, version_id)
        assert version is not None

        parsed_version = document_parser.parse_document_version(session, version)
        active_parse_run_id = parsed_version.active_parse_run_id
        table = session.scalar(
            select(DocumentTable).where(DocumentTable.parse_run_id == active_parse_run_id)
        )
        assert table is not None
        columns = list(
            session.scalars(
                select(DocumentTableColumn)
                .where(DocumentTableColumn.table_id == table.id)
                .order_by(DocumentTableColumn.column_index)
            )
        )
        row = session.scalar(
            select(DocumentTableRow)
            .where(DocumentTableRow.table_id == table.id)
            .where(DocumentTableRow.row_index == 1)
        )
        assert row is not None
        cells = list(
            session.scalars(
                select(DocumentTableCell)
                .where(DocumentTableCell.row_id == row.id)
                .order_by(DocumentTableCell.column_index)
            )
        )

    assert parsed_version.parse_status == "parsed"
    assert [column.column_key for column in columns] == ["user_id", "user_id_2", "user_id_3"]
    assert [cell.column_id for cell in cells] == [column.id for column in columns]
