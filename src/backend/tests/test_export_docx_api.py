"""Tests for the DOCX export endpoint."""

from io import BytesIO

from docx import Document as DocxDocument


def _build_compare_docx(*, body_paragraphs: list[tuple[str, str | None]]) -> bytes:
    document = DocxDocument()
    for text, style in body_paragraphs:
        paragraph = document.add_paragraph(text)
        if style is not None:
            paragraph.style = style
    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _create_project_and_document(client, auth_headers, *, project_name, document_title):
    project_response = client.post(
        "/api/v1/projects",
        json={"name": project_name, "description": "Export test project"},
        headers=auth_headers,
    )
    assert project_response.status_code == 201
    project_id = project_response.json()["data"]["id"]

    document_response = client.post(
        f"/api/v1/projects/{project_id}/documents",
        json={"title": document_title, "document_type": "SRS", "description": "Test"},
        headers=auth_headers,
    )
    assert document_response.status_code == 201
    return document_response.json()["data"]["id"]


def _upload_and_parse(client, auth_headers, *, document_id, version_label, payload):
    response = client.post(
        f"/api/v1/documents/{document_id}/versions",
        files={
            "file": (
                f"{version_label}.docx",
                payload,
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
        data={"version_label": version_label},
        headers=auth_headers,
    )
    assert response.status_code == 201
    version_id = response.json()["data"]["id"]

    parse_response = client.post(
        f"/api/v1/document-versions/{version_id}/parse",
        headers=auth_headers,
    )
    assert parse_response.status_code == 200
    return version_id


def _create_compare_run(client, auth_headers):
    """Create a full compare run with change items for testing export."""
    document_id = _create_project_and_document(
        client, auth_headers, project_name="Export Project", document_title="Export Spec"
    )

    source_payload = _build_compare_docx(
        body_paragraphs=[
            ("Requirements", "Heading 1"),
            ("The system shall support login.", None),
            ("The system shall write audit logs.", None),
        ]
    )
    target_payload = _build_compare_docx(
        body_paragraphs=[
            ("Requirements", "Heading 1"),
            ("The system shall support secure login with MFA.", None),
            ("The system shall write audit logs.", None),
            ("The system shall enforce session timeout.", None),
        ]
    )

    source_id = _upload_and_parse(
        client, auth_headers, document_id=document_id, version_label="v1.0", payload=source_payload
    )
    target_id = _upload_and_parse(
        client, auth_headers, document_id=document_id, version_label="v2.0", payload=target_payload
    )

    response = client.post(
        f"/api/v1/documents/{document_id}/compare-runs",
        json={"source_version_id": source_id, "target_version_id": target_id},
        headers=auth_headers,
    )
    assert response.status_code == 201
    return response.json()["data"]


def test_export_docx_returns_valid_docx_file(client, auth_headers):
    compare_run = _create_compare_run(client, auth_headers)

    response = client.get(
        f"/api/v1/compare-runs/{compare_run['id']}/export/docx",
        headers=auth_headers,
    )

    assert response.status_code == 200
    assert "application/vnd.openxmlformats-officedocument.wordprocessingml.document" in response.headers["content-type"]
    assert "attachment" in response.headers.get("content-disposition", "")
    assert f"CR-{compare_run['id']:04d}" in response.headers["content-disposition"]

    # Verify it's a valid DOCX by parsing it
    docx_file = DocxDocument(BytesIO(response.content))
    full_text = "\n".join(p.text for p in docx_file.paragraphs)

    assert "Redline Review Report" in full_text
    assert "Export Spec" in full_text
    assert "v1.0" in full_text
    assert "v2.0" in full_text
    assert "Change Summary" in full_text
    assert "Change Items" in full_text


def test_export_docx_includes_summary_text_when_provided(client, auth_headers):
    compare_run = _create_compare_run(client, auth_headers)

    response = client.get(
        f"/api/v1/compare-runs/{compare_run['id']}/export/docx",
        params={"summary_text": "This is a custom executive summary for the report."},
        headers=auth_headers,
    )

    assert response.status_code == 200
    docx_file = DocxDocument(BytesIO(response.content))
    full_text = "\n".join(p.text for p in docx_file.paragraphs)

    assert "Executive Summary" in full_text
    assert "custom executive summary" in full_text


def test_export_docx_accepts_summary_text_in_post_body(client, auth_headers):
    compare_run = _create_compare_run(client, auth_headers)

    response = client.post(
        f"/api/v1/compare-runs/{compare_run['id']}/export/docx",
        json={"summary_text": "This summary is sent in the request body."},
        headers=auth_headers,
    )

    assert response.status_code == 200
    docx_file = DocxDocument(BytesIO(response.content))
    full_text = "\n".join(p.text for p in docx_file.paragraphs)

    assert "Executive Summary" in full_text
    assert "sent in the request body" in full_text


def test_export_docx_requires_authentication(client, auth_headers):
    compare_run = _create_compare_run(client, auth_headers)

    response = client.get(
        f"/api/v1/compare-runs/{compare_run['id']}/export/docx",
    )

    assert response.status_code == 401
