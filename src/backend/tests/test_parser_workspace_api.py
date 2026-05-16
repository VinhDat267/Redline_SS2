import json
from io import BytesIO

import fitz
from docx import Document as DocxDocument
from sqlalchemy import select

from app.models import DocumentParseRun, DocumentSurface, DocumentTable, DocumentVersion
from app.services import document_pdf_parser


def _build_parser_workspace_docx() -> bytes:
    document = DocxDocument()
    document.add_paragraph("Requirements", style="Heading 1")
    document.add_paragraph("Body paragraph for parser workspace.")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Requirement ID"
    table.cell(0, 1).text = "Title"
    table.cell(1, 0).text = "REQ-001"
    table.cell(1, 1).text = "Login"
    document.sections[0].header.add_paragraph("Release Notes")

    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _build_blank_pdf() -> bytes:
    document = fitz.open()
    page = document.new_page(width=612, height=792)
    page.draw_rect(fitz.Rect(72, 72, 220, 220), color=(0, 0, 0), fill=(0, 0, 0))
    payload = document.tobytes()
    document.close()
    return payload


def _create_parsed_version(client, auth_headers, session_factory) -> tuple[int, int]:
    project_response = client.post(
        "/api/v1/projects",
        json={"name": "Parser Workspace Project", "description": "Parser workspace parent"},
        headers=auth_headers,
    )
    project_id = project_response.json()["data"]["id"]
    document_response = client.post(
        f"/api/v1/projects/{project_id}/documents",
        json={
            "title": "Parser Workspace Spec",
            "document_type": "SPEC",
            "description": "Parser workspace target",
        },
        headers=auth_headers,
    )
    document_id = document_response.json()["data"]["id"]
    create_response = client.post(
        f"/api/v1/documents/{document_id}/versions",
        files={
            "file": (
                "parser-workspace.docx",
                _build_parser_workspace_docx(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
        data={"version_label": "v1.0", "notes": "Parser workspace seed"},
        headers=auth_headers,
    )
    version_id = create_response.json()["data"]["id"]
    parse_response = client.post(
        f"/api/v1/document-versions/{version_id}/parse",
        headers=auth_headers,
    )
    assert parse_response.status_code == 200

    with session_factory() as session:
        version = session.get(DocumentVersion, version_id)
        assert version is not None
        assert version.active_parse_run_id is not None

    return document_id, version_id


def test_get_parser_workspace_returns_render_ready_shell(client, auth_headers, session_factory):
    document_id, version_id = _create_parsed_version(client, auth_headers, session_factory)

    response = client.get(
        f"/api/v1/documents/{document_id}/parser-workspace",
        params={"version_id": version_id},
        headers=auth_headers,
    )

    assert response.status_code == 200
    payload = response.json()["data"]

    assert payload["document"]["id"] == document_id
    assert payload["selected_version"]["id"] == version_id
    assert payload["selected_version"]["active_parse_run_id"] is not None
    assert payload["selected_version"]["warning_count"] == 0
    assert payload["selected_version"]["parser_version"] == "v1"
    assert payload["parse_run"]["id"] == payload["selected_version"]["active_parse_run_id"]
    summary = payload["summary"]
    assert {key: summary[key] for key in ("total_surfaces", "total_blocks", "table_count", "row_count", "warning_count")} == {
        "total_surfaces": 2,
        "total_blocks": 5,
        "table_count": 1,
        "row_count": 2,
        "warning_count": 0,
    }
    assert summary["coverage"]["policy_result"] == "pass"
    assert summary["coverage"]["coverage_ratio"] == 1.0
    assert summary["coverage"]["expected_token_count"] > 0
    assert summary["diagnostics"] == []
    assert payload["surface_groups"] == {
        "body": [
            {
                "id": payload["surface_groups"]["body"][0]["id"],
                "surface_key": "body-main",
                "surface_type": "body",
                "label": "Body",
                "item_count": 4,
            }
        ],
        "headers": [
            {
                "id": payload["surface_groups"]["headers"][0]["id"],
                "surface_key": "header-section-1-default",
                "surface_type": "header",
                "label": "Header / section-1",
                "item_count": 1,
            }
        ],
        "footers": [],
        "footnotes": [],
        "endnotes": [],
        "pages": [],
    }
    assert payload["compare_readiness"] == {
        "is_ready": True,
        "status": "ready",
        "message": "Version is parsed and ready for compare setup.",
    }


def test_get_parser_workspace_returns_quality_diagnostics(client, auth_headers, session_factory):
    document_id, version_id = _create_parsed_version(client, auth_headers, session_factory)

    with session_factory() as session:
        version = session.get(DocumentVersion, version_id)
        assert version is not None
        parse_run = session.get(DocumentParseRun, version.active_parse_run_id)
        assert parse_run is not None

        parse_run.warning_count = 1
        parse_run.summary_json = json.dumps(
            {
                "warning_count": 1,
                "warnings": ["DOCX contains text box content that is not part of parser truth."],
                "coverage": {
                    "policy_result": "warn",
                    "canonical_text_length": 124,
                    "secondary_text_length": 156,
                    "expected_token_count": 20,
                    "matched_expected_token_count": 20,
                    "diagnostic_only_token_count": 6,
                    "ignored_token_count": 0,
                    "coverage_ratio": 1.0,
                    "unmatched_text_samples": [],
                },
                "diagnostics": [
                    {
                        "code": "unsupported_textbox",
                        "severity": "warning",
                        "category": "unsupported_content",
                        "policy_impact": "warn",
                        "source_part": "word/document.xml",
                        "source_path": "/w:document[1]/w:body[1]/w:p[3]/w:r[1]/w:pict[1]",
                        "relationship_id": None,
                        "occurrence_key": "unsupported_textbox:word/document.xml:/w:document[1]/w:body[1]/w:p[3]/w:r[1]/w:pict[1]",
                        "surface_type": "body",
                        "message": "DOCX contains text box content that is not part of parser truth.",
                        "count": 1,
                        "text_samples": ["Text box side letter"],
                    }
                ],
            }
        )
        session.add(parse_run)
        session.commit()

    response = client.get(
        f"/api/v1/documents/{document_id}/parser-workspace",
        params={"version_id": version_id},
        headers=auth_headers,
    )

    assert response.status_code == 200
    summary = response.json()["data"]["summary"]

    assert summary["warning_count"] == 1
    assert summary["coverage"] == {
        "policy_result": "warn",
        "canonical_text_length": 124,
        "secondary_text_length": 156,
        "expected_token_count": 20,
        "matched_expected_token_count": 20,
        "diagnostic_only_token_count": 6,
        "ignored_token_count": 0,
        "coverage_ratio": 1.0,
        "unmatched_text_samples": [],
    }
    assert summary["diagnostics"] == [
        {
            "code": "unsupported_textbox",
            "severity": "warning",
            "category": "unsupported_content",
            "policy_impact": "warn",
            "source_part": "word/document.xml",
            "source_path": "/w:document[1]/w:body[1]/w:p[3]/w:r[1]/w:pict[1]",
            "occurrence_key": "unsupported_textbox:word/document.xml:/w:document[1]/w:body[1]/w:p[3]/w:r[1]/w:pict[1]",
            "surface_type": "body",
            "message": "DOCX contains text box content that is not part of parser truth.",
            "count": 1,
            "text_samples": ["Text box side letter"],
        }
    ]


def test_get_parser_workspace_returns_pdf_pages_and_typed_quality_summary(
    monkeypatch,
    client,
    auth_headers,
):
    project_response = client.post(
        "/api/v1/projects",
        json={"name": "PDF Parser Workspace Project", "description": "PDF workspace parent"},
        headers=auth_headers,
    )
    project_id = project_response.json()["data"]["id"]
    document_response = client.post(
        f"/api/v1/projects/{project_id}/documents",
        json={
            "title": "PDF Contract",
            "document_type": "CONTRACT",
            "description": "PDF parser workspace target",
        },
        headers=auth_headers,
    )
    document_id = document_response.json()["data"]["id"]
    create_response = client.post(
        f"/api/v1/documents/{document_id}/versions",
        files={"file": ("contract-scan.pdf", _build_blank_pdf(), "application/pdf")},
        data={"version_label": "v1.0", "notes": "Scanned PDF"},
        headers=auth_headers,
    )
    version_id = create_response.json()["data"]["id"]

    def fake_ocr(page, page_index, settings):
        return document_pdf_parser.OcrPageResult(
            text="1. Payment Terms\nInvoices are due within thirty days.",
            average_confidence=93.0,
            retained_token_count=12,
            low_confidence_token_ratio=0.04,
            languages="eng+vie",
            dpi=200,
        )

    monkeypatch.setattr(document_pdf_parser, "run_ocr_for_page", fake_ocr)

    parse_response = client.post(
        f"/api/v1/document-versions/{version_id}/parse",
        headers=auth_headers,
    )
    assert parse_response.status_code == 200

    response = client.get(
        f"/api/v1/documents/{document_id}/parser-workspace",
        params={"version_id": version_id},
        headers=auth_headers,
    )

    assert response.status_code == 200
    payload = response.json()["data"]
    assert payload["surface_groups"]["pages"] == [
        {
            "id": payload["surface_groups"]["pages"][0]["id"],
            "surface_key": "pdf-page-1",
            "surface_type": "page",
            "label": "Page 1",
            "item_count": 2,
        }
    ]
    assert payload["summary"]["coverage"]["policy_result"] == "warn"
    assert payload["summary"]["pdf"] == {
        "page_count": 1,
        "text_layer_page_count": 0,
        "ocr_page_count": 1,
        "failed_page_count": 0,
        "table_like_page_count": 0,
        "extraction_modes_by_page": {"1": "ocr"},
        "ocr_languages": "eng+vie",
        "average_ocr_confidence": 93.0,
    }
    assert payload["summary"]["diagnostics"][0]["code"] == "pdf_ocr_used"
    assert payload["summary"]["diagnostics"][0]["metadata"]["confidence"] == 93.0


def test_get_parser_surface_detail_returns_render_ready_payload(client, auth_headers, session_factory):
    document_id, version_id = _create_parsed_version(client, auth_headers, session_factory)

    with session_factory() as session:
        version = session.get(DocumentVersion, version_id)
        assert version is not None
        body_surface = session.scalar(
            select(DocumentSurface)
            .where(DocumentSurface.parse_run_id == version.active_parse_run_id)
            .where(DocumentSurface.surface_type == "body")
        )
        body_table = session.scalar(
            select(DocumentTable)
            .where(DocumentTable.parse_run_id == version.active_parse_run_id)
            .where(DocumentTable.surface_id == body_surface.id)
        )

    assert body_surface is not None
    assert body_table is not None

    response = client.get(
        f"/api/v1/document-versions/{version_id}/parser-surfaces/{body_surface.id}",
        headers=auth_headers,
    )

    assert response.status_code == 200
    payload = response.json()["data"]

    assert payload["surface"] == {
        "id": body_surface.id,
        "surface_key": "body-main",
        "surface_type": "body",
        "label": "Body",
        "logical_order_index": 0,
    }
    assert payload["items"] == [
        {
            "kind": "block",
            "block_id": payload["items"][0]["block_id"],
            "block_type": "heading",
            "section_title": "Requirements",
            "surface_order_index": 0,
            "raw_content": "Requirements",
            "normalized_content": "Requirements",
        },
        {
            "kind": "block",
            "block_id": payload["items"][1]["block_id"],
            "block_type": "paragraph",
            "section_title": "Requirements",
            "surface_order_index": 1,
            "raw_content": "Body paragraph for parser workspace.",
            "normalized_content": "Body paragraph for parser workspace.",
        },
        {
            "kind": "table",
            "table_id": body_table.id,
            "table_key": "tbl-0000",
            "surface_order_index": 2,
            "row_count": 2,
        },
    ]
    assert payload["tables"] == [
        {
            "id": body_table.id,
            "table_key": "tbl-0000",
            "header_strategy": "explicit_first_row",
            "section_title": "Requirements",
            "columns": [
                {
                    "column_key": "requirement_id",
                    "column_index": 0,
                    "header_text": "Requirement ID",
                },
                {
                    "column_key": "title",
                    "column_index": 1,
                    "header_text": "Title",
                },
            ],
            "rows": [
                {
                    "row_key": payload["tables"][0]["rows"][0]["row_key"],
                    "row_index": 0,
                    "is_header_row": True,
                    "cells": [
                        {
                            "column_key": "requirement_id",
                            "column_index": 0,
                            "raw_value": "Requirement ID",
                            "normalized_value": "Requirement ID",
                            "merge_origin_key": None,
                            "row_span": 1,
                            "col_span": 1,
                        },
                        {
                            "column_key": "title",
                            "column_index": 1,
                            "raw_value": "Title",
                            "normalized_value": "Title",
                            "merge_origin_key": None,
                            "row_span": 1,
                            "col_span": 1,
                        },
                    ],
                },
                {
                    "row_key": payload["tables"][0]["rows"][1]["row_key"],
                    "row_index": 1,
                    "is_header_row": False,
                    "cells": [
                        {
                            "column_key": "requirement_id",
                            "column_index": 0,
                            "raw_value": "REQ-001",
                            "normalized_value": "REQ-001",
                            "merge_origin_key": None,
                            "row_span": 1,
                            "col_span": 1,
                        },
                        {
                            "column_key": "title",
                            "column_index": 1,
                            "raw_value": "Login",
                            "normalized_value": "Login",
                            "merge_origin_key": None,
                            "row_span": 1,
                            "col_span": 1,
                        },
                    ],
                },
            ],
        }
    ]
