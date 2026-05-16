from dataclasses import dataclass
from io import BytesIO

from docx import Document as DocxDocument


def _build_docx_bytes() -> bytes:
    document = DocxDocument()
    document.add_paragraph("Functional Requirements", style="Heading 1")
    document.add_paragraph("REQ-AUTH-001 The system shall require MFA for administrator login.")
    document.add_paragraph("REQ-AUDIT-002 The system shall retain audit logs for 365 days.")

    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _create_parsed_version(client, auth_headers) -> tuple[int, int, int]:
    project_response = client.post(
        "/api/v1/projects",
        json={"name": "AI Requirement Extraction", "description": "Candidate flow"},
        headers=auth_headers,
    )
    project_id = project_response.json()["data"]["id"]
    document_response = client.post(
        f"/api/v1/projects/{project_id}/documents",
        json={
            "title": "SRS",
            "document_type": "SRS",
            "description": "Requirement extraction target",
        },
        headers=auth_headers,
    )
    document_id = document_response.json()["data"]["id"]

    version_response = client.post(
        f"/api/v1/documents/{document_id}/versions",
        files={
            "file": (
                "srs-requirements.docx",
                _build_docx_bytes(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
        data={"version_label": "v1.0"},
        headers=auth_headers,
    )
    version_id = version_response.json()["data"]["id"]

    parse_response = client.post(
        f"/api/v1/document-versions/{version_id}/parse",
        headers=auth_headers,
    )
    assert parse_response.status_code == 200
    return project_id, document_id, version_id


@dataclass(slots=True)
class FakeExtractionCandidate:
    requirement_code: str
    title: str
    description: str | None
    source_section: str | None
    source_block_key: str | None
    confidence: float | None


@dataclass(slots=True)
class FakeExtractionResult:
    candidates: list[FakeExtractionCandidate]
    provider_used: str
    fallback_used: bool
    error_message: str | None


class FakeRequirementExtractionAdapter:
    def generate_requirement_candidates(self, payload):
        assert payload["document_version_id"] > 0
        assert payload["parse_run_id"] > 0
        assert any("REQ-AUTH-001" in block["content"] for block in payload["blocks"])
        return FakeExtractionResult(
            candidates=[
                FakeExtractionCandidate(
                    requirement_code="REQ-AUTH-001",
                    title="Administrator MFA",
                    description="The system shall require MFA for administrator login.",
                    source_section="Functional Requirements",
                    source_block_key="body-main-block-0001",
                    confidence=0.91,
                ),
                FakeExtractionCandidate(
                    requirement_code="REQ-AUDIT-002",
                    title="Audit log retention",
                    description="The system shall retain audit logs for 365 days.",
                    source_section="Functional Requirements",
                    source_block_key="body-main-block-0002",
                    confidence=0.84,
                ),
            ],
            provider_used="fake",
            fallback_used=False,
            error_message=None,
        )


def test_generate_requirement_candidates_persists_ai_suggestions(client, auth_headers, monkeypatch):
    from app.services import requirement_candidates as candidate_service

    monkeypatch.setattr(candidate_service, "get_llm_adapter", lambda: FakeRequirementExtractionAdapter())
    _, _, version_id = _create_parsed_version(client, auth_headers)

    response = client.post(
        f"/api/v1/document-versions/{version_id}/requirement-candidates/generate",
        json={"force_regenerate": False},
        headers=auth_headers,
    )

    assert response.status_code == 200
    payload = response.json()["data"]
    assert payload["summary"] == {
        "total": 2,
        "pending": 2,
        "accepted": 0,
        "rejected": 0,
    }
    assert payload["provider_used"] == "fake"
    assert payload["fallback_used"] is False
    assert [candidate["requirement_code"] for candidate in payload["candidates"]] == [
        "REQ-AUTH-001",
        "REQ-AUDIT-002",
    ]
    assert payload["candidates"][0]["status"] == "pending"
    assert payload["candidates"][0]["source_block_key"] == "body-main-block-0001"
    assert payload["candidates"][0]["accepted_requirement_id"] is None

    list_response = client.get(
        f"/api/v1/document-versions/{version_id}/requirement-candidates",
        headers=auth_headers,
    )
    assert list_response.status_code == 200
    assert [candidate["requirement_code"] for candidate in list_response.json()["data"]["candidates"]] == [
        "REQ-AUTH-001",
        "REQ-AUDIT-002",
    ]


def test_accept_and_reject_requirement_candidates_update_truth_safely(client, auth_headers, monkeypatch):
    from app.services import requirement_candidates as candidate_service

    monkeypatch.setattr(candidate_service, "get_llm_adapter", lambda: FakeRequirementExtractionAdapter())
    project_id, _, version_id = _create_parsed_version(client, auth_headers)

    generate_response = client.post(
        f"/api/v1/document-versions/{version_id}/requirement-candidates/generate",
        json={"force_regenerate": False},
        headers=auth_headers,
    )
    candidates = generate_response.json()["data"]["candidates"]
    accepted_candidate_id = candidates[0]["id"]
    rejected_candidate_id = candidates[1]["id"]

    accept_response = client.post(
        f"/api/v1/requirement-candidates/{accepted_candidate_id}/accept",
        headers=auth_headers,
    )

    assert accept_response.status_code == 200
    accepted_candidate = accept_response.json()["data"]
    assert accepted_candidate["status"] == "accepted"
    assert accepted_candidate["accepted_requirement_id"] is not None

    duplicate_accept_response = client.post(
        f"/api/v1/requirement-candidates/{accepted_candidate_id}/accept",
        headers=auth_headers,
    )
    assert duplicate_accept_response.status_code == 200
    assert duplicate_accept_response.json()["data"]["accepted_requirement_id"] == accepted_candidate["accepted_requirement_id"]

    reject_response = client.post(
        f"/api/v1/requirement-candidates/{rejected_candidate_id}/reject",
        json={"reason": "Not in final demo scope"},
        headers=auth_headers,
    )
    assert reject_response.status_code == 200
    assert reject_response.json()["data"]["status"] == "rejected"
    assert reject_response.json()["data"]["rejection_reason"] == "Not in final demo scope"

    project_requirements_response = client.get(
        f"/api/v1/projects/{project_id}/requirements",
        headers=auth_headers,
    )
    assert project_requirements_response.status_code == 200
    assert [item["requirement_code"] for item in project_requirements_response.json()["data"]] == [
        "REQ-AUTH-001"
    ]


def test_generate_requirement_candidates_requires_parsed_version(client, auth_headers, monkeypatch):
    from app.services import requirement_candidates as candidate_service

    monkeypatch.setattr(candidate_service, "get_llm_adapter", lambda: FakeRequirementExtractionAdapter())
    project_response = client.post(
        "/api/v1/projects",
        json={"name": "Unparsed Candidate Guard", "description": "No parse yet"},
        headers=auth_headers,
    )
    project_id = project_response.json()["data"]["id"]
    document_response = client.post(
        f"/api/v1/projects/{project_id}/documents",
        json={"title": "Unparsed SRS", "document_type": "SRS"},
        headers=auth_headers,
    )
    document_id = document_response.json()["data"]["id"]
    version_response = client.post(
        f"/api/v1/documents/{document_id}/versions",
        files={
            "file": (
                "unparsed.docx",
                _build_docx_bytes(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
        data={"version_label": "v1.0"},
        headers=auth_headers,
    )
    version_id = version_response.json()["data"]["id"]

    response = client.post(
        f"/api/v1/document-versions/{version_id}/requirement-candidates/generate",
        json={"force_regenerate": False},
        headers=auth_headers,
    )

    assert response.status_code == 409
    assert response.json()["detail"] == "Document version must be parsed before extracting requirements"


def test_reject_accepted_candidate_returns_conflict(client, auth_headers, monkeypatch):
    from app.services import requirement_candidates as candidate_service

    monkeypatch.setattr(candidate_service, "get_llm_adapter", lambda: FakeRequirementExtractionAdapter())
    _, _, version_id = _create_parsed_version(client, auth_headers)

    generate_response = client.post(
        f"/api/v1/document-versions/{version_id}/requirement-candidates/generate",
        json={"force_regenerate": False},
        headers=auth_headers,
    )
    candidate_id = generate_response.json()["data"]["candidates"][0]["id"]

    # Accept first
    accept_response = client.post(
        f"/api/v1/requirement-candidates/{candidate_id}/accept",
        headers=auth_headers,
    )
    assert accept_response.status_code == 200
    assert accept_response.json()["data"]["status"] == "accepted"

    # Reject the already-accepted candidate should fail
    reject_response = client.post(
        f"/api/v1/requirement-candidates/{candidate_id}/reject",
        json={"reason": "Changed my mind"},
        headers=auth_headers,
    )
    assert reject_response.status_code == 409
